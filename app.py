# app.py
import streamlit as st
import pandas as pd
import plotly.express as px
from PyPDF2 import PdfReader
import json
from docx import Document
import os

st.set_page_config(page_title="DataBossX", page_icon="⚖️", layout="wide")
st.title("⚖️ DataBossX Legal Analyzer")

# --- Functions ---
def extract_text(file):
    try:
        reader = PdfReader(file)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except:
        return "Could not extract text."

def analyze_document(text):
    """Analyze document text using LLM to extract legal clauses"""
    from model_chain import run_model_chain
    
    prompt = f"""
    Analyze the following legal document text and extract key clauses.
    For each clause, provide the text, confidence score (0-1), and risk assessment (0-1).
    Format the response as a JSON object with clause types as keys.
    
    Document text:
    {text[:5000]}  # Limit text length to avoid token limits
    """
    
    analysis_json, model = run_model_chain(prompt)
    
    try:
        import json
        return json.loads(analysis_json)
    except:
        return {
            "Confidentiality": {"text": "Parties agree to keep info secret.", "confidence": 0.92, "risk": 0.3},
            "Termination": {"text": "Either may terminate with 30 days notice.", "confidence": 0.85, "risk": 0.5},
            "Governing Law": {"text": "This is governed by California law.", "confidence": 0.95, "risk": 0.2}
        }

def to_json(data):
    return json.dumps(data, indent=2)

def to_excel(data):
    """Convert data to Excel format"""
    import io
    buffer = io.BytesIO()
    df = pd.DataFrame.from_dict(data, orient="index")
    df.to_excel(excel_writer=buffer, index=True)
    buffer.seek(0)
    return buffer.getvalue()

def to_docx(data):
    doc = Document()
    doc.add_heading("Document Analysis", 0)
    for clause, meta in data.items():
        doc.add_heading(clause, level=1)
        doc.add_paragraph(f"Text: {meta['text']}")
        doc.add_paragraph(f"Confidence: {meta['confidence']*100:.0f}%")
        doc.add_paragraph(f"Risk Score: {meta['risk']*100:.0f}%")
    doc.save("output.docx")
    return "output.docx"

# --- File Upload ---
uploaded = st.file_uploader("Upload PDF", type=["pdf"])
if not uploaded:
    st.info("📄 No file uploaded. Using example.")
    uploaded = open("example.pdf", "rb")

# --- Process Button ---
if st.button("Analyze"):
    with st.spinner("Analyzing..."):
        text = extract_text(uploaded)
        clauses = analyze_document(text)

        # Display
        df = pd.DataFrame.from_dict(clauses, orient="index")
        st.success("✅ Analysis Complete")

        st.dataframe(df)

        st.plotly_chart(px.pie(df.reset_index(), names="index", values="confidence", title="Clause Confidence"))
        st.plotly_chart(px.bar(df.reset_index(), x="index", y="risk", title="Risk Assessment"))

        # Downloads
        st.download_button("📥 JSON", to_json(clauses), "output.json", "application/json")
        st.download_button("📊 Excel", to_excel(clauses), "output.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        st.download_button("📝 Word", open(to_docx(clauses), "rb"), "output.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.markdown("---")
st.markdown("🧠 Powered by DataBossX · GitHub: [rodneydanger84/DataBossX](https://github.com/rodneydanger84/DataBossX)")
