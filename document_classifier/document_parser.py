"""
document_parser.py
------------------
Extract plain text from PDF and DOCX files.

Each parser returns a single string so the rest of the pipeline is
format-agnostic.  Table content from DOCX is included so the AI sees
structured data (e.g. invoice line-items) that would otherwise be lost.
"""

import io
import logging
from pathlib import Path

import pdfplumber
from docx import Document

from config import Config

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Dispatch to the correct parser based on file extension.

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename:   Original filename (used only to determine the format).

    Returns:
        The extracted text as a single UTF-8 string.

    Raises:
        ValueError:   Unsupported extension.
        RuntimeError: Extraction failure (wraps the underlying exception).
    """
    ext = Path(filename).suffix.lower()
    if ext not in Config.SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Accepted: {sorted(Config.SUPPORTED_EXTENSIONS)}"
        )

    try:
        if ext == ".pdf":
            return _extract_pdf(file_bytes, filename)
        else:
            return _extract_docx(file_bytes, filename)
    except (ValueError, RuntimeError):
        raise
    except Exception as exc:
        logger.error(
            "Text extraction failed for '%s': %s", filename, exc, exc_info=True
        )
        raise RuntimeError(
            f"Could not extract text from '{filename}': {exc}"
        ) from exc


# ── Private helpers ────────────────────────────────────────────────────────────


def _extract_pdf(file_bytes: bytes, filename: str) -> str:
    """Page-by-page text extraction using pdfplumber."""
    pages: list[str] = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        total = len(pdf.pages)
        logger.debug("PDF '%s' has %d page(s)", filename, total)

        for idx, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            stripped = text.strip()
            if stripped:
                pages.append(stripped)
            else:
                logger.debug(
                    "Page %d/%d of '%s' yielded no text "
                    "(possibly image-based or blank)",
                    idx,
                    total,
                    filename,
                )

    full_text = "\n\n".join(pages)
    logger.info(
        "PDF '%s': extracted %d chars across %d/%d text pages",
        filename,
        len(full_text),
        len(pages),
        total,
    )
    return full_text


def _extract_docx(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from paragraphs and tables in a DOCX file.

    Tables are rendered as pipe-separated rows so structured data
    (line-items, form fields) survives the extraction step.
    """
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables — render each row as "cell1 | cell2 | ..."
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts)
    logger.info(
        "DOCX '%s': extracted %d chars from %d paragraph(s) and %d table(s)",
        filename,
        len(full_text),
        len(doc.paragraphs),
        len(doc.tables),
    )
    return full_text
