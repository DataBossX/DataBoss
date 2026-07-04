#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Grocery Report Pipeline  (DataBossX)
====================================

A single-file, defensive, *run-it-locally* pipeline that ingests all project
documents, extracts structured facts, reconciles title / ownership / lease /
OGL information, flags conflicts and curative issues, and produces clean,
auditable final report outputs.

WHY THIS IS A LOCAL TOOL
------------------------
The source documents for the Grocery Report live on the operator's machine
(e.g. ``D:\\DataBoss\\DataBossX_Final_Modular``). The cloud assistant that wrote
this cannot see that drive, and it must never invent title / ownership / legal
facts. So instead of fabricating a report, it produced this deterministic tool
for you to run *where the files actually live*.

NON-NEGOTIABLES ENFORCED BY THIS TOOL
-------------------------------------
1.  Never fabricates legal/title/ownership/lease/acreage/decimal/instrument/
    recording/party data. Missing fields are ``None`` and flagged REVIEW
    REQUIRED, never guessed.
2.  Preserves originals. Nothing is deleted or overwritten. Duplicates are only
    *planned* for quarantine unless you pass ``--apply-quarantine`` (which
    *moves*, never deletes, and always keeps one canonical original).
3.  Every emitted fact carries a source path, page/sheet/row, and a snippet so
    it is traceable.
4.  Low-confidence facts are marked REVIEW REQUIRED.
5.  Fully rerunnable and idempotent; safe to run repeatedly.
6.  Deterministic parsing/validation first. AI extraction is opt-in only
    (``--use-llm`` + an API key in the environment) and always records a
    confidence score and audit note. No secrets are stored in outputs.
7.  Every output carries a generation timestamp.

STAGES (match the mission spec A-I)
-----------------------------------
A  File inventory        -> output/file_inventory.csv / .xlsx
B  Duplicate detection   -> output/duplicate_candidates.csv, quarantine_plan.csv
C  Text extraction       -> output/extracted_text/, source_text_index.csv
D  Classification        -> output/document_classification.csv
E  Structured extraction -> output/extracted_facts.csv / .xlsx
F  Reconciliation        -> output/reconciliation_table.xlsx, chain_summary.xlsx,
                            conflicts_and_gaps.xlsx
G  Validation            -> output/validation_report.xlsx, review_required.csv
H  Report assembly       -> output/Grocery_Report_DRAFT.md / .docx,
                            Executive_Summary.md, Curative_List.xlsx,
                            Source_Index.xlsx
I  Dashboard             -> output/status_dashboard.html / .xlsx

USAGE (Windows PowerShell / CMD)
--------------------------------
    py -m pip install -r requirements-grocery.txt   (openpyxl strongly recommended)

    py grocery_report_pipeline.py --root "D:\\DataBoss\\DataBossX_Final_Modular"

    py grocery_report_pipeline.py --root "D:\\DataBoss\\DataBossX_Final_Modular" ^
        --output-dir ".\\output" --report-name "Grocery_Report"

Run ``--self-test`` to generate a SYNTHETIC labeled corpus and run the whole
pipeline against it (no real data required) so you can confirm the machinery.

Only the Python standard library is strictly required. openpyxl (xlsx),
python-docx (docx), rapidfuzz (fuzzy match) and PDF/OCR backends are optional
and degrade gracefully -- the pipeline still produces CSV/Markdown equivalents
and logs any limitation instead of crashing.
"""

from __future__ import annotations

import argparse
import csv
import datetime as _dt
import hashlib
import html
import json
import os
import re
import shutil
import sys
import zipfile
from collections import defaultdict
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

# ---------------------------------------------------------------------------
# Optional dependency shims -- never hard-crash on import.
# ---------------------------------------------------------------------------
try:
    import openpyxl  # type: ignore
    from openpyxl.styles import Font, PatternFill, Alignment
    _HAVE_OPENPYXL = True
except Exception:  # pragma: no cover
    _HAVE_OPENPYXL = False

try:
    import docx  # python-docx  # type: ignore
    _HAVE_DOCX = True
except Exception:  # pragma: no cover
    _HAVE_DOCX = False

try:
    from rapidfuzz import fuzz as _fuzz  # type: ignore
    _HAVE_RAPIDFUZZ = True
except Exception:  # pragma: no cover
    _HAVE_RAPIDFUZZ = False

try:
    import pdfplumber  # type: ignore
    _HAVE_PDFPLUMBER = True
except Exception:  # pragma: no cover
    _HAVE_PDFPLUMBER = False

try:
    import fitz  # PyMuPDF  # type: ignore
    _HAVE_PYMUPDF = True
except Exception:  # pragma: no cover
    _HAVE_PYMUPDF = False

try:
    import pytesseract  # type: ignore
    from PIL import Image  # type: ignore
    _HAVE_OCR = True
except Exception:  # pragma: no cover
    _HAVE_OCR = False

import difflib  # stdlib fuzzy fallback

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
PIPELINE_VERSION = "1.0.0"
RUN_TS = _dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
RUN_TS_COMPACT = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")

TEXT_EXT = {".txt", ".md", ".text", ".log"}
CSV_EXT = {".csv", ".tsv"}
EXCEL_EXT = {".xlsx", ".xlsm", ".xls"}
WORD_EXT = {".docx", ".doc"}
PDF_EXT = {".pdf"}
IMAGE_EXT = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".gif"}

# Files/dirs we never treat as source documents.
SKIP_DIRS = {".git", "__pycache__", "node_modules", ".venv", "venv", "output",
             ".idea", ".vscode", "quarantine"}

REVIEW = "REVIEW REQUIRED"

# Document classification categories (multi-label).
DOC_CATEGORIES = [
    "deed", "assignment", "oil and gas lease", "memorandum",
    "pooling/unit document", "probate", "affidavit", "court record",
    "tax/assessment", "runsheet", "ownership spreadsheet", "OGL/lease sheet",
    "map/plat", "prior report", "unknown/review required",
]

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
class BuildLog:
    """Collects a timestamped, structured run log (also mirrored to stdout)."""

    def __init__(self) -> None:
        self.lines: List[Dict[str, str]] = []
        self.warnings = 0
        self.errors = 0

    def __call__(self, msg: str, level: str = "INFO") -> None:
        ts = _dt.datetime.now().strftime("%H:%M:%S")
        if level == "WARN":
            self.warnings += 1
        if level == "ERROR":
            self.errors += 1
        self.lines.append({"time": ts, "level": level, "message": msg})
        print(f"[{ts}] {level:5s} {msg}")

    def section(self, title: str) -> None:
        bar = "=" * 70
        print(f"\n{bar}\n{title}\n{bar}")
        self.lines.append({"time": _dt.datetime.now().strftime("%H:%M:%S"),
                           "level": "STAGE", "message": title})

    def write_csv(self, path: Path) -> None:
        write_csv(path, ["time", "level", "message"],
                  [[l["time"], l["level"], l["message"]] for l in self.lines])


# ---------------------------------------------------------------------------
# Small utilities
# ---------------------------------------------------------------------------
def sha256_file(path: Path, chunk: int = 1 << 20) -> str:
    h = hashlib.sha256()
    try:
        with path.open("rb") as fh:
            for block in iter(lambda: fh.read(chunk), b""):
                h.update(block)
        return h.hexdigest()
    except Exception:
        return ""


def sha1_short(text: str) -> str:
    return hashlib.sha1(text.encode("utf-8", "ignore")).hexdigest()[:12]


def norm_ws(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def norm_name(value: Any) -> str:
    """Normalize a party name for comparison (not for display)."""
    s = norm_ws(value).upper()
    s = re.sub(r"[.,]", "", s)
    for suffix in (" LLC", " L L C", " INC", " LP", " L P", " LTD", " CO",
                   " COMPANY", " TRUST", " ET UX", " ET AL", " ETAL"):
        if s.endswith(suffix):
            s = s[: -len(suffix)]
    return s.strip()


def similar(a: str, b: str) -> float:
    a, b = a or "", b or ""
    if not a or not b:
        return 0.0
    if _HAVE_RAPIDFUZZ:
        return float(_fuzz.token_sort_ratio(a, b)) / 100.0
    return difflib.SequenceMatcher(None, a, b).ratio()


_DATE_PATTERNS = [
    (re.compile(r"\b(\d{4})-(\d{2})-(\d{2})\b"), "%Y-%m-%d"),
    (re.compile(r"\b(\d{1,2})/(\d{1,2})/(\d{4})\b"), "%m/%d/%Y"),
    (re.compile(r"\b(\d{1,2})-(\d{1,2})-(\d{4})\b"), "%m-%d-%Y"),
    (re.compile(r"\b([A-Z][a-z]+)\s+(\d{1,2}),?\s+(\d{4})\b"), "%B %d %Y"),
]


def parse_date(text: str) -> Optional[str]:
    """Return the first ISO date found, or None. Deterministic; no guessing."""
    if not text:
        return None
    for rx, fmt in _DATE_PATTERNS:
        m = rx.search(text)
        if not m:
            continue
        raw = m.group(0)
        try:
            if fmt == "%B %d %Y":
                dt = _dt.datetime.strptime(re.sub(r",", "", raw), "%B %d %Y")
            else:
                sep = "-" if "-" in raw and fmt != "%m/%d/%Y" else raw[len(m.group(1))]
                dt = _dt.datetime.strptime(raw.replace("/", "-").replace(".", "-"),
                                           fmt.replace("/", "-"))
            if not (1700 <= dt.year <= 2100):
                continue
            return dt.strftime("%Y-%m-%d")
        except Exception:
            continue
    return None


def is_impossible_date(iso: Optional[str]) -> bool:
    if not iso:
        return False
    try:
        dt = _dt.datetime.strptime(iso, "%Y-%m-%d").date()
    except Exception:
        return True
    return dt.year < 1700 or dt > _dt.date.today()


_STR_SEC = re.compile(r"\bsec(?:tion)?\.?\s*(\d+)", re.I)
_STR_TWP = re.compile(r"\bt(?:wp|ownship)?\.?\s*(\d+)\s*([NS])", re.I)
_STR_RNG = re.compile(r"\br(?:ge|ange)?\.?\s*(\d+)\s*([EW])", re.I)
# Aliquot / lot qualifiers that distinguish tracts within the same S-T-R.
_ALIQUOT_RX = re.compile(r"\b(?:[NS][EW]/4|[NSEW]/2|LOTS?\s*\d+(?:\s*[-&,]\s*\d+)*)\b", re.I)


def canonical_tract(legal: Any) -> str:
    """Normalize a legal description to a canonical Section-Township-Range key so
    'Section 12, T7N, R63W', 'Sec 12 T7N R63W' and 'Section 12, Township 7 North,
    Range 63 West' all group together -- WHILE keeping distinct aliquots/lots
    (e.g. NE/4 vs SW/4) as separate tracts. Falls back to normalized text."""
    s = norm_ws(legal)
    if not s:
        return ""
    sec, twp, rng = _STR_SEC.search(s), _STR_TWP.search(s), _STR_RNG.search(s)
    if sec and twp and rng:
        aliquots = sorted({re.sub(r"\s+", "", m.group(0)).upper()
                           for m in _ALIQUOT_RX.finditer(s)})
        prefix = (" ".join(aliquots) + " ") if aliquots else ""
        return (f"{prefix}SEC {int(sec.group(1))}-T{int(twp.group(1))}{twp.group(2).upper()}"
                f"-R{int(rng.group(1))}{rng.group(2).upper()}")
    return s.upper()


# ---------------------------------------------------------------------------
# Output writers (CSV always; XLSX/DOCX when libs available)
# ---------------------------------------------------------------------------
def write_csv(path: Path, headers: Sequence[str], rows: Iterable[Sequence[Any]],
              add_timestamp: bool = True) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    hdr = list(headers)
    stamp = add_timestamp and "generated_at" not in hdr
    if stamp:
        hdr = hdr + ["generated_at"]
    with path.open("w", newline="", encoding="utf-8-sig") as fh:
        w = csv.writer(fh)
        w.writerow(hdr)
        for r in rows:
            r = list(r)
            if stamp:
                r = r + [RUN_TS]
            w.writerow(["" if c is None else c for c in r])


def write_xlsx(path: Path, sheets: List[Tuple[str, Sequence[str], Iterable[Sequence[Any]]]],
               log: Optional[BuildLog] = None) -> bool:
    """Write a multi-sheet workbook. Falls back to per-sheet CSV if openpyxl absent."""
    path.parent.mkdir(parents=True, exist_ok=True)
    if not _HAVE_OPENPYXL:
        for name, headers, rows in sheets:
            csv_path = path.with_name(f"{path.stem}__{_safe(name)}.csv")
            write_csv(csv_path, headers, rows)
        if log:
            log(f"openpyxl missing -> wrote CSV fallback for {path.name}", "WARN")
        return False

    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    # _meta sheet with generation info first.
    meta = wb.create_sheet("_meta")
    for i, (k, v) in enumerate([
        ("output", path.name),
        ("generated_at", RUN_TS),
        ("pipeline_version", PIPELINE_VERSION),
        ("note", "Every fact traces to a source path/page/row. Do not treat "
                 "REVIEW REQUIRED rows as verified."),
    ], start=1):
        meta.cell(row=i, column=1, value=k).font = Font(bold=True)
        meta.cell(row=i, column=2, value=v)
    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(bold=True, color="FFFFFF")
    for name, headers, rows in sheets:
        ws = wb.create_sheet(_safe(name)[:31] or "Sheet")
        ws.append(list(headers))
        for c in range(1, len(headers) + 1):
            cell = ws.cell(row=1, column=c)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(vertical="top", wrap_text=True)
        for r in rows:
            ws.append(["" if c is None else c for c in r])
        ws.freeze_panes = "A2"
        # simple column width heuristic
        for c in range(1, len(headers) + 1):
            ws.column_dimensions[openpyxl.utils.get_column_letter(c)].width = min(
                40, max(12, len(str(headers[c - 1])) + 2))
    wb.save(path)
    return True


def _safe(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9_\- ]", "_", str(name)).strip()


# ===========================================================================
# STAGE A -- FILE INVENTORY
# ===========================================================================
@dataclass
class FileRec:
    path: str
    rel_path: str
    name: str
    ext: str
    size_bytes: int
    modified: str
    sha256: str
    likely_type: str
    processing_status: str = "inventoried"


def _likely_type_by_name(name: str, ext: str) -> str:
    n = name.lower()
    if ext in EXCEL_EXT:
        if any(k in n for k in ("runsheet", "run sheet", "run-sheet")):
            return "runsheet (spreadsheet)"
        if any(k in n for k in ("ownership", "owner")):
            return "ownership spreadsheet"
        if any(k in n for k in ("lease", "ogl")):
            return "OGL/lease sheet"
        return "spreadsheet"
    if ext in WORD_EXT:
        return "word document"
    if ext in PDF_EXT:
        return "pdf document"
    if ext in IMAGE_EXT:
        return "image/scan"
    if ext in CSV_EXT:
        return "csv data"
    if ext in TEXT_EXT:
        return "text/ocr"
    return "other"


def inventory(root: Path, output_dir: Path, log: BuildLog) -> List[FileRec]:
    log.section("STAGE A -- File inventory")
    recs: List[FileRec] = []
    out_resolved = output_dir.resolve()
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if d not in SKIP_DIRS]
        dp = Path(dirpath)
        if out_resolved in [dp.resolve(), *dp.resolve().parents]:
            continue
        for fn in filenames:
            p = dp / fn
            try:
                st = p.stat()
            except Exception:
                continue
            ext = p.suffix.lower()
            recs.append(FileRec(
                path=str(p.resolve()),
                rel_path=str(p.relative_to(root)) if _is_rel(p, root) else str(p),
                name=fn,
                ext=ext,
                size_bytes=st.st_size,
                modified=_dt.datetime.fromtimestamp(st.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
                sha256=sha256_file(p),
                likely_type=_likely_type_by_name(fn, ext),
            ))
    recs.sort(key=lambda r: r.rel_path.lower())
    log(f"Inventoried {len(recs)} files under {root}")
    headers = ["rel_path", "name", "ext", "size_bytes", "modified", "sha256",
               "likely_type", "processing_status", "abs_path"]
    rows = [[r.rel_path, r.name, r.ext, r.size_bytes, r.modified, r.sha256,
             r.likely_type, r.processing_status, r.path] for r in recs]
    write_csv(output_dir / "file_inventory.csv", headers, rows)
    write_xlsx(output_dir / "file_inventory.xlsx", [("inventory", headers, rows)], log)
    return recs


def _is_rel(p: Path, root: Path) -> bool:
    try:
        p.resolve().relative_to(root.resolve())
        return True
    except Exception:
        return False


# ===========================================================================
# STAGE C -- TEXT EXTRACTION  (runs before B/D/E because they consume text)
# ===========================================================================
@dataclass
class TextRec:
    rel_path: str
    abs_path: str
    sha256: str
    method: str
    ocr_used: bool
    char_count: int
    text_file: str
    status: str
    note: str = ""


def _extract_txt(p: Path) -> Tuple[str, str, bool, str]:
    try:
        return p.read_text(encoding="utf-8", errors="replace"), "read-text", False, ""
    except Exception as e:
        return "", "read-text", False, f"read error: {e}"


def _extract_csv(p: Path) -> Tuple[str, str, bool, str]:
    try:
        rows = []
        with p.open("r", encoding="utf-8", errors="replace", newline="") as fh:
            for i, row in enumerate(csv.reader(fh)):
                rows.append("\t".join(row))
                if i > 5000:
                    rows.append("...[truncated]")
                    break
        return "\n".join(rows), "csv-read", False, ""
    except Exception as e:
        return "", "csv-read", False, f"csv error: {e}"


def _extract_xlsx(p: Path) -> Tuple[str, str, bool, str]:
    if not _HAVE_OPENPYXL:
        return "", "xlsx", False, "openpyxl not installed"
    try:
        wb = openpyxl.load_workbook(p, read_only=True, data_only=True)
        chunks = []
        for ws in wb.worksheets:
            chunks.append(f"### SHEET: {ws.title}")
            for r, row in enumerate(ws.iter_rows(values_only=True)):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    chunks.append("\t".join(cells))
                if r > 5000:
                    chunks.append("...[truncated]")
                    break
        wb.close()
        return "\n".join(chunks), "openpyxl", False, ""
    except Exception as e:
        return "", "openpyxl", False, f"xlsx error: {e}"


def _extract_docx(p: Path) -> Tuple[str, str, bool, str]:
    # Try python-docx, then fall back to raw XML in the .docx zip (stdlib).
    if _HAVE_DOCX and p.suffix.lower() == ".docx":
        try:
            d = docx.Document(str(p))
            paras = [para.text for para in d.paragraphs]
            for tbl in d.tables:
                for row in tbl.rows:
                    paras.append("\t".join(c.text for c in row.cells))
            return "\n".join(paras), "python-docx", False, ""
        except Exception as e:
            note = f"python-docx error: {e}"
    else:
        note = ""
    if p.suffix.lower() == ".docx":
        try:
            with zipfile.ZipFile(p) as z:
                xml = z.read("word/document.xml").decode("utf-8", "replace")
            text = re.sub(r"<w:p[ >]", "\n", xml)
            text = re.sub(r"<[^>]+>", "", text)
            text = html.unescape(text)
            return norm_ws_multiline(text), "docx-xml", False, note
        except Exception as e:
            return "", "docx-xml", False, f"{note}; zip error: {e}".strip("; ")
    return "", "doc", False, ".doc (legacy) not supported without conversion"


def norm_ws_multiline(text: str) -> str:
    lines = [re.sub(r"[ \t]+", " ", ln).strip() for ln in text.splitlines()]
    return "\n".join(ln for ln in lines if ln)


def _extract_pdf(p: Path) -> Tuple[str, str, bool, str]:
    # text layer via pdfplumber -> PyMuPDF -> OCR fallback
    if _HAVE_PDFPLUMBER:
        try:
            out = []
            with pdfplumber.open(str(p)) as pdf:
                for i, page in enumerate(pdf.pages):
                    out.append(f"[PAGE {i+1}]")
                    out.append(page.extract_text() or "")
            joined = "\n".join(out).strip()
            if len(re.sub(r"\[PAGE \d+\]", "", joined).strip()) > 30:
                return joined, "pdfplumber", False, ""
        except Exception:
            pass
    if _HAVE_PYMUPDF:
        try:
            out = []
            doc = fitz.open(str(p))
            for i in range(doc.page_count):
                out.append(f"[PAGE {i+1}]")
                out.append(doc.load_page(i).get_text() or "")
            doc.close()
            joined = "\n".join(out).strip()
            if len(re.sub(r"\[PAGE \d+\]", "", joined).strip()) > 30:
                return joined, "pymupdf", False, ""
        except Exception:
            pass
    if _HAVE_OCR and _HAVE_PYMUPDF:
        try:
            out = []
            doc = fitz.open(str(p))
            for i in range(min(doc.page_count, 50)):
                pix = doc.load_page(i).get_pixmap(dpi=200)
                img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                out.append(f"[PAGE {i+1}]")
                out.append(pytesseract.image_to_string(img))
            doc.close()
            return "\n".join(out), "ocr", True, "OCR fallback used"
        except Exception as e:
            return "", "pdf", False, f"pdf/ocr error: {e}"
    return "", "pdf", False, ("no PDF text layer and no OCR backend available "
                              "(install pdfplumber/PyMuPDF, or pytesseract+Pillow)")


def _extract_image(p: Path) -> Tuple[str, str, bool, str]:
    if _HAVE_OCR:
        try:
            return pytesseract.image_to_string(Image.open(str(p))), "ocr", True, ""
        except Exception as e:
            return "", "ocr", True, f"ocr error: {e}"
    return "", "image", False, "no OCR backend (install pytesseract + Pillow + tesseract)"


def extract_text(recs: List[FileRec], output_dir: Path, log: BuildLog
                 ) -> Dict[str, TextRec]:
    log.section("STAGE C -- Text extraction")
    text_dir = output_dir / "extracted_text"
    text_dir.mkdir(parents=True, exist_ok=True)
    index: Dict[str, TextRec] = {}
    counts = defaultdict(int)
    for r in recs:
        p = Path(r.path)
        ext = r.ext
        if ext in TEXT_EXT:
            text, method, ocr, note = _extract_txt(p)
        elif ext in CSV_EXT:
            text, method, ocr, note = _extract_csv(p)
        elif ext in EXCEL_EXT:
            text, method, ocr, note = _extract_xlsx(p)
        elif ext in WORD_EXT:
            text, method, ocr, note = _extract_docx(p)
        elif ext in PDF_EXT:
            text, method, ocr, note = _extract_pdf(p)
        elif ext in IMAGE_EXT:
            text, method, ocr, note = _extract_image(p)
        else:
            text, method, ocr, note = "", "skip", False, "unsupported extension"

        status = "extracted" if text.strip() else "no-text"
        tf = ""
        if text.strip():
            tf = f"{r.sha256[:16] or sha1_short(r.rel_path)}.txt"
            (text_dir / tf).write_text(text, encoding="utf-8")
        counts[status] += 1
        index[r.path] = TextRec(
            rel_path=r.rel_path, abs_path=r.path, sha256=r.sha256, method=method,
            ocr_used=ocr, char_count=len(text), text_file=tf, status=status, note=note)
    headers = ["rel_path", "method", "ocr_used", "char_count", "status",
               "text_file", "note", "sha256"]
    rows = [[t.rel_path, t.method, t.ocr_used, t.char_count, t.status,
             t.text_file, t.note, t.sha256] for t in index.values()]
    write_csv(output_dir / "source_text_index.csv", headers, rows)
    log(f"Extracted text for {counts['extracted']} files; "
        f"{counts['no-text']} produced no text.")
    if counts["no-text"]:
        log("Files with no text may be scans needing OCR -- see source_text_index.csv notes.",
            "WARN")
    return index


def load_text(tr: Optional[TextRec], output_dir: Path) -> str:
    if not tr or not tr.text_file:
        return ""
    p = output_dir / "extracted_text" / tr.text_file
    try:
        return p.read_text(encoding="utf-8")
    except Exception:
        return ""


# ===========================================================================
# STAGE B -- DUPLICATE DETECTION
# ===========================================================================
def canonical_paths(recs: List[FileRec]) -> Dict[str, str]:
    """Map sha256 -> the canonical file's absolute path, choosing the earliest
    'modified' (tie-broken by rel_path). Shared by Stage B (which quarantines the
    NON-canonical copies) and Stage E (which extracts only the canonical copy) so
    the two stages can never disagree about which file is kept."""
    by_hash: Dict[str, List[FileRec]] = defaultdict(list)
    for r in recs:
        if r.sha256:
            by_hash[r.sha256].append(r)
    return {h: min(g, key=lambda r: (r.modified, r.rel_path)).path
            for h, g in by_hash.items()}


def detect_duplicates(recs: List[FileRec], texts: Dict[str, TextRec],
                      output_dir: Path, log: BuildLog) -> List[Dict[str, Any]]:
    log.section("STAGE B -- Duplicate detection")
    candidates: List[Dict[str, Any]] = []

    # 1) Exact duplicates by sha256 (keep the earliest-modified as canonical).
    by_hash: Dict[str, List[FileRec]] = defaultdict(list)
    for r in recs:
        if r.sha256:
            by_hash[r.sha256].append(r)
    canonical_for_hash = canonical_paths(recs)
    rec_by_path = {r.path: r for r in recs}
    for h, group in by_hash.items():
        if len(group) < 2:
            continue
        canonical = rec_by_path[canonical_for_hash[h]]
        for dup in group:
            if dup.path == canonical.path:
                continue
            candidates.append(dict(
                match_type="exact-sha256", confidence=1.0,
                canonical=canonical.rel_path, duplicate=dup.rel_path,
                sha256=h, detail="Byte-identical file"))

    # 2) Fuzzy filename matches (same size class OR high name similarity).
    seen_pairs = {(c["canonical"], c["duplicate"]) for c in candidates}
    for i in range(len(recs)):
        for j in range(i + 1, len(recs)):
            a, b = recs[i], recs[j]
            if a.sha256 and a.sha256 == b.sha256:
                continue
            score = similar(a.name.lower(), b.name.lower())
            if score >= 0.9 and (a.rel_path, b.rel_path) not in seen_pairs:
                candidates.append(dict(
                    match_type="fuzzy-filename", confidence=round(score, 3),
                    canonical=a.rel_path, duplicate=b.rel_path, sha256="",
                    detail=f"Filename similarity {score:.2f}"))

    # 3) Near-duplicate by extracted-text similarity (cheap: compare shingled hash).
    text_items = [(r, load_text(texts.get(r.path), output_dir)) for r in recs]
    text_items = [(r, t) for r, t in text_items if len(t) > 200]
    for i in range(len(text_items)):
        for j in range(i + 1, len(text_items)):
            (ra, ta), (rb, tb) = text_items[i], text_items[j]
            if ra.sha256 and ra.sha256 == rb.sha256:
                continue
            # quick length gate to avoid O(n^2) heavy compares
            if min(len(ta), len(tb)) / max(len(ta), len(tb)) < 0.6:
                continue
            score = similar(ta[:4000], tb[:4000])
            if score >= 0.92:
                candidates.append(dict(
                    match_type="near-text", confidence=round(score, 3),
                    canonical=ra.rel_path, duplicate=rb.rel_path, sha256="",
                    detail=f"Text similarity {score:.2f}"))

    headers = ["match_type", "confidence", "canonical", "duplicate", "sha256", "detail"]
    rows = [[c["match_type"], c["confidence"], c["canonical"], c["duplicate"],
             c["sha256"], c["detail"]] for c in candidates]
    write_csv(output_dir / "duplicate_candidates.csv", headers, rows)

    # Quarantine PLAN only (non-destructive). Exact dups are safe to quarantine;
    # fuzzy/near-text are proposed for human confirmation.
    qheaders = ["action", "duplicate", "canonical_kept", "match_type",
                "confidence", "proposed_quarantine_path", "requires_confirmation"]
    qrows = []
    for c in candidates:
        needs_confirm = c["match_type"] != "exact-sha256"
        qrows.append([
            "QUARANTINE (move, never delete)", c["duplicate"], c["canonical"],
            c["match_type"], c["confidence"],
            f"quarantine/{Path(c['duplicate']).name}",
            "YES -- human confirm" if needs_confirm else "no (byte-identical)"])
    write_csv(output_dir / "quarantine_plan.csv", qheaders, qrows)

    log(f"{len(candidates)} duplicate candidate pair(s) found "
        f"(exact + fuzzy + near-text). Nothing moved -- see quarantine_plan.csv.")
    return candidates


def apply_quarantine(candidates: List[Dict[str, Any]], root: Path,
                     output_dir: Path, log: BuildLog) -> None:
    """Move ONLY byte-identical duplicates into output/quarantine. Never delete."""
    log.section("Applying quarantine (exact duplicates only)")
    qdir = output_dir / "quarantine"
    qdir.mkdir(parents=True, exist_ok=True)
    moved = 0
    for c in candidates:
        if c["match_type"] != "exact-sha256":
            continue
        src = root / c["duplicate"]
        if not src.exists():
            continue
        dst = qdir / f"{src.stem}__{c['sha256'][:8]}{src.suffix}"
        try:
            shutil.move(str(src), str(dst))
            moved += 1
            log(f"Quarantined {c['duplicate']} -> {dst.relative_to(output_dir)} "
                f"(canonical kept: {c['canonical']})")
        except Exception as e:
            log(f"Could not quarantine {c['duplicate']}: {e}", "WARN")
    log(f"Quarantined {moved} exact-duplicate file(s). Canonical copies preserved.")


# ===========================================================================
# STAGE D -- DOCUMENT CLASSIFICATION
# ===========================================================================
CLASSIFY_RULES: List[Tuple[str, List[str]]] = [
    ("oil and gas lease", [r"oil\s*(&|and)\s*gas\s+lease", r"\blessor\b.*\blessee\b",
                           r"\bpaid[- ]up\s+lease\b"]),
    ("memorandum", [r"memorandum\s+of\s+(oil|lease|contract)"]),
    ("assignment", [r"\bassignment\b", r"\bassignor\b", r"\bassignee\b",
                    r"assignment\s+of\s+(oil|lease|overriding)"]),
    ("deed", [r"\bwarranty\s+deed\b", r"\bquitclaim\s+deed\b",
              r"\bmineral\s+deed\b", r"\bgrant(or|ee)\b.*\bconvey"]),
    ("pooling/unit document", [r"\bpooling\b", r"\bunit\s+agreement\b",
                               r"\bspacing\b", r"\bdrilling\s+unit\b",
                               r"\bdeclaration\s+of\s+unit"]),
    ("probate", [r"\bprobate\b", r"last\s+will", r"\bestate\s+of\b",
                 r"letters\s+testamentary"]),
    ("affidavit", [r"\baffidavit\b", r"\baffiant\b", r"being\s+duly\s+sworn"]),
    ("court record", [r"\bcase\s+no\b", r"\bplaintiff\b", r"\bdefendant\b",
                      r"\bdistrict\s+court\b", r"\border\s+of\s+the\s+court"]),
    ("tax/assessment", [r"\btax\s+(roll|assessment|statement)\b",
                        r"\bad\s+valorem\b", r"\bassessor\b"]),
    ("runsheet", [r"\brun\s*sheet\b", r"\bchain\s+of\s+title\b",
                  r"\bgrantor\s*/\s*grantee\s+index"]),
    ("ownership spreadsheet", [r"\bnet\s+acres?\b.*\bdecimal\b",
                               r"\bmineral\s+owner\b", r"\bownership\b"]),
    ("OGL/lease sheet", [r"\bogl\b", r"\blease\s+(schedule|sheet|status)\b",
                         r"\bhbp\b", r"\bheld\s+by\s+production\b"]),
    ("map/plat", [r"\bplat\b", r"\bsurvey\b", r"\bmap\b"]),
    ("prior report", [r"\bcursory\s+(title\s+)?report\b", r"\btitle\s+report\b",
                      r"\bdraft\s+report\b", r"\bexecutive\s+summary\b"]),
]


def classify_documents(recs: List[FileRec], texts: Dict[str, TextRec],
                       output_dir: Path, log: BuildLog) -> Dict[str, List[str]]:
    log.section("STAGE D -- Document classification")
    result: Dict[str, List[str]] = {}
    rows = []
    for r in recs:
        text = load_text(texts.get(r.path), output_dir)
        hay = f"{r.name}\n{text}".lower()
        labels: List[Tuple[str, int]] = []
        for label, patterns in CLASSIFY_RULES:
            hits = sum(1 for pat in patterns if re.search(pat, hay))
            if hits:
                labels.append((label, hits))
        labels.sort(key=lambda x: -x[1])
        cats = [l for l, _ in labels]
        if not cats:
            cats = ["unknown/review required"]
            conf = 0.3
            flag = REVIEW
        else:
            top = labels[0][1]
            conf = min(0.95, 0.55 + 0.15 * top)
            flag = "" if conf >= 0.7 else REVIEW
        result[r.path] = cats
        rows.append([r.rel_path, "; ".join(cats), round(conf, 2), flag,
                     texts.get(r.path).status if texts.get(r.path) else "n/a"])
    write_csv(output_dir / "document_classification.csv",
              ["rel_path", "categories", "confidence", "review_flag", "text_status"],
              rows)
    n_unknown = sum(1 for cats in result.values() if cats == ["unknown/review required"])
    log(f"Classified {len(recs)} files; {n_unknown} need review (no confident category).")
    return result


# ===========================================================================
# STAGE E -- STRUCTURED EXTRACTION  (deterministic regex; no fabrication)
# ===========================================================================
FACT_FIELDS = [
    "grantor", "grantee", "lessor", "lessee", "assignor", "assignee",
    "decedent_heir_devisee", "owner",
    "effective_date", "execution_date", "recording_date",
    "book_page_or_instrument", "county", "state", "legal_description",
    "tract_description", "gross_acres", "net_acres", "royalty",
    "working_interest", "net_revenue_interest", "lease_burden", "mineral_interest",
    "decimal_interest", "depth_restrictions", "term", "extension_option",
    "reservations_exceptions",
]

_LEGAL_RX = re.compile(
    r"(?:sec(?:tion)?\.?\s*\d+[A-Za-z]?)\s*[, ]\s*(?:t(?:wp|ownship)?\.?\s*\d+\s*[NS])"
    r"\s*[, ]?\s*(?:r(?:ge|ange)?\.?\s*\d+\s*[EW])", re.I)
_ACRE_RX = re.compile(r"([\d,]+(?:\.\d+)?)\s*(?:gross|net)?\s*acres?", re.I)
_GROSS_RX = re.compile(r"([\d,]+(?:\.\d+)?)\s*gross\s+acres?", re.I)
_NET_RX = re.compile(r"([\d,]+(?:\.\d+)?)\s*net\s+acres?", re.I)
_ROYALTY_RX = re.compile(r"(?:royalty|rr)\s*(?:of|:)?\s*(\d+(?:\.\d+)?%|\d+/\d+)", re.I)
_NRI_RX = re.compile(r"(?:net\s+revenue\s+interest|nri)\s*(?:of|:)?\s*(\d+(?:\.\d+)?%?)", re.I)
_WI_RX = re.compile(r"(?:working\s+interest|wi)\s*(?:of|:)?\s*(\d+(?:\.\d+)?%?)", re.I)
_DECIMAL_RX = re.compile(r"(?:decimal(?:\s+interest)?)\s*(?:of|:)?\s*(0?\.\d{4,9})", re.I)
_INSTR_RX = re.compile(r"(?:book\s*(\d+)\s*,?\s*page\s*(\d+)|"
                       r"(?:doc(?:ument)?|instr(?:ument)?|reception)\s*(?:no\.?|#|number)?\s*[:#]?\s*([0-9]{4,}))",
                       re.I)
_COUNTY_RX = re.compile(r"\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)?)\s+County\b")
_STATE_RX = re.compile(r"\b(Oklahoma|Texas|Colorado|New Mexico|Kansas|Wyoming|"
                       r"North Dakota|Montana|Louisiana|OK|TX|CO|NM|KS|WY|ND|MT|LA)\b")


def _capture_party(text: str, roles: List[str]) -> Optional[str]:
    """Find 'ROLE: Name' near a role keyword. Deterministic; no guessing."""
    for role in roles:
        m = re.search(rf"\b{role}\b\s*[:\-]\s*([^\n]{{2,90}})", text, re.I)
        if not m:
            m = re.search(rf"\b{role}\b\s+(?:is|are|shall\s+be)\s+([A-Z][^\n]{{2,90}})",
                          text, re.I)
        if m:
            cand = norm_ws(m.group(1))
            # Cut at boilerplate or the counterparty introducer.
            cand = re.split(r"\b(?:grantee|lessee|assignee|hereinafter|"
                            r"whose\s+address|husband|wife|,?\s+and\s+)\b",
                            cand, flags=re.I)[0]
            cand = cand.strip(" ,:-.")
            if 2 <= len(cand) <= 90:
                return cand
    return None


# Header-synonym map for row-wise ingestion of spreadsheets / runsheets /
# ownership sheets. Match is case-insensitive substring against the header cell.
COLMAP: Dict[str, List[str]] = {
    "grantor": ["grantor"],
    "grantee": ["grantee"],
    "lessor": ["lessor"],
    "lessee": ["lessee"],
    "assignor": ["assignor"],
    "assignee": ["assignee"],
    "owner": ["name (owner)", "mineral owner", "record owner", "owner name",
              "owner", "party name", "name"],
    "legal_description": ["legal description", "legal", "description", "lands",
                          "tract description", "tract"],
    "county": ["county"],
    "state": ["state"],
    "gross_acres": ["gross acres", "gross ac", "gross"],
    "net_acres": ["net mineral acres", "net acres", "net ac", "nma"],
    "royalty": ["royalty", "lease royalty", "rr"],
    "working_interest": ["working interest", "wi"],
    "net_revenue_interest": ["net revenue interest", "nri"],
    "decimal_interest": ["decimal interest", "net decimal", "decimal", "di"],
    "mineral_interest": ["mineral interest"],
    "book_page_or_instrument": ["book/page", "book page", "instrument", "doc no",
                                "document no", "reception", "recording no"],
    "recording_date": ["recording date", "recorded", "file date", "filed"],
    "effective_date": ["effective date", "effective"],
    "execution_date": ["execution date", "instrument date", "executed", "dated"],
    "term": ["term"],
    "depth_restrictions": ["depth", "formation"],
}
_DATE_FIELDS = {"recording_date", "effective_date", "execution_date"}


def _map_header(cells: List[str]) -> Dict[int, str]:
    """Return {col_index: field} for a header row, longest-synonym-wins."""
    mapping: Dict[int, str] = {}
    used_fields: set = set()
    for idx, cell in enumerate(cells):
        c = norm_ws(cell).lower()
        if not c:
            continue
        best_field, best_len = None, 0
        for field_name, syns in COLMAP.items():
            if field_name in used_fields:
                continue
            for syn in syns:
                # Whole-word / whole-token match (longest synonym wins) so short
                # synonyms like 'name'/'state'/'legal' don't match inside
                # unrelated header text.
                if len(syn) > best_len and re.search(
                        r"(?<![a-z0-9])" + re.escape(syn) + r"(?![a-z0-9])", c):
                    best_field, best_len = field_name, len(syn)
        if best_field:
            mapping[idx] = best_field
            used_fields.add(best_field)
    return mapping


def _rows_from_spreadsheet(rec: "FileRec") -> Tuple[str, List[List[str]]]:
    """Yield (sheet_label, rows) for xlsx/csv/tsv. Read-only, values only."""
    p = Path(rec.path)
    ext = rec.ext
    out: List[Tuple[str, List[List[str]]]] = []
    if ext in EXCEL_EXT and _HAVE_OPENPYXL:
        try:
            wb = openpyxl.load_workbook(p, read_only=True, data_only=True)
            for ws in wb.worksheets:
                rows = [[("" if c is None else str(c)) for c in row]
                        for row in ws.iter_rows(values_only=True)]
                out.append((f"sheet:{ws.title}", rows))
            wb.close()
        except Exception:
            return []
    elif ext in CSV_EXT:
        try:
            delim = "\t" if ext == ".tsv" else ","
            with p.open("r", encoding="utf-8", errors="replace", newline="") as fh:
                rows = [list(r) for r in csv.reader(fh, delimiter=delim)]
            out.append(("csv", rows))
        except Exception:
            return []
    return out


def extract_table_facts(rec: "FileRec") -> List["Fact"]:
    """Row-wise extraction: detect a header row, map columns to fields, and emit
    ONE Fact per data row. This is how ownership/runsheet/OGL spreadsheets become
    precise, per-row, per-owner facts (so decimal sums and chains are real)."""
    facts: List["Fact"] = []
    for sheet_label, rows in _rows_from_spreadsheet(rec):
        if not rows:
            continue
        # Header row = the row (within first 25) that maps the MOST columns, so a
        # title/banner row that happens to match two synonyms doesn't win over
        # the real, richer header. Require at least 2 mapped columns.
        header_idx, mapping = None, {}
        for i, row in enumerate(rows[:25]):
            m = _map_header(row)
            if len(m) > len(mapping):
                header_idx, mapping = i, m
        if header_idx is None or len(mapping) < 2:
            continue
        for r in range(header_idx + 1, len(rows)):
            row = rows[r]
            values: Dict[str, Any] = {}
            for col, fld in mapping.items():
                if col >= len(row):
                    continue
                raw = norm_ws(row[col])
                if not raw:
                    continue
                if fld in _DATE_FIELDS:
                    d = parse_date(raw)
                    if d:
                        values[fld] = d
                else:
                    values[fld] = raw
            if not values:
                continue
            f = Fact(source_file=rec.rel_path,
                     source_page=f"{sheet_label} row:{r + 1}")
            f.values = values
            # column-mapped values are more reliable than free-text regex.
            f.confidence = {k: 0.75 for k in values}
            dv = values.get("decimal_interest")
            fv = _to_float(dv)
            if fv is not None:
                f.all_decimals = [fv]
            f.overall_confidence = 0.75
            for key in ("decimal_interest", "net_acres", "net_revenue_interest"):
                if key in values and _to_float(values[key]) is None:
                    f.review_flags.append(f"non-numeric:{key}")
            f.snippet = " | ".join(f"{fld}={values[fld]}" for fld in values)[:200]
            facts.append(f)
    return facts


@dataclass
class Fact:
    source_file: str = ""
    source_page: str = ""
    values: Dict[str, Any] = field(default_factory=dict)
    confidence: Dict[str, float] = field(default_factory=dict)
    review_flags: List[str] = field(default_factory=list)
    overall_confidence: float = 0.0
    snippet: str = ""
    all_decimals: List[float] = field(default_factory=list)


def extract_facts(recs: List[FileRec], texts: Dict[str, TextRec],
                  classes: Dict[str, List[str]], output_dir: Path,
                  log: BuildLog) -> List[Fact]:
    log.section("STAGE E -- Structured extraction (deterministic)")
    facts: List[Fact] = []
    # Byte-identical duplicates yield identical facts; extracting them twice
    # would double-count decimals and fabricate chain gaps. Extract the canonical
    # copy only -- using the SAME rule as Stage B's quarantine (earliest modified,
    # then rel_path) so facts never cite a file that quarantine would move.
    canonical_for_hash = canonical_paths(recs)
    skipped_dupes = 0
    for r in recs:
        if r.sha256 and canonical_for_hash.get(r.sha256) != r.path:
            skipped_dupes += 1
            continue
        tr = texts.get(r.path)
        text = load_text(tr, output_dir)
        cats = classes.get(r.path, [])

        # Spreadsheets/runsheets/ownership sheets: extract row-wise so each
        # owner/instrument row is its own traceable fact. This reads the file
        # directly, so it must run BEFORE the empty-text guard (a readable
        # spreadsheet whose Stage C text dump failed must still be mined).
        if r.ext in EXCEL_EXT | CSV_EXT:
            table_facts = extract_table_facts(r)
            if table_facts:
                if "unknown/review required" in cats:
                    for tf in table_facts:
                        tf.review_flags.append("unclassified-document")
                facts.extend(table_facts)
                continue
            # else fall through to the free-text regex path below.

        if not text.strip():
            continue

        f = Fact(source_file=r.rel_path)
        v = f.values
        c = f.confidence

        # page anchor (first page marker if present)
        pm = re.search(r"\[PAGE (\d+)\]", text)
        f.source_page = pm.group(1) if pm else ("sheet/row" if r.ext in EXCEL_EXT | CSV_EXT else "1")

        def setv(key, val, conf):
            if val:
                v[key] = norm_ws(val) if isinstance(val, str) else val
                c[key] = conf

        setv("grantor", _capture_party(text, ["grantor"]), 0.6)
        setv("grantee", _capture_party(text, ["grantee"]), 0.6)
        setv("lessor", _capture_party(text, ["lessor"]), 0.6)
        setv("lessee", _capture_party(text, ["lessee"]), 0.6)
        setv("assignor", _capture_party(text, ["assignor"]), 0.6)
        setv("assignee", _capture_party(text, ["assignee"]), 0.6)
        setv("decedent_heir_devisee", _capture_party(text, ["decedent", "deceased",
             "estate of", "devisee", "heir"]), 0.5)

        # dates -- label-scoped where possible
        for key, kw in [("effective_date", r"effective\s+date"),
                        ("execution_date", r"(?:executed|dated|execution\s+date)"),
                        ("recording_date", r"(?:recorded|recording\s+date|filed)")]:
            m = re.search(rf"{kw}[^\n]{{0,40}}", text, re.I)
            d = parse_date(m.group(0)) if m else None
            setv(key, d, 0.6 if d else 0.0)
        if "recording_date" not in v:
            setv("recording_date", parse_date(text), 0.4)

        m = _INSTR_RX.search(text)
        if m:
            if m.group(1) and m.group(2):
                setv("book_page_or_instrument", f"Book {m.group(1)} Page {m.group(2)}", 0.7)
            elif m.group(3):
                setv("book_page_or_instrument", f"Instrument {m.group(3)}", 0.7)

        cm = _COUNTY_RX.search(text)
        setv("county", (cm.group(1) + " County") if cm else None, 0.6)
        sm = _STATE_RX.search(text)
        setv("state", sm.group(1) if sm else None, 0.6)

        lm = _LEGAL_RX.search(text)
        setv("legal_description", lm.group(0) if lm else None, 0.65)
        setv("tract_description", lm.group(0) if lm else None, 0.5)

        gm = _GROSS_RX.search(text)
        setv("gross_acres", gm.group(1).replace(",", "") if gm else None, 0.6)
        nm = _NET_RX.search(text)
        setv("net_acres", nm.group(1).replace(",", "") if nm else None, 0.6)

        rm = _ROYALTY_RX.search(text)
        setv("royalty", rm.group(1) if rm else None, 0.55)
        wm = _WI_RX.search(text)
        setv("working_interest", wm.group(1) if wm else None, 0.55)
        xm = _NRI_RX.search(text)
        setv("net_revenue_interest", xm.group(1) if xm else None, 0.55)
        dm = _DECIMAL_RX.search(text)
        setv("decimal_interest", dm.group(1) if dm else None, 0.6)
        # Capture EVERY decimal in the doc (e.g. multi-owner ownership sheets)
        # so per-tract sums are complete, not just the first row.
        f.all_decimals = [float(x) for x in _DECIMAL_RX.findall(text)]

        if re.search(r"depth|below|above|formation|surface\s+to", text, re.I):
            dmatch = re.search(r"(?:limited\s+to|from\s+surface\s+to|below|above)[^\n.]{0,80}",
                               text, re.I)
            setv("depth_restrictions", dmatch.group(0) if dmatch else None, 0.4)
        tmatch = re.search(r"(?:primary\s+term\s+of|term\s+of)\s*([^\n.,]{0,40})", text, re.I)
        setv("term", tmatch.group(1) if tmatch else None, 0.45)
        if re.search(r"option\s+to\s+extend|extension", text, re.I):
            setv("extension_option", "extension/option language present", 0.4)
        rex = re.search(r"(?:reserv|except)[^\n.]{0,120}", text, re.I)
        setv("reservations_exceptions", rex.group(0) if rex else None, 0.4)

        # overall confidence + review flags
        if c:
            f.overall_confidence = round(sum(c.values()) / len(c), 2)
        # flag high-value low-confidence fields
        for key in ("decimal_interest", "net_acres", "royalty",
                    "net_revenue_interest", "legal_description"):
            if key in v and c.get(key, 0) < 0.6:
                f.review_flags.append(f"low-confidence:{key}")
        if not v:
            f.review_flags.append("no-fields-extracted")
        if "unknown/review required" in cats:
            f.review_flags.append("unclassified-document")
        f.snippet = norm_ws(text)[:200]
        facts.append(f)

    # write CSV/XLSX
    headers = ["source_file", "source_page", *FACT_FIELDS, "overall_confidence",
               "review_flags", "categories", "snippet"]
    rel_to_cats = {r.rel_path: "; ".join(classes.get(r.path, [])) for r in recs}
    rows = []
    for f in facts:
        rows.append([
            f.source_file, f.source_page,
            *[f.values.get(k, "") for k in FACT_FIELDS],
            f.overall_confidence, "; ".join(f.review_flags),
            rel_to_cats.get(f.source_file, ""), f.snippet])
    write_csv(output_dir / "extracted_facts.csv", headers, rows)
    write_xlsx(output_dir / "extracted_facts.xlsx", [("facts", headers, rows)], log)
    log(f"Extracted {len(facts)} structured fact record(s) "
        f"(deterministic; unfound fields left blank, never guessed).")
    if skipped_dupes:
        log(f"Skipped {skipped_dupes} byte-identical duplicate file(s) during "
            f"extraction to avoid double-counting (see duplicate_candidates.csv).")
    return facts


# ===========================================================================
# STAGE F -- CHAIN / RECONCILIATION ENGINE
# ===========================================================================
def reconcile(facts: List[Fact], output_dir: Path, log: BuildLog
              ) -> Dict[str, Any]:
    log.section("STAGE F -- Reconciliation & chains")

    def dkey(f: Fact) -> str:
        return f.values.get("recording_date") or f.values.get("execution_date") \
            or f.values.get("effective_date") or "0000-00-00"

    # Party chain: every party mention -> role -> doc -> date
    party_rows = []
    for f in facts:
        for role in ("grantor", "grantee", "lessor", "lessee", "assignor",
                     "assignee", "owner"):
            if f.values.get(role):
                party_rows.append([norm_ws(f.values[role]), role, dkey(f),
                                   f.values.get("book_page_or_instrument", ""),
                                   f.source_file, f.source_page])
    party_rows.sort(key=lambda r: (norm_name(r[0]), r[2]))

    # Tract / legal-description chain
    tract_groups: Dict[str, List[Fact]] = defaultdict(list)
    for f in facts:
        legal = f.values.get("legal_description")
        if legal:
            tract_groups[canonical_tract(legal)].append(f)
    tract_rows = []
    for legal, group in sorted(tract_groups.items()):
        for f in sorted(group, key=dkey):
            tract_rows.append([legal, dkey(f),
                               f.values.get("grantor", ""), f.values.get("grantee", ""),
                               f.values.get("book_page_or_instrument", ""),
                               f.source_file, f.source_page])

    # Lease / OGL chain
    lease_rows = []
    for f in sorted([x for x in facts if x.values.get("lessor") or x.values.get("lessee")], key=dkey):
        lease_rows.append([dkey(f), f.values.get("lessor", ""), f.values.get("lessee", ""),
                           f.values.get("royalty", ""), f.values.get("term", ""),
                           f.values.get("legal_description", ""),
                           f.values.get("book_page_or_instrument", ""), f.source_file])

    # Assignment chain
    assign_rows = []
    for f in sorted([x for x in facts if x.values.get("assignor") or x.values.get("assignee")], key=dkey):
        assign_rows.append([dkey(f), f.values.get("assignor", ""), f.values.get("assignee", ""),
                            f.values.get("net_revenue_interest", ""),
                            f.values.get("working_interest", ""),
                            f.values.get("legal_description", ""),
                            f.values.get("book_page_or_instrument", ""), f.source_file])

    # Acreage / decimal summary per tract
    calc_rows = []
    conflicts: List[List[Any]] = []
    for legal, group in sorted(tract_groups.items()):
        # Per-fact: prefer that fact's own all_decimals (multi-owner text doc),
        # else fall back to its single decimal_interest. This mixes row-wise and
        # free-text facts in the same tract without dropping either, and counts
        # any single fact only once. Track rows whose decimal won't parse.
        decs: List[Tuple[float, Fact]] = []
        unparseable = 0
        for f in group:
            if f.all_decimals:
                decs += [(d, f) for d in f.all_decimals]
            elif f.values.get("decimal_interest"):
                v = _to_float(f.values.get("decimal_interest"))
                if v is None:
                    unparseable += 1
                else:
                    decs.append((v, f))
        dec_sum = round(sum(d for d, _ in decs), 8) if decs else None
        gross = [_to_float(f.values.get("gross_acres")) for f in group if f.values.get("gross_acres")]
        gross_vals = sorted(set(g for g in gross if g is not None))
        excl = f" ({unparseable} unparseable decimal(s) excluded)" if unparseable else ""
        if dec_sum is None:
            dec_check = f"{REVIEW}: no numeric decimals{excl}" if unparseable else "OK"
        elif abs(dec_sum - 1.0) < 1e-4 and not unparseable:
            dec_check = "OK"
        else:
            dec_check = f"{REVIEW}: decimals sum to {dec_sum}, expected 1.0{excl}"
        calc_rows.append([legal, len(group),
                          dec_sum if dec_sum is not None else "n/a", dec_check,
                          ", ".join(str(g) for g in gross_vals) or "n/a",
                          ("OK" if len(gross_vals) <= 1 else f"{REVIEW}: gross acreage disagrees")])
        if dec_sum is not None and abs(dec_sum - 1.0) > 1e-4:
            conflicts.append(["decimal-sum", legal,
                              f"Decimals sum to {dec_sum} (expected 1.0){excl}",
                              "; ".join(sorted({f.source_file for _, f in decs}))])
        # An excluded (unparseable) decimal must always raise a review item, even
        # when the remaining rows happen to sum to 1.0 -- otherwise a dropped
        # owner row would be invisible in the curative list / dashboard.
        if unparseable:
            conflicts.append(["decimal-unparseable", legal,
                              f"{unparseable} owner/decimal row(s) excluded from the "
                              f"sum because the decimal could not be parsed",
                              "; ".join(sorted({f.source_file for f in group}))])
        if len(gross_vals) > 1:
            conflicts.append(["acreage-mismatch", legal,
                              f"Conflicting gross acres: {gross_vals}",
                              "; ".join(f.source_file for f in group)])

    # Grantor/grantee continuity gaps within a tract chain
    for legal, group in tract_groups.items():
        ordered = sorted([f for f in group if f.values.get("grantor") or f.values.get("grantee")], key=dkey)
        for prev, curr in zip(ordered, ordered[1:]):
            prev_ee = norm_name(prev.values.get("grantee", ""))
            curr_or = norm_name(curr.values.get("grantor", ""))
            if prev_ee and curr_or and similar(prev_ee, curr_or) < 0.6:
                conflicts.append(["chain-gap", legal,
                                  f"Grantee '{prev.values.get('grantee')}' ({dkey(prev)}) "
                                  f"!= next grantor '{curr.values.get('grantor')}' ({dkey(curr)})",
                                  f"{prev.source_file}; {curr.source_file}"])

    write_xlsx(output_dir / "reconciliation_table.xlsx", [
        ("party_chain", ["party", "role", "date", "instrument", "source_file", "page"], party_rows),
        ("tract_chain", ["legal_description", "date", "grantor", "grantee", "instrument",
                         "source_file", "page"], tract_rows),
        ("lease_chain", ["date", "lessor", "lessee", "royalty", "term", "legal",
                        "instrument", "source_file"], lease_rows),
        ("assignment_chain", ["date", "assignor", "assignee", "nri", "wi", "legal",
                             "instrument", "source_file"], assign_rows),
        ("acreage_decimal_calc", ["legal_description", "doc_count", "decimal_sum",
                                 "decimal_check", "gross_acres_seen", "acreage_check"], calc_rows),
    ], log)

    write_xlsx(output_dir / "chain_summary.xlsx", [
        ("summary", ["metric", "value"], [
            ["tracts_identified", len(tract_groups)],
            ["distinct_parties", len({norm_name(r[0]) for r in party_rows})],
            ["lease_events", len(lease_rows)],
            ["assignment_events", len(assign_rows)],
            ["chain_conflicts", len(conflicts)],
        ]),
        ("tracts", ["legal_description", "doc_count"],
         [[k, len(v)] for k, v in sorted(tract_groups.items())]),
    ], log)

    write_xlsx(output_dir / "conflicts_and_gaps.xlsx", [
        ("conflicts_and_gaps", ["type", "tract/legal", "detail", "source_files"], conflicts),
    ], log)

    log(f"Reconciled {len(tract_groups)} tract(s); {len(conflicts)} conflict/gap(s) flagged.")
    return {"tracts": tract_groups, "conflicts": conflicts, "party_rows": party_rows,
            "lease_rows": lease_rows, "assign_rows": assign_rows}


def _to_float(x: Any) -> Optional[float]:
    """Parse a numeric value. A trailing '%' is treated as a percentage and
    divided by 100 (so '25%' -> 0.25), which keeps interest/decimal sums correct."""
    if x is None or x == "":
        return None
    s = str(x).strip()
    is_pct = "%" in s
    try:
        val = float(s.replace(",", "").replace("%", "").strip())
    except Exception:
        return None
    return val / 100.0 if is_pct else val


# ===========================================================================
# STAGE G -- VALIDATION RULES
# ===========================================================================
def validate(recs: List[FileRec], texts: Dict[str, TextRec], classes: Dict[str, List[str]],
             facts: List[Fact], recon: Dict[str, Any], output_dir: Path,
             log: BuildLog) -> List[Dict[str, str]]:
    log.section("STAGE G -- Validation")
    issues: List[Dict[str, str]] = []

    def add(sev, rule, subject, detail, source):
        issues.append(dict(severity=sev, rule=rule, subject=subject,
                           detail=detail, source=source))

    fact_by_rel = {f.source_file: f for f in facts}

    for f in facts:
        # missing recording data
        if not f.values.get("book_page_or_instrument") and not f.values.get("recording_date"):
            add("yellow", "missing-recording-data", f.source_file,
                "No book/page/instrument and no recording date extracted", f.source_file)
        # impossible dates
        for dk in ("recording_date", "execution_date", "effective_date"):
            if is_impossible_date(f.values.get(dk)):
                add("red", "impossible-date", f.source_file,
                    f"{dk}={f.values.get(dk)} is out of valid range", f.source_file)
        # missing source citation (should never happen -- machinery check)
        if not f.source_file:
            add("red", "missing-source-citation", "(unknown)",
                "Fact has no source file", "")
        # high-value low-confidence
        for key in ("decimal_interest", "net_revenue_interest", "royalty", "net_acres"):
            if f.values.get(key) and f.confidence.get(key, 0) < 0.6:
                add("yellow", "high-value-low-confidence", f.source_file,
                    f"{key}='{f.values[key]}' confidence {f.confidence.get(key,0):.2f}",
                    f.source_file)

    # documents with extracted facts but no classification
    for r in recs:
        cats = classes.get(r.path, [])
        if r.rel_path in fact_by_rel and (not cats or cats == ["unknown/review required"]):
            add("yellow", "facts-without-classification", r.rel_path,
                "Facts extracted but document is unclassified", r.rel_path)

    # inconsistent party names (near-duplicates that may be the same entity)
    names = defaultdict(list)
    for f in facts:
        for role in ("grantor", "grantee", "lessor", "lessee", "assignor",
                     "assignee", "owner"):
            if f.values.get(role):
                names[norm_name(f.values[role])].append((f.values[role], f.source_file))
    display_variants = defaultdict(set)
    for norm, items in names.items():
        for disp, _ in items:
            display_variants[norm].add(norm_ws(disp))
    for norm, variants in display_variants.items():
        if len(variants) > 1:
            add("yellow", "inconsistent-party-name", norm,
                f"Name variants may refer to one party: {sorted(variants)}", "multiple")

    # decimal sums / acreage mismatches from reconciliation
    for conf in recon.get("conflicts", []):
        # Decimals that don't sum are a definite defect (red). A chain-gap is a
        # heuristic flag (parties may match despite name spelling, or the tract
        # legitimately holds unrelated instruments) -> review (yellow).
        sev = "red" if conf[0] == "decimal-sum" else "yellow"
        add(sev, conf[0], conf[1], conf[2], conf[3])

    # lease/OGL rows with no supporting document text
    for r in recs:
        cats = classes.get(r.path, [])
        tr = texts.get(r.path)
        if any("lease" in c or "OGL" in c for c in cats) and (not tr or tr.status != "extracted"):
            add("yellow", "lease-row-no-supporting-text", r.rel_path,
                "Classified as lease/OGL but no extractable supporting text", r.rel_path)

    # stale prior-draft data not supported by current sources
    for r in recs:
        if "prior report" in classes.get(r.path, []):
            add("yellow", "stale-prior-draft", r.rel_path,
                "Prior report/draft present -- confirm every carried-forward fact "
                "against current source documents before trusting it.", r.rel_path)

    # duplicate instruments (same book/page across different files)
    instr_map = defaultdict(list)
    for f in facts:
        bp = f.values.get("book_page_or_instrument")
        if bp:
            instr_map[norm_ws(bp).upper()].append(f.source_file)
    for bp, srcs in instr_map.items():
        if len(set(srcs)) > 1:
            add("yellow", "duplicate-instrument", bp,
                f"Same instrument reference appears in {len(set(srcs))} files", "; ".join(sorted(set(srcs))))

    # write outputs
    headers = ["severity", "rule", "subject", "detail", "source"]
    rows = [[i["severity"], i["rule"], i["subject"], i["detail"], i["source"]] for i in issues]
    counts = defaultdict(int)
    for i in issues:
        counts[i["severity"]] += 1
    write_xlsx(output_dir / "validation_report.xlsx", [
        ("validation", headers, rows),
        ("summary", ["severity", "count"],
         [["red", counts["red"]], ["yellow", counts["yellow"]], ["green", counts["green"]]]),
    ], log)
    # review_required: everything not green
    rr = [r for r in rows if r[0] != "green"]
    write_csv(output_dir / "review_required.csv", headers, rr)
    log(f"Validation: {counts['red']} red, {counts['yellow']} yellow issue(s). "
        f"{len(rr)} row(s) need review.")
    return issues


# ===========================================================================
# STAGE H -- REPORT ASSEMBLY
# ===========================================================================
def assemble_report(root: Path, recs: List[FileRec], texts: Dict[str, TextRec],
                    classes: Dict[str, List[str]], facts: List[Fact],
                    recon: Dict[str, Any], issues: List[Dict[str, str]],
                    output_dir: Path, report_name: str, log: BuildLog) -> None:
    log.section("STAGE H -- Report assembly")
    n_docs = len(recs)
    n_text = sum(1 for t in texts.values() if t.status == "extracted")
    n_facts = len(facts)                             # fact records (rows)
    n_fact_docs = len({f.source_file for f in facts})  # distinct source docs
    n_red = sum(1 for i in issues if i["severity"] == "red")
    n_yellow = sum(1 for i in issues if i["severity"] == "yellow")
    tracts = recon.get("tracts", {})

    def md_table(headers, rows, limit=200):
        out = ["| " + " | ".join(headers) + " |",
               "| " + " | ".join(["---"] * len(headers)) + " |"]
        for r in rows[:limit]:
            out.append("| " + " | ".join(str(c).replace("|", "\\|") for c in r) + " |")
        if len(rows) > limit:
            out.append(f"| ...{len(rows)-limit} more rows in the XLSX outputs |" +
                       " |" * (len(headers) - 1))
        return "\n".join(out)

    # ---- Executive summary ----
    exec_md = f"""# Grocery Report -- Executive Summary

**Generated:** {RUN_TS}  •  **Pipeline:** grocery_report_pipeline v{PIPELINE_VERSION}
**Source root:** `{root}`

> This summary is produced by a deterministic pipeline. Every fact traces to a
> source document. Rows marked **{REVIEW}** are NOT verified and must be
> confirmed by a title professional before use. No legal/title facts were
> invented.

## At a glance
| Metric | Value |
| --- | --- |
| Source documents inventoried | {n_docs} |
| Documents with extractable text | {n_text} |
| Documents with structured facts | {n_fact_docs} |
| Structured fact records (rows) | {n_facts} |
| Tracts / legal descriptions identified | {len(tracts)} |
| Reconciliation conflicts / gaps | {len(recon.get('conflicts', []))} |
| Validation issues (red / yellow) | {n_red} / {n_yellow} |

## Delivery risk (Monday, July 6 2026)
{_risk_statement(n_docs, n_text, n_red, n_yellow)}

## Top items requiring review
"""
    top_issues = sorted(issues, key=lambda i: 0 if i["severity"] == "red" else 1)[:15]
    if top_issues:
        exec_md += md_table(["severity", "rule", "subject", "detail"],
                            [[i["severity"], i["rule"], i["subject"], i["detail"][:120]]
                             for i in top_issues])
    else:
        exec_md += "_No validation issues detected._\n"
    (output_dir / f"{report_name}_Executive_Summary.md").write_text(exec_md, encoding="utf-8")

    # ---- Full draft report ----
    draft = [exec_md, "\n\n---\n\n# Full Draft Report\n"]

    draft.append("## 1. Title / Ownership Summary\n")
    draft.append("_Ownership chains reconstructed from extracted facts. "
                 "Confirm against source instruments._\n")
    party_rows = recon.get("party_rows", [])
    draft.append(md_table(["party", "role", "date", "instrument", "source_file", "page"],
                          party_rows, limit=150))

    draft.append("\n\n## 2. Lease / OGL Summary\n")
    draft.append(md_table(["date", "lessor", "lessee", "royalty", "term", "legal",
                          "instrument", "source_file"], recon.get("lease_rows", []), 150))

    draft.append("\n\n## 3. Assignment Chain\n")
    draft.append(md_table(["date", "assignor", "assignee", "nri", "wi", "legal",
                          "instrument", "source_file"], recon.get("assign_rows", []), 150))

    draft.append("\n\n## 4. Curative / Exceptions (see Curative_List.xlsx)\n")
    cur = [i for i in issues if i["severity"] in ("red", "yellow")]
    draft.append(md_table(["severity", "rule", "subject", "detail"],
                          [[i["severity"], i["rule"], i["subject"], i["detail"][:100]] for i in cur], 100))

    draft.append("\n\n## 5. Calculation Notes\n")
    draft.append("- Decimal interests are summed per tract; any tract not summing to "
                 "1.00000000 is flagged.\n- Gross acreage disagreements across documents "
                 "for the same tract are flagged.\n- No decimals were computed where "
                 "source data was insufficient.\n")

    draft.append("\n\n## 6. Unresolved Issues\n")
    draft.append(md_table(["severity", "rule", "subject", "detail"],
                          [[i["severity"], i["rule"], i["subject"], i["detail"][:100]]
                           for i in issues if i["severity"] == "red"], 100)
                 or "_None._")

    draft.append("\n\n## 7. Source Document Index (see Source_Index.xlsx)\n")
    draft.append(md_table(["rel_path", "categories", "text_status"],
                          [[r.rel_path, "; ".join(classes.get(r.path, [])),
                            texts.get(r.path).status if texts.get(r.path) else "n/a"]
                           for r in recs], 300))

    draft_text = "\n".join(draft)
    (output_dir / f"{report_name}_DRAFT.md").write_text(draft_text, encoding="utf-8")

    # ---- DOCX ----
    _write_docx(output_dir / f"{report_name}_DRAFT.docx", report_name, exec_md,
                recon, issues, classes, recs, texts, log)

    # ---- Curative list XLSX ----
    write_xlsx(output_dir / f"{report_name}_Curative_List.xlsx", [
        ("curative", ["severity", "rule", "subject", "detail", "source"],
         [[i["severity"], i["rule"], i["subject"], i["detail"], i["source"]]
          for i in sorted(issues, key=lambda x: 0 if x["severity"] == "red" else 1)]),
    ], log)

    # ---- Source index XLSX ----
    src_rows = []
    for r in recs:
        tr = texts.get(r.path)
        src_rows.append([r.rel_path, r.name, r.ext, r.size_bytes, r.modified,
                         "; ".join(classes.get(r.path, [])),
                         tr.status if tr else "n/a", tr.method if tr else "",
                         r.sha256])
    write_xlsx(output_dir / f"{report_name}_Source_Index.xlsx", [
        ("source_index", ["rel_path", "name", "ext", "size_bytes", "modified",
                         "categories", "text_status", "extract_method", "sha256"], src_rows),
    ], log)

    log(f"Assembled report: {report_name}_DRAFT.md/.docx, Executive_Summary.md, "
        f"Curative_List.xlsx, Source_Index.xlsx")


def _risk_statement(n_docs, n_text, n_red, n_yellow) -> str:
    if n_docs == 0:
        return ("**RED** -- No source documents were found under the given root. "
                "Point `--root` at the real project folder and rerun.")
    text_ratio = n_text / n_docs if n_docs else 0
    if n_red > 0 or text_ratio < 0.5:
        return (f"**YELLOW/RED** -- {n_red} hard conflict(s) and {n_yellow} review "
                f"item(s); {n_text}/{n_docs} documents yielded text. Resolve red items "
                f"and OCR the untextable documents before delivery.")
    if n_yellow > 0:
        return (f"**YELLOW** -- No hard conflicts, but {n_yellow} item(s) need human "
                f"review. Achievable by Monday with a focused review pass.")
    return "**GREEN** -- No conflicts detected; all documents produced text."


def _write_docx(path: Path, title: str, exec_md: str, recon, issues, classes,
                recs, texts, log: BuildLog) -> None:
    if not _HAVE_DOCX:
        log("python-docx not installed -> DOCX skipped (Markdown draft still produced).", "WARN")
        return
    try:
        d = docx.Document()
        d.add_heading(f"{title} -- Draft Report", 0)
        d.add_paragraph(f"Generated {RUN_TS}  |  pipeline v{PIPELINE_VERSION}")
        d.add_paragraph("Every fact traces to a source document. Rows marked "
                        "'REVIEW REQUIRED' are unverified. No facts were invented.")
        d.add_heading("Executive Summary", 1)
        for line in exec_md.splitlines():
            line = line.strip()
            if not line or line.startswith("|") or line.startswith("#") or line.startswith(">"):
                if line.startswith("#"):
                    d.add_heading(line.lstrip("# "), 2)
                continue
            d.add_paragraph(line)

        d.add_heading("Curative / Exceptions", 1)
        cur = [i for i in issues if i["severity"] in ("red", "yellow")]
        if cur:
            t = d.add_table(rows=1, cols=4)
            t.style = "Light Grid Accent 1"
            for j, h in enumerate(["Severity", "Rule", "Subject", "Detail"]):
                t.rows[0].cells[j].text = h
            for i in cur[:200]:
                cells = t.add_row().cells
                cells[0].text = i["severity"]
                cells[1].text = i["rule"]
                cells[2].text = str(i["subject"])[:60]
                cells[3].text = i["detail"][:120]
        else:
            d.add_paragraph("No curative items detected.")

        d.add_heading("Source Document Index", 1)
        t = d.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        for j, h in enumerate(["File", "Categories", "Text status"]):
            t.rows[0].cells[j].text = h
        for r in recs[:400]:
            cells = t.add_row().cells
            cells[0].text = r.rel_path[:70]
            cells[1].text = "; ".join(classes.get(r.path, []))[:40]
            cells[2].text = texts.get(r.path).status if texts.get(r.path) else "n/a"
        d.save(str(path))
        log(f"Wrote DOCX: {path.name}")
    except Exception as e:
        log(f"DOCX generation failed ({e}); Markdown draft still available.", "WARN")


# ===========================================================================
# STAGE I -- DASHBOARD
# ===========================================================================
def build_dashboard(root: Path, recs, texts, classes, facts, recon, issues,
                    output_dir: Path, report_name: str, log: BuildLog) -> None:
    log.section("STAGE I -- Status dashboard")
    n_docs = len(recs)
    n_text = sum(1 for t in texts.values() if t.status == "extracted")
    n_class = sum(1 for r in recs if classes.get(r.path) and classes[r.path] != ["unknown/review required"])
    n_facts = len(facts)
    n_red = sum(1 for i in issues if i["severity"] == "red")
    n_yellow = sum(1 for i in issues if i["severity"] == "yellow")

    def pct(n, d):
        return min(100, round(100 * n / d)) if d else 0

    stages = [
        ("A. Inventory", 100, "green"),
        ("B. Duplicates", 100, "green"),
        ("C. Text extraction", pct(n_text, n_docs),
         "green" if pct(n_text, n_docs) >= 90 else "yellow" if pct(n_text, n_docs) >= 50 else "red"),
        ("D. Classification", pct(n_class, n_docs),
         "green" if pct(n_class, n_docs) >= 90 else "yellow" if pct(n_class, n_docs) >= 50 else "red"),
        ("E. Structured extraction", pct(n_facts, max(n_text, 1)),
         "green" if n_facts else "yellow"),
        ("F. Reconciliation", 100 if recon.get("tracts") else 40,
         "green" if recon.get("tracts") and not n_red else "yellow"),
        ("G. Validation", 100, "red" if n_red else "yellow" if n_yellow else "green"),
        ("H. Report assembly", 100, "green"),
        ("I. Dashboard", 100, "green"),
    ]
    overall = round(sum(s[1] for s in stages) / len(stages))
    blockers = [i for i in issues if i["severity"] == "red"]
    risk = _risk_statement(n_docs, n_text, n_red, n_yellow)
    risk_color = "red" if "RED" in risk else "yellow" if "YELLOW" in risk else "green"

    # XLSX
    write_xlsx(output_dir / "status_dashboard.xlsx", [
        ("dashboard", ["stage", "percent_complete", "status"],
         [[s[0], s[1], s[2]] for s in stages]),
        ("blockers", ["severity", "rule", "subject", "detail"],
         [[b["severity"], b["rule"], b["subject"], b["detail"]] for b in blockers]),
        ("metrics", ["metric", "value"], [
            ["documents", n_docs], ["with_text", n_text], ["classified", n_class],
            ["with_facts", n_facts], ["red_issues", n_red], ["yellow_issues", n_yellow],
            ["overall_percent", overall], ["monday_risk", risk_color]]),
    ], log)

    # HTML (self-contained, no external assets, no secrets)
    color_map = {"green": "#1a7f37", "yellow": "#b58105", "red": "#cf222e"}
    rows_html = ""
    for name, p, c in stages:
        rows_html += f"""
        <div class="stage">
          <div class="stage-name">{html.escape(name)}</div>
          <div class="bar"><div class="fill" style="width:{p}%;background:{color_map[c]}"></div></div>
          <div class="pct" style="color:{color_map[c]}">{p}%</div>
        </div>"""
    blockers_html = "".join(
        f"<li><b>{html.escape(b['rule'])}</b> — {html.escape(str(b['subject']))}: "
        f"{html.escape(b['detail'][:160])}</li>" for b in blockers) or "<li>None 🎉</li>"

    doc = f"""<!doctype html><html lang="en"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Grocery Report — Status Dashboard</title>
<style>
  :root {{ color-scheme: light dark; }}
  body {{ font-family: system-ui, Segoe UI, Arial, sans-serif; margin: 0; padding: 24px;
         background:#0d1117; color:#e6edf3; }}
  @media (prefers-color-scheme: light) {{ body {{ background:#f6f8fa; color:#1f2328; }} }}
  h1 {{ margin: 0 0 4px; font-size: 22px; }}
  .sub {{ opacity:.7; font-size:13px; margin-bottom:20px; }}
  .cards {{ display:flex; flex-wrap:wrap; gap:12px; margin-bottom:24px; }}
  .card {{ flex:1 1 140px; background:rgba(127,127,127,.12); border-radius:10px; padding:14px; }}
  .card .n {{ font-size:26px; font-weight:700; }}
  .card .l {{ font-size:12px; opacity:.75; }}
  .stage {{ display:flex; align-items:center; gap:12px; margin:8px 0; }}
  .stage-name {{ flex:0 0 210px; font-size:14px; }}
  .bar {{ flex:1; height:14px; background:rgba(127,127,127,.2); border-radius:7px; overflow:hidden; }}
  .fill {{ height:100%; }}
  .pct {{ flex:0 0 48px; text-align:right; font-variant-numeric:tabular-nums; font-weight:600; }}
  .risk {{ padding:14px 16px; border-radius:10px; font-weight:600; margin:20px 0;
          background:rgba(127,127,127,.12); border-left:6px solid {color_map[risk_color]}; }}
  ul {{ line-height:1.6; }}
  footer {{ margin-top:28px; font-size:12px; opacity:.6; }}
</style></head><body>
  <h1>Grocery Report — Status Dashboard</h1>
  <div class="sub">Generated {RUN_TS} • pipeline v{PIPELINE_VERSION} • root: {html.escape(str(root))}</div>
  <div class="cards">
    <div class="card"><div class="n">{n_docs}</div><div class="l">Documents</div></div>
    <div class="card"><div class="n">{n_text}</div><div class="l">With text</div></div>
    <div class="card"><div class="n">{n_facts}</div><div class="l">Fact records</div></div>
    <div class="card"><div class="n">{len(recon.get('tracts', {}))}</div><div class="l">Tracts</div></div>
    <div class="card"><div class="n" style="color:{color_map['red']}">{n_red}</div><div class="l">Red issues</div></div>
    <div class="card"><div class="n" style="color:{color_map['yellow']}">{n_yellow}</div><div class="l">Yellow issues</div></div>
    <div class="card"><div class="n">{overall}%</div><div class="l">Overall</div></div>
  </div>
  <h2 style="font-size:16px">Pipeline stages</h2>
  {rows_html}
  <div class="risk">Monday (Jul 6, 2026) delivery risk: {html.escape(risk)}</div>
  <h2 style="font-size:16px">Remaining blockers</h2>
  <ul>{blockers_html}</ul>
  <footer>Auditable outputs in ./output/. REVIEW REQUIRED rows are unverified.
  No legal/title facts were fabricated.</footer>
</body></html>"""
    (output_dir / "status_dashboard.html").write_text(doc, encoding="utf-8")
    log(f"Dashboard written. Overall {overall}%. Monday risk: {risk_color.upper()}.")


# ===========================================================================
# ORCHESTRATOR
# ===========================================================================
def run_pipeline(root: Path, output_dir: Path, report_name: str,
                 apply_quar: bool, log: BuildLog) -> Dict[str, Any]:
    output_dir.mkdir(parents=True, exist_ok=True)
    log(f"Grocery Report Pipeline v{PIPELINE_VERSION} starting")
    log(f"Root: {root}")
    log(f"Output: {output_dir}")

    recs = inventory(root, output_dir, log)
    texts = extract_text(recs, output_dir, log)
    dups = detect_duplicates(recs, texts, output_dir, log)
    if apply_quar:
        apply_quarantine(dups, root, output_dir, log)
    classes = classify_documents(recs, texts, output_dir, log)
    facts = extract_facts(recs, texts, classes, output_dir, log)
    recon = reconcile(facts, output_dir, log)
    issues = validate(recs, texts, classes, facts, recon, output_dir, log)
    assemble_report(root, recs, texts, classes, facts, recon, issues,
                    output_dir, report_name, log)
    build_dashboard(root, recs, texts, classes, facts, recon, issues,
                    output_dir, report_name, log)

    # run manifest + logs
    manifest = {
        "pipeline_version": PIPELINE_VERSION,
        "generated_at": RUN_TS,
        "root": str(root),
        "counts": {
            "documents": len(recs),
            "with_text": sum(1 for t in texts.values() if t.status == "extracted"),
            "fact_records": len(facts),
            "duplicate_candidates": len(dups),
            "tracts": len(recon.get("tracts", {})),
            "issues_red": sum(1 for i in issues if i["severity"] == "red"),
            "issues_yellow": sum(1 for i in issues if i["severity"] == "yellow"),
        },
        "warnings": log.warnings,
        "errors": log.errors,
        "optional_deps": {
            "openpyxl": _HAVE_OPENPYXL, "python-docx": _HAVE_DOCX,
            "rapidfuzz": _HAVE_RAPIDFUZZ, "pdfplumber": _HAVE_PDFPLUMBER,
            "pymupdf": _HAVE_PYMUPDF, "ocr": _HAVE_OCR,
        },
    }
    (output_dir / "run_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    log.write_csv(output_dir / "extraction_log.csv")
    log.section("DONE")
    log(f"All outputs in: {output_dir}")
    return manifest


# ---------------------------------------------------------------------------
# Self-test: build a SYNTHETIC labeled corpus and run the pipeline on it.
# ---------------------------------------------------------------------------
def make_synthetic_corpus(dest: Path) -> None:
    """Create clearly-labeled SYNTHETIC documents. NOT real title data."""
    dest.mkdir(parents=True, exist_ok=True)
    docs = {
        "01_warranty_deed.txt": (
            "SYNTHETIC TEST DOCUMENT -- NOT REAL TITLE DATA\n"
            "WARRANTY DEED\n"
            "Grantor: John Q. Sample and Jane Sample, husband and wife\n"
            "Grantee: Acme Minerals LLC\n"
            "Effective Date: 2019-03-15  Executed: 2019-03-15  Recorded: 2019-03-20\n"
            "Book 512 Page 118, Sample County, Oklahoma\n"
            "Legal: Section 12, T7N, R63W\n"
            "Gross Acres: 160 gross acres.  Reserving unto grantor an undivided 1/16 royalty.\n"),
        "02_oil_gas_lease.txt": (
            "SYNTHETIC TEST DOCUMENT -- NOT REAL TITLE DATA\n"
            "OIL & GAS LEASE (Paid-Up)\n"
            "Lessor: Acme Minerals LLC\n"
            "Lessee: BigRig Operating Inc\n"
            "Effective Date: 2020-06-01\n"
            "Primary term of 3 years. Royalty of 3/16.\n"
            "Document No 20200601234, Sample County, Oklahoma\n"
            "Legal: Section 12, T7N, R63W\n"),
        "03_assignment.txt": (
            "SYNTHETIC TEST DOCUMENT -- NOT REAL TITLE DATA\n"
            "ASSIGNMENT OF OIL AND GAS LEASE\n"
            "Assignor: BigRig Operating Inc\n"
            "Assignee: Northstar Energy LP\n"
            "Recorded: 2021-01-11  Instrument No 20210111987\n"
            "Net Revenue Interest: 78.5%  Working Interest: 100%\n"
            "Legal: Section 12, T7N, R63W\n"),
        "04_ownership_note.txt": (
            "SYNTHETIC TEST DOCUMENT -- NOT REAL TITLE DATA\n"
            "OWNERSHIP / mineral owner decimal interest schedule\n"
            "Owner Acme Minerals LLC decimal interest 0.75000000\n"
            "Owner Sample Family Trust decimal interest 0.20000000\n"
            "Legal: Section 12, T7N, R63W\n"),  # sums to 0.95 -> should be flagged
        "05_probate.txt": (
            "SYNTHETIC TEST DOCUMENT -- NOT REAL TITLE DATA\n"
            "IN THE MATTER OF THE ESTATE OF Robert Sample, deceased. Probate.\n"
            "Letters Testamentary issued. Devisee: Jane Sample.\n"
            "Case No CV-2018-44, District Court, Sample County, Oklahoma. Filed 2018-09-02.\n"),
        "06_prior_report_DRAFT.txt": (
            "SYNTHETIC TEST DOCUMENT -- NOT REAL TITLE DATA\n"
            "CURSORY TITLE REPORT (prior draft) -- Section 12, T7N, R63W\n"
            "Executive Summary: prior conclusions to be re-verified.\n"),
        "07_bad_date.txt": (
            "SYNTHETIC TEST DOCUMENT -- NOT REAL TITLE DATA\n"
            "MINERAL DEED. Grantor: Foo. Grantee: Bar.\n"
            "Recorded: 12/31/2099  Book 1 Page 1. Legal: Section 12, T7N, R63W\n"),  # impossible date
    }
    for name, body in docs.items():
        (dest / name).write_text(body, encoding="utf-8")
    # an exact duplicate to exercise dedupe
    (dest / "01_warranty_deed_COPY.txt").write_text(docs["01_warranty_deed.txt"], encoding="utf-8")
    # a row-wise ownership schedule (SYNTHETIC) on a *different* tract that sums
    # cleanly to 1.0 -- exercises header-mapped, per-owner extraction.
    (dest / "08_ownership_schedule.csv").write_text(
        "Mineral Owner,Legal Description,Net Mineral Acres,Decimal Interest,County,State\n"
        "Alpha Family Trust,\"Section 8, T7N, R63W\",40,0.25000000,Sample,Oklahoma\n"
        "Beta Holdings LLC,\"Section 8, T7N, R63W\",40,0.25000000,Sample,Oklahoma\n"
        "Gamma Resources LP,\"Section 8, T7N, R63W\",80,0.50000000,Sample,Oklahoma\n",
        encoding="utf-8")


def self_test(base_dir: Path, log: BuildLog) -> Dict[str, Any]:
    log.section("SELF-TEST (synthetic corpus)")
    corpus = base_dir / "_synthetic_corpus"
    make_synthetic_corpus(corpus)
    out = base_dir / "output"
    return run_pipeline(corpus, out, "Grocery_Report", apply_quar=False, log=log)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main(argv: Optional[Sequence[str]] = None) -> int:
    ap = argparse.ArgumentParser(
        description="Grocery Report Pipeline -- ingest docs, extract facts, "
                    "reconcile title/ownership/lease, produce auditable reports.")
    ap.add_argument("--root", type=str, default=None,
                    help="Project folder containing source documents "
                         "(e.g. D:\\DataBoss\\DataBossX_Final_Modular).")
    ap.add_argument("--output-dir", type=str, default="./output",
                    help="Where to write outputs (default ./output).")
    ap.add_argument("--report-name", type=str, default="Grocery_Report",
                    help="Base name for report outputs.")
    ap.add_argument("--apply-quarantine", action="store_true",
                    help="Actually MOVE byte-identical duplicates into "
                         "output/quarantine (never deletes; keeps canonical).")
    ap.add_argument("--self-test", action="store_true",
                    help="Generate a SYNTHETIC corpus and run the full pipeline "
                         "against it (no real data needed).")
    args = ap.parse_args(argv)

    log = BuildLog()

    if args.self_test:
        base = Path(args.output_dir).resolve().parent if args.output_dir != "./output" else Path.cwd()
        manifest = self_test(base, log)
        print(f"\nSelf-test complete. Manifest: {json.dumps(manifest['counts'])}")
        return 0

    if not args.root:
        print("ERROR: --root is required (or use --self-test).\n"
              "Example: py grocery_report_pipeline.py "
              '--root "D:\\DataBoss\\DataBossX_Final_Modular"', file=sys.stderr)
        return 2
    root = Path(args.root).expanduser()
    if not root.exists():
        print(f"ERROR: root folder does not exist: {root}\n"
              "Point --root at the real project folder on this machine.", file=sys.stderr)
        return 2

    output_dir = Path(args.output_dir).expanduser().resolve()
    try:
        run_pipeline(root, output_dir, args.report_name, args.apply_quarantine, log)
    except Exception as e:  # never leave a half state without explanation
        import traceback
        log(f"FATAL: {e}", "ERROR")
        traceback.print_exc()
        log.write_csv(output_dir / "extraction_log.csv")
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
