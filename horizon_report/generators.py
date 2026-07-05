"""Report generators for creating DOCX, XLSX, and PDF outputs."""

from pathlib import Path
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from .models import ReportBundle


class DocxGenerator:
    """Generator for Word document reports."""
    
    def __init__(self, logger):
        """Initialize DOCX generator.
        
        Args:
            logger: Logger instance
        """
        self.logger = logger
    
    def generate(self, bundle: ReportBundle, output_path: Path) -> bool:
        """Generate DOCX report.
        
        Args:
            bundle: Report data bundle
            output_path: Output file path
            
        Returns:
            True if successful
        """
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading(bundle.metadata.report_type, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata section
            self._add_metadata_section(doc, bundle)
            
            # Property information
            if bundle.metadata.section or bundle.metadata.county:
                self._add_property_section(doc, bundle)
            
            # Input files summary
            if bundle.source_documents:
                self._add_sources_section(doc, bundle)
            
            # Ownership data
            self._add_ownership_section(doc, bundle)
            
            # Chain of title
            if bundle.chain_events:
                self._add_chain_section(doc, bundle)
            
            # Missing items and exceptions
            if bundle.missing_items or bundle.exceptions:
                self._add_issues_section(doc, bundle)
            
            # QA summary
            self._add_qa_section(doc, bundle)
            
            # Analyst notes
            if bundle.analyst_notes:
                self._add_notes_section(doc, bundle)
            
            # Save
            doc.save(output_path)
            self.logger.info(f"DOCX report generated: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error generating DOCX: {e}")
            return False
    
    def _add_metadata_section(self, doc: Document, bundle: ReportBundle) -> None:
        """Add metadata section."""
        doc.add_heading('Report Information', 1)
        
        p = doc.add_paragraph()
        p.add_run(f"Report Date: ").bold = True
        p.add_run(bundle.metadata.report_date)
        
        p = doc.add_paragraph()
        p.add_run(f"Generated: ").bold = True
        p.add_run(bundle.metadata.generated_at.strftime("%Y-%m-%d %H:%M:%S"))
        
        if bundle.metadata.input_folder:
            p = doc.add_paragraph()
            p.add_run(f"Input Folder: ").bold = True
            p.add_run(bundle.metadata.input_folder)
        
        doc.add_paragraph()
    
    def _add_property_section(self, doc: Document, bundle: ReportBundle) -> None:
        """Add property information section."""
        doc.add_heading('Property Information', 1)
        
        if bundle.metadata.section:
            p = doc.add_paragraph()
            p.add_run("Section: ").bold = True
            p.add_run(f"{bundle.metadata.section}")
            if bundle.metadata.township:
                p.add_run(f"-{bundle.metadata.township}")
            if bundle.metadata.range:
                p.add_run(f"-{bundle.metadata.range}")
        
        if bundle.metadata.county:
            p = doc.add_paragraph()
            p.add_run("County: ").bold = True
            p.add_run(bundle.metadata.county)
        
        if bundle.metadata.state:
            p = doc.add_paragraph()
            p.add_run("State: ").bold = True
            p.add_run(bundle.metadata.state)
        
        doc.add_paragraph()
    
    def _add_sources_section(self, doc: Document, bundle: ReportBundle) -> None:
        """Add source documents section."""
        doc.add_heading('Source Documents', 1)
        
        doc.add_paragraph(f"Total sources analyzed: {len(bundle.source_documents)}")
        
        for i, source in enumerate(bundle.source_documents[:10], 1):  # Limit to 10
            p = doc.add_paragraph(f"{i}. ", style='List Number')
            p.add_run(f"{source.path.name} ").bold = True
            p.add_run(f"({source.file_type}, {source.relevance})")
        
        if len(bundle.source_documents) > 10:
            doc.add_paragraph(f"... and {len(bundle.source_documents) - 10} more files")
        
        doc.add_paragraph()
    
    def _add_ownership_section(self, doc: Document, bundle: ReportBundle) -> None:
        """Add ownership data section."""
        doc.add_heading('Ownership Summary', 1)
        
        if not bundle.ownership_entries:
            warning = doc.add_paragraph()
            warning.add_run("⚠️ NO VERIFIED OWNERSHIP DATA EXTRACTED — SOURCE REVIEW REQUIRED").bold = True
            doc.add_paragraph()
            return
        
        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        headers = table.rows[0].cells
        headers[0].text = "Owner"
        headers[1].text = "Interest"
        headers[2].text = "Net Acres"
        headers[3].text = "Instrument"
        headers[4].text = "Status"
        
        for cell in headers:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data rows
        for entry in bundle.ownership_entries[:50]:  # Limit to 50
            row = table.add_row().cells
            row[0].text = entry.owner_name or ""
            row[1].text = entry.interest_fraction or (f"{entry.interest_decimal:.6f}" if entry.interest_decimal else "")
            row[2].text = f"{entry.net_acres:.2f}" if entry.net_acres else ""
            row[3].text = entry.source_instrument or ""
            row[4].text = entry.status or ""
        
        if len(bundle.ownership_entries) > 50:
            doc.add_paragraph(f"\nNote: Showing first 50 of {len(bundle.ownership_entries)} entries. See XLSX for complete data.")
        
        doc.add_paragraph()
    
    def _add_chain_section(self, doc: Document, bundle: ReportBundle) -> None:
        """Add chain of title section."""
        doc.add_heading('Chain of Title', 1)
        
        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        headers = table.rows[0].cells
        headers[0].text = "Date"
        headers[1].text = "Instrument"
        headers[2].text = "Grantor → Grantee"
        headers[3].text = "Interest"
        
        for cell in headers:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        for event in bundle.chain_events[:30]:  # Limit to 30
            row = table.add_row().cells
            row[0].text = event.date or ""
            row[1].text = f"{event.instrument_type} {event.instrument_number}"
            row[2].text = f"{event.grantor} → {event.grantee}"
            row[3].text = event.conveyed_interest or ""
        
        doc.add_paragraph()
    
    def _add_issues_section(self, doc: Document, bundle: ReportBundle) -> None:
        """Add missing items and exceptions section."""
        doc.add_heading('Issues and Exceptions', 1)
        
        if bundle.missing_items:
            doc.add_heading('Missing Items', 2)
            for item in bundle.missing_items:
                doc.add_paragraph(f"• {item}", style='List Bullet')
        
        if bundle.exceptions:
            doc.add_heading('Exceptions', 2)
            for exc in bundle.exceptions:
                doc.add_paragraph(f"• {exc}", style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_qa_section(self, doc: Document, bundle: ReportBundle) -> None:
        """Add QA summary section."""
        doc.add_heading('Quality Assurance', 1)
        
        p = doc.add_paragraph()
        p.add_run("Confidence Score: ").bold = True
        score_text = f"{bundle.confidence_score:.1f}/100"
        run = p.add_run(score_text)
        
        # Color code the score
        if bundle.confidence_score >= 80:
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif bundle.confidence_score >= 60:
            run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        if bundle.qa_results:
            doc.add_heading('QA Checks', 2)
            passed = sum(1 for r in bundle.qa_results if r.passed)
            total = len(bundle.qa_results)
            doc.add_paragraph(f"Passed: {passed}/{total} checks")
            
            # Show failed checks
            failed = [r for r in bundle.qa_results if not r.passed and r.severity in ['error', 'critical']]
            if failed:
                doc.add_paragraph("Critical Issues:")
                for result in failed:
                    doc.add_paragraph(f"❌ {result.check_name}: {result.message}", style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_notes_section(self, doc: Document, bundle: ReportBundle) -> None:
        """Add analyst notes section."""
        doc.add_heading('Analyst Notes', 1)
        
        for note in bundle.analyst_notes:
            doc.add_paragraph(f"• {note}", style='List Bullet')


class XlsxGenerator:
    """Generator for Excel workbook reports."""
    
    def __init__(self, logger):
        """Initialize XLSX generator.
        
        Args:
            logger: Logger instance
        """
        self.logger = logger
    
    def generate(self, bundle: ReportBundle, output_path: Path) -> bool:
        """Generate XLSX report.
        
        Args:
            bundle: Report data bundle
            output_path: Output file path
            
        Returns:
            True if successful
        """
        try:
            wb = openpyxl.Workbook()
            wb.remove(wb.active)  # Remove default sheet
            
            # Create sheets
            self._create_summary_sheet(wb, bundle)
            self._create_ownership_sheet(wb, bundle)
            self._create_chain_sheet(wb, bundle)
            self._create_sources_sheet(wb, bundle)
            self._create_missing_sheet(wb, bundle)
            self._create_exceptions_sheet(wb, bundle)
            self._create_qa_sheet(wb, bundle)
            self._create_calculations_sheet(wb, bundle)
            self._create_log_sheet(wb, bundle)
            
            # Save
            wb.save(output_path)
            self.logger.info(f"XLSX report generated: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error generating XLSX: {e}")
            return False
    
    def _create_summary_sheet(self, wb: openpyxl.Workbook, bundle: ReportBundle) -> None:
        """Create summary sheet."""
        ws = wb.create_sheet("Summary")
        
        # Header style
        header_font = Font(bold=True, size=12)
        
        # Report metadata
        ws['A1'] = "Ownership Report Summary"
        ws['A1'].font = Font(bold=True, size=14)
        
        row = 3
        ws[f'A{row}'] = "Report Date:"
        ws[f'B{row}'] = bundle.metadata.report_date
        row += 1
        
        ws[f'A{row}'] = "Generated:"
        ws[f'B{row}'] = bundle.metadata.generated_at.strftime("%Y-%m-%d %H:%M:%S")
        row += 1
        
        if bundle.metadata.section:
            ws[f'A{row}'] = "Section:"
            ws[f'B{row}'] = f"{bundle.metadata.section}-{bundle.metadata.township}-{bundle.metadata.range}"
            row += 1
        
        if bundle.metadata.county:
            ws[f'A{row}'] = "County:"
            ws[f'B{row}'] = bundle.metadata.county
            row += 1
        
        if bundle.metadata.state:
            ws[f'A{row}'] = "State:"
            ws[f'B{row}'] = bundle.metadata.state
            row += 1
        
        row += 1
        ws[f'A{row}'] = "Statistics:"
        ws[f'A{row}'].font = header_font
        row += 1
        
        ws[f'A{row}'] = "Source Documents:"
        ws[f'B{row}'] = len(bundle.source_documents)
        row += 1
        
        ws[f'A{row}'] = "Ownership Entries:"
        ws[f'B{row}'] = len(bundle.ownership_entries)
        row += 1
        
        ws[f'A{row}'] = "Chain Events:"
        ws[f'B{row}'] = len(bundle.chain_events)
        row += 1
        
        ws[f'A{row}'] = "Confidence Score:"
        ws[f'B{row}'] = f"{bundle.confidence_score:.1f}/100"
        
        # Format columns
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 40
    
    def _create_ownership_sheet(self, wb: openpyxl.Workbook, bundle: ReportBundle) -> None:
        """Create ownership data sheet."""
        ws = wb.create_sheet("Ownership")
        
        # Headers
        headers = ["Owner Name", "Interest Type", "Interest Decimal", "Interest Fraction", 
                   "Net Acres", "Gross Acres", "Royalty", "Legal Description", 
                   "Instrument", "Recording Info", "Book", "Page", "Recording Date",
                   "Grantor", "Grantee", "Status", "Confidence", "Notes"]
        
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
        
        # Data
        for row_num, entry in enumerate(bundle.ownership_entries, 2):
            ws.cell(row=row_num, column=1, value=entry.owner_name)
            ws.cell(row=row_num, column=2, value=entry.interest_type)
            ws.cell(row=row_num, column=3, value=entry.interest_decimal)
            ws.cell(row=row_num, column=4, value=entry.interest_fraction)
            ws.cell(row=row_num, column=5, value=entry.net_acres)
            ws.cell(row=row_num, column=6, value=entry.gross_acres)
            ws.cell(row=row_num, column=7, value=entry.royalty)
            ws.cell(row=row_num, column=8, value=entry.legal_description)
            ws.cell(row=row_num, column=9, value=entry.source_instrument)
            ws.cell(row=row_num, column=10, value=entry.recording_info)
            ws.cell(row=row_num, column=11, value=entry.book)
            ws.cell(row=row_num, column=12, value=entry.page)
            ws.cell(row=row_num, column=13, value=entry.recording_date)
            ws.cell(row=row_num, column=14, value=entry.grantor)
            ws.cell(row=row_num, column=15, value=entry.grantee)
            ws.cell(row=row_num, column=16, value=entry.status)
            ws.cell(row=row_num, column=17, value=entry.confidence)
            ws.cell(row=row_num, column=18, value=entry.notes)
        
        # Freeze first row
        ws.freeze_panes = 'A2'
        
        # Auto-size columns (approximate)
        for col in range(1, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col)].width = 15
    
    def _create_chain_sheet(self, wb: openpyxl.Workbook, bundle: ReportBundle) -> None:
        """Create chain of title sheet."""
        ws = wb.create_sheet("Chain of Title")
        
        headers = ["Sequence", "Date", "Instrument Type", "Instrument Number",
                   "Grantor", "Grantee", "Conveyed Interest", "Retained Interest",
                   "Recording Info", "Verified", "Notes"]
        
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
        
        for row_num, event in enumerate(bundle.chain_events, 2):
            ws.cell(row=row_num, column=1, value=event.sequence)
            ws.cell(row=row_num, column=2, value=event.date)
            ws.cell(row=row_num, column=3, value=event.instrument_type)
            ws.cell(row=row_num, column=4, value=event.instrument_number)
            ws.cell(row=row_num, column=5, value=event.grantor)
            ws.cell(row=row_num, column=6, value=event.grantee)
            ws.cell(row=row_num, column=7, value=event.conveyed_interest)
            ws.cell(row=row_num, column=8, value=event.retained_interest)
            ws.cell(row=row_num, column=9, value=event.recording_info)
            ws.cell(row=row_num, column=10, value="Yes" if event.verified else "No")
            ws.cell(row=row_num, column=11, value=event.notes)
        
        ws.freeze_panes = 'A2'
        for col in range(1, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col)].width = 15
    
    def _create_sources_sheet(self, wb: openpyxl.Workbook, bundle: ReportBundle) -> None:
        """Create source documents sheet."""
        ws = wb.create_sheet("Source Documents")
        
        headers = ["File Name", "Path", "Type", "Size", "Modified", "Relevance", "Warnings"]
        
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
        
        for row_num, doc in enumerate(bundle.source_documents, 2):
            ws.cell(row=row_num, column=1, value=doc.path.name)
            ws.cell(row=row_num, column=2, value=str(doc.path))
            ws.cell(row=row_num, column=3, value=doc.file_type)
            ws.cell(row=row_num, column=4, value=doc.size)
            ws.cell(row=row_num, column=5, value=doc.modified_time.strftime("%Y-%m-%d %H:%M:%S"))
            ws.cell(row=row_num, column=6, value=doc.relevance)
            ws.cell(row=row_num, column=7, value=", ".join(doc.warnings) if doc.warnings else "")
        
        ws.freeze_panes = 'A2'
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 50
        for col in range(3, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col)].width = 15
    
    def _create_missing_sheet(self, wb: openpyxl.Workbook, bundle: ReportBundle) -> None:
        """Create missing items sheet."""
        ws = wb.create_sheet("Missing Items")
        
        ws['A1'] = "Missing Item"
        ws['A1'].font = Font(bold=True)
        ws['A1'].fill = PatternFill(start_color="FFE5E5", end_color="FFE5E5", fill_type="solid")
        
        for row_num, item in enumerate(bundle.missing_items, 2):
            ws.cell(row=row_num, column=1, value=item)
        
        ws.column_dimensions['A'].width = 80
    
    def _create_exceptions_sheet(self, wb: openpyxl.Workbook, bundle: ReportBundle) -> None:
        """Create exceptions sheet."""
        ws = wb.create_sheet("Exceptions")
        
        ws['A1'] = "Exception"
        ws['A1'].font = Font(bold=True)
        ws['A1'].fill = PatternFill(start_color="FFF5E5", end_color="FFF5E5", fill_type="solid")
        
        for row_num, exc in enumerate(bundle.exceptions, 2):
            ws.cell(row=row_num, column=1, value=exc)
        
        ws.column_dimensions['A'].width = 80
    
    def _create_qa_sheet(self, wb: openpyxl.Workbook, bundle: ReportBundle) -> None:
        """Create QA checks sheet."""
        ws = wb.create_sheet("QA Checks")
        
        headers = ["Check Name", "Status", "Message", "Severity"]
        
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
        
        for row_num, result in enumerate(bundle.qa_results, 2):
            ws.cell(row=row_num, column=1, value=result.check_name)
            ws.cell(row=row_num, column=2, value="PASS" if result.passed else "FAIL")
            ws.cell(row=row_num, column=3, value=result.message)
            ws.cell(row=row_num, column=4, value=result.severity)
        
        ws.freeze_panes = 'A2'
        for col in range(1, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col)].width = 25
    
    def _create_calculations_sheet(self, wb: openpyxl.Workbook, bundle: ReportBundle) -> None:
        """Create calculations sheet."""
        ws = wb.create_sheet("Calculations")
        
        ws['A1'] = "Calculation"
        ws['A1'].font = Font(bold=True)
        ws['B1'] = "Result"
        ws['B1'].font = Font(bold=True)
        
        row = 2
        
        # Calculate totals if data available
        if bundle.ownership_entries:
            total_interest = sum(e.interest_decimal for e in bundle.ownership_entries if e.interest_decimal)
            ws[f'A{row}'] = "Total Interest (Decimal)"
            ws[f'B{row}'] = total_interest
            row += 1
            
            ws[f'A{row}'] = "Interest Check (should be ~1.0)"
            ws[f'B{row}'] = "OK" if abs(total_interest - 1.0) < 0.01 else "VARIANCE"
            row += 1
            
            total_net_acres = sum(e.net_acres for e in bundle.ownership_entries if e.net_acres)
            ws[f'A{row}'] = "Total Net Acres"
            ws[f'B{row}'] = total_net_acres
            row += 1
        
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 20
    
    def _create_log_sheet(self, wb: openpyxl.Workbook, bundle: ReportBundle) -> None:
        """Create run log sheet."""
        ws = wb.create_sheet("Run Log")
        
        ws['A1'] = "Run Log Entry"
        ws['A1'].font = Font(bold=True)
        
        row = 2
        ws[f'A{row}'] = f"Report generated: {bundle.metadata.generated_at}"
        row += 1
        ws[f'A{row}'] = f"Input folder: {bundle.metadata.input_folder}"
        row += 1
        ws[f'A{row}'] = f"Sources processed: {len(bundle.source_documents)}"
        row += 1
        ws[f'A{row}'] = f"Ownership entries: {len(bundle.ownership_entries)}"
        row += 1
        ws[f'A{row}'] = f"Chain events: {len(bundle.chain_events)}"
        row += 1
        ws[f'A{row}'] = f"Confidence score: {bundle.confidence_score:.1f}"
        
        ws.column_dimensions['A'].width = 60


class PdfGenerator:
    """Generator for PDF reports."""
    
    def __init__(self, logger):
        """Initialize PDF generator.
        
        Args:
            logger: Logger instance
        """
        self.logger = logger
    
    def generate(self, bundle: ReportBundle, output_path: Path) -> bool:
        """Generate PDF report.
        
        Args:
            bundle: Report data bundle
            output_path: Output file path
            
        Returns:
            True if successful
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib import colors
            
            doc = SimpleDocTemplate(str(output_path), pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#003366'),
                spaceAfter=30,
                alignment=1  # Center
            )
            story.append(Paragraph(bundle.metadata.report_type, title_style))
            story.append(Spacer(1, 0.2 * inch))
            
            # Metadata
            story.append(Paragraph(f"<b>Report Date:</b> {bundle.metadata.report_date}", styles['Normal']))
            story.append(Paragraph(f"<b>Generated:</b> {bundle.metadata.generated_at.strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            
            if bundle.metadata.section:
                section_str = f"{bundle.metadata.section}-{bundle.metadata.township}-{bundle.metadata.range}"
                story.append(Paragraph(f"<b>Section:</b> {section_str}", styles['Normal']))
            
            if bundle.metadata.county:
                story.append(Paragraph(f"<b>County:</b> {bundle.metadata.county}, {bundle.metadata.state}", styles['Normal']))
            
            story.append(Spacer(1, 0.3 * inch))
            
            # Summary
            story.append(Paragraph("<b>Summary</b>", styles['Heading2']))
            story.append(Paragraph(f"Source Documents: {len(bundle.source_documents)}", styles['Normal']))
            story.append(Paragraph(f"Ownership Entries: {len(bundle.ownership_entries)}", styles['Normal']))
            story.append(Paragraph(f"Chain Events: {len(bundle.chain_events)}", styles['Normal']))
            story.append(Paragraph(f"Confidence Score: {bundle.confidence_score:.1f}/100", styles['Normal']))
            story.append(Spacer(1, 0.3 * inch))
            
            # Ownership table (limited)
            if bundle.ownership_entries:
                story.append(Paragraph("<b>Ownership Data (First 20 Entries)</b>", styles['Heading2']))
                
                table_data = [["Owner", "Interest", "Net Acres", "Instrument"]]
                for entry in bundle.ownership_entries[:20]:
                    table_data.append([
                        entry.owner_name or "",
                        entry.interest_fraction or (f"{entry.interest_decimal:.4f}" if entry.interest_decimal else ""),
                        f"{entry.net_acres:.2f}" if entry.net_acres else "",
                        entry.source_instrument or ""
                    ])
                
                table = Table(table_data, colWidths=[2*inch, 1*inch, 1*inch, 1.5*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
            else:
                story.append(Paragraph("<b>⚠️ NO VERIFIED OWNERSHIP DATA EXTRACTED — SOURCE REVIEW REQUIRED</b>", styles['Normal']))
            
            story.append(Spacer(1, 0.3 * inch))
            
            # QA Summary
            story.append(Paragraph("<b>Quality Assurance</b>", styles['Heading2']))
            story.append(Paragraph(f"Confidence Score: {bundle.confidence_score:.1f}/100", styles['Normal']))
            
            if bundle.qa_results:
                passed = sum(1 for r in bundle.qa_results if r.passed)
                story.append(Paragraph(f"QA Checks Passed: {passed}/{len(bundle.qa_results)}", styles['Normal']))
            
            # Build PDF
            doc.build(story)
            self.logger.info(f"PDF report generated: {output_path}")
            return True
            
        except ImportError:
            self.logger.warning("reportlab not installed - PDF generation skipped")
            self.logger.info("Install with: pip install reportlab")
            return False
        except Exception as e:
            self.logger.error(f"Error generating PDF: {e}")
            return False
