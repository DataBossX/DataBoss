"""Document parsers for extracting data from various file formats."""

import pandas as pd
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document as DocxDocument
import pypdf
from .models import SourceDocument, OwnershipEntry, ChainEvent
from .utils import normalize_column_name


class DocumentParser:
    """Base class for document parsers."""
    
    def __init__(self, logger):
        """Initialize parser.
        
        Args:
            logger: Logger instance
        """
        self.logger = logger
    
    def parse(self, doc: SourceDocument) -> Dict[str, Any]:
        """Parse a document and extract data.
        
        Args:
            doc: Source document to parse
            
        Returns:
            Dictionary containing parsed data
        """
        raise NotImplementedError("Subclasses must implement parse()")


class ExcelParser(DocumentParser):
    """Parser for Excel and CSV files."""
    
    # Column name mappings for common variations
    COLUMN_MAPPINGS = {
        "owner": ["owner", "owner_name", "ownername", "party", "party_name"],
        "interest": ["interest", "interest_decimal", "interest_pct", "ownership"],
        "interest_fraction": ["interest_fraction", "fraction", "interest_frac"],
        "net_acres": ["net_acres", "netacres", "nma", "net_mineral_acres"],
        "gross_acres": ["gross_acres", "grossacres", "gross", "total_acres"],
        "royalty": ["royalty", "roi", "royalty_interest"],
        "legal": ["legal", "legal_description", "description", "tract"],
        "instrument": ["instrument", "instrument_number", "inst_no", "recording_no"],
        "book": ["book", "book_no"],
        "page": ["page", "page_no"],
        "recording_date": ["recording_date", "recorded", "date_recorded", "rec_date"],
        "grantor": ["grantor", "from"],
        "grantee": ["grantee", "to"],
    }
    
    def parse(self, doc: SourceDocument) -> Dict[str, Any]:
        """Parse Excel or CSV file.
        
        Args:
            doc: Source document
            
        Returns:
            Parsed data dictionary
        """
        try:
            if doc.file_type == '.csv':
                return self._parse_csv(doc.path)
            else:
                return self._parse_excel(doc.path)
        except Exception as e:
            self.logger.error(f"Error parsing {doc.path}: {e}")
            return {"error": str(e), "sheets": []}
    
    def _parse_csv(self, path: Path) -> Dict[str, Any]:
        """Parse CSV file.
        
        Args:
            path: File path
            
        Returns:
            Parsed data
        """
        try:
            df = pd.read_csv(path)
            return {
                "sheets": [{
                    "name": "Sheet1",
                    "rows": len(df),
                    "columns": list(df.columns),
                    "data": df.to_dict('records'),
                    "normalized_columns": self._normalize_columns(df.columns)
                }]
            }
        except Exception as e:
            self.logger.error(f"CSV parse error: {e}")
            return {"error": str(e), "sheets": []}
    
    def _parse_excel(self, path: Path) -> Dict[str, Any]:
        """Parse Excel file.
        
        Args:
            path: File path
            
        Returns:
            Parsed data
        """
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(path)
            sheets_data = []
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(path, sheet_name=sheet_name)
                    
                    sheets_data.append({
                        "name": sheet_name,
                        "rows": len(df),
                        "columns": list(df.columns),
                        "data": df.to_dict('records'),
                        "normalized_columns": self._normalize_columns(df.columns)
                    })
                except Exception as e:
                    self.logger.warning(f"Error reading sheet {sheet_name}: {e}")
                    sheets_data.append({
                        "name": sheet_name,
                        "error": str(e)
                    })
            
            return {"sheets": sheets_data}
        except Exception as e:
            self.logger.error(f"Excel parse error: {e}")
            return {"error": str(e), "sheets": []}
    
    def _normalize_columns(self, columns: List[str]) -> Dict[str, str]:
        """Normalize column names and map to standard fields.
        
        Args:
            columns: List of column names
            
        Returns:
            Dictionary mapping normalized names to standard fields
        """
        mapping = {}
        
        for col in columns:
            normalized = normalize_column_name(col)
            
            # Try to match to standard field
            for standard_field, variations in self.COLUMN_MAPPINGS.items():
                for variation in variations:
                    if variation in normalized:
                        mapping[col] = standard_field
                        break
        
        return mapping
    
    def extract_ownership_entries(self, parsed_data: Dict[str, Any]) -> List[OwnershipEntry]:
        """Extract ownership entries from parsed Excel/CSV data.
        
        Args:
            parsed_data: Parsed data dictionary
            
        Returns:
            List of ownership entries
        """
        entries = []
        
        for sheet in parsed_data.get("sheets", []):
            if "error" in sheet:
                continue
            
            data = sheet.get("data", [])
            col_mapping = sheet.get("normalized_columns", {})
            
            for row in data:
                entry = OwnershipEntry()
                
                # Map data using column mappings
                for original_col, standard_field in col_mapping.items():
                    value = row.get(original_col)
                    
                    if value is None or (isinstance(value, float) and pd.isna(value)):
                        continue
                    
                    if standard_field == "owner":
                        entry.owner_name = str(value)
                    elif standard_field == "interest":
                        entry.interest_decimal = float(value) if isinstance(value, (int, float)) else None
                    elif standard_field == "interest_fraction":
                        entry.interest_fraction = str(value)
                    elif standard_field == "net_acres":
                        entry.net_acres = float(value) if isinstance(value, (int, float)) else None
                    elif standard_field == "gross_acres":
                        entry.gross_acres = float(value) if isinstance(value, (int, float)) else None
                    elif standard_field == "royalty":
                        entry.royalty = float(value) if isinstance(value, (int, float)) else None
                    elif standard_field == "legal":
                        entry.legal_description = str(value)
                    elif standard_field == "instrument":
                        entry.source_instrument = str(value)
                    elif standard_field == "book":
                        entry.book = str(value)
                    elif standard_field == "page":
                        entry.page = str(value)
                    elif standard_field == "recording_date":
                        entry.recording_date = str(value)
                    elif standard_field == "grantor":
                        entry.grantor = str(value)
                    elif standard_field == "grantee":
                        entry.grantee = str(value)
                
                # Only add if we found at least an owner or instrument
                if entry.owner_name or entry.source_instrument:
                    entries.append(entry)
        
        return entries


class DocxParser(DocumentParser):
    """Parser for Word documents."""
    
    def parse(self, doc: SourceDocument) -> Dict[str, Any]:
        """Parse DOCX file.
        
        Args:
            doc: Source document
            
        Returns:
            Parsed data
        """
        try:
            document = DocxDocument(doc.path)
            
            # Extract paragraphs
            paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
            
            # Extract tables
            tables = []
            for table in document.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables)
            }
        except Exception as e:
            self.logger.error(f"DOCX parse error: {e}")
            return {"error": str(e), "paragraphs": [], "tables": []}


class PdfParser(DocumentParser):
    """Parser for PDF files."""
    
    def parse(self, doc: SourceDocument) -> Dict[str, Any]:
        """Parse PDF file.
        
        Args:
            doc: Source document
            
        Returns:
            Parsed data
        """
        try:
            with open(doc.path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                
                pages = []
                for page_num, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        pages.append({
                            "page_number": page_num + 1,
                            "text": text,
                            "text_length": len(text)
                        })
                    except Exception as e:
                        pages.append({
                            "page_number": page_num + 1,
                            "error": str(e)
                        })
                
                # Check if PDF is mostly images (low text content)
                avg_text_length = sum(p.get("text_length", 0) for p in pages) / len(pages) if pages else 0
                is_image_based = avg_text_length < 100
                
                return {
                    "page_count": len(pages),
                    "pages": pages,
                    "is_image_based": is_image_based,
                    "needs_ocr": is_image_based
                }
        except Exception as e:
            self.logger.error(f"PDF parse error: {e}")
            return {"error": str(e), "pages": []}


class ParserFactory:
    """Factory for creating appropriate parser based on file type."""
    
    @staticmethod
    def create_parser(doc: SourceDocument, logger) -> Optional[DocumentParser]:
        """Create appropriate parser for document.
        
        Args:
            doc: Source document
            logger: Logger instance
            
        Returns:
            Parser instance or None if unsupported
        """
        if doc.file_type in ['.xlsx', '.xls', '.csv']:
            return ExcelParser(logger)
        elif doc.file_type == '.docx':
            return DocxParser(logger)
        elif doc.file_type == '.pdf':
            return PdfParser(logger)
        else:
            return None
