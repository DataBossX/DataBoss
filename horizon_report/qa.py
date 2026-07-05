"""Quality assurance and validation for generated reports."""

from pathlib import Path
from typing import List, Tuple
import openpyxl
from .models import QAResult, ReportBundle


class QAEngine:
    """Quality assurance engine for report validation."""
    
    def __init__(self, logger, config):
        """Initialize QA engine.
        
        Args:
            logger: Logger instance
            config: Configuration instance
        """
        self.logger = logger
        self.config = config
        self.results: List[QAResult] = []
    
    def run_all_checks(self, bundle: ReportBundle, generated_files: dict) -> Tuple[List[QAResult], float]:
        """Run all QA checks.
        
        Args:
            bundle: Report data bundle
            generated_files: Dictionary of generated file paths
            
        Returns:
            Tuple of (QA results list, confidence score)
        """
        self.results = []
        
        # File existence checks
        self._check_file_exists(generated_files.get('docx'), "DOCX file exists")
        self._check_file_exists(generated_files.get('xlsx'), "XLSX file exists")
        self._check_file_exists(generated_files.get('pdf'), "PDF file exists")
        self._check_file_exists(generated_files.get('qa_md'), "QA markdown exists")
        self._check_file_exists(generated_files.get('log'), "Log file exists")
        
        # File content checks
        if generated_files.get('docx'):
            self._check_docx_not_empty(generated_files['docx'])
        
        if generated_files.get('xlsx'):
            self._check_xlsx_valid(generated_files['xlsx'])
        
        # Input validation
        self._check_input_folder_exists()
        self._check_input_files_found(bundle)
        
        # Data quality checks
        self._check_ownership_data(bundle)
        self._check_data_completeness(bundle)
        self._check_data_integrity(bundle)
        
        # Security checks
        self._check_no_secrets_in_output(bundle)
        
        # Google Drive check
        self._check_google_drive_status(generated_files)
        
        # Calculate confidence score
        confidence = self._calculate_confidence_score(bundle)
        
        return self.results, confidence
    
    def _add_result(self, check_name: str, passed: bool, message: str, 
                   severity: str = "info", details: dict = None) -> None:
        """Add a QA result.
        
        Args:
            check_name: Name of the check
            passed: Whether check passed
            message: Result message
            severity: Severity level (info, warning, error, critical)
            details: Additional details dictionary
        """
        result = QAResult(
            check_name=check_name,
            passed=passed,
            message=message,
            severity=severity,
            details=details or {}
        )
        self.results.append(result)
        
        level = self.logger.info if passed else self.logger.warning
        if severity in ['error', 'critical']:
            level = self.logger.error
        
        level(f"QA Check: {check_name} - {'PASS' if passed else 'FAIL'} - {message}")
    
    def _check_file_exists(self, file_path: Path, check_name: str) -> None:
        """Check if a file exists.
        
        Args:
            file_path: Path to check
            check_name: Name of check
        """
        if file_path and file_path.exists():
            size = file_path.stat().st_size
            self._add_result(
                check_name, 
                True, 
                f"File exists ({size} bytes)",
                "info",
                {"path": str(file_path), "size": size}
            )
        else:
            severity = "error" if "DOCX" in check_name or "XLSX" in check_name else "warning"
            self._add_result(
                check_name, 
                False, 
                f"File not found: {file_path}" if file_path else "File path not provided",
                severity
            )
    
    def _check_docx_not_empty(self, file_path: Path) -> None:
        """Check that DOCX file is not empty.
        
        Args:
            file_path: Path to DOCX file
        """
        try:
            from docx import Document
            doc = Document(file_path)
            
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            table_count = len(doc.tables)
            
            if paragraph_count > 0 or table_count > 0:
                self._add_result(
                    "DOCX content check",
                    True,
                    f"DOCX has {paragraph_count} paragraphs and {table_count} tables",
                    "info",
                    {"paragraphs": paragraph_count, "tables": table_count}
                )
            else:
                self._add_result(
                    "DOCX content check",
                    False,
                    "DOCX appears to be empty",
                    "error"
                )
        except Exception as e:
            self._add_result(
                "DOCX content check",
                False,
                f"Error reading DOCX: {e}",
                "error"
            )
    
    def _check_xlsx_valid(self, file_path: Path) -> None:
        """Check that XLSX file is valid and has required sheets.
        
        Args:
            file_path: Path to XLSX file
        """
        try:
            wb = openpyxl.load_workbook(file_path, read_only=True)
            sheet_names = wb.sheetnames
            wb.close()
            
            required_sheets = ["Summary", "Ownership", "QA Checks"]
            missing_sheets = [s for s in required_sheets if s not in sheet_names]
            
            if not missing_sheets:
                self._add_result(
                    "XLSX structure check",
                    True,
                    f"XLSX has all required sheets ({len(sheet_names)} total)",
                    "info",
                    {"sheets": sheet_names}
                )
            else:
                self._add_result(
                    "XLSX structure check",
                    False,
                    f"Missing required sheets: {', '.join(missing_sheets)}",
                    "warning",
                    {"missing": missing_sheets, "found": sheet_names}
                )
        except Exception as e:
            self._add_result(
                "XLSX structure check",
                False,
                f"Error reading XLSX: {e}",
                "error"
            )
    
    def _check_input_folder_exists(self) -> None:
        """Check that input folder exists."""
        if self.config.input_path.exists():
            file_count = len(list(self.config.input_path.glob("*")))
            self._add_result(
                "Input folder exists",
                True,
                f"Input folder found with {file_count} items",
                "info"
            )
        else:
            self._add_result(
                "Input folder exists",
                False,
                f"Input folder not found: {self.config.input_path}",
                "error"
            )
    
    def _check_input_files_found(self, bundle: ReportBundle) -> None:
        """Check that input files were found.
        
        Args:
            bundle: Report bundle
        """
        count = len(bundle.source_documents)
        
        if count > 0:
            self._add_result(
                "Input files found",
                True,
                f"Found {count} source files",
                "info",
                {"count": count}
            )
        else:
            self._add_result(
                "Input files found",
                False,
                "No source files were found in input folder",
                "error"
            )
    
    def _check_ownership_data(self, bundle: ReportBundle) -> None:
        """Check ownership data quality.
        
        Args:
            bundle: Report bundle
        """
        count = len(bundle.ownership_entries)
        
        if count == 0:
            self._add_result(
                "Ownership data extracted",
                False,
                "No ownership data was extracted from sources",
                "critical",
                {"count": 0}
            )
            return
        
        # Check for verified data
        with_owner = sum(1 for e in bundle.ownership_entries if e.owner_name)
        with_interest = sum(1 for e in bundle.ownership_entries if e.interest_decimal or e.interest_fraction)
        
        if with_owner > 0 and with_interest > 0:
            self._add_result(
                "Ownership data extracted",
                True,
                f"Extracted {count} entries ({with_owner} with owner, {with_interest} with interest)",
                "info",
                {"total": count, "with_owner": with_owner, "with_interest": with_interest}
            )
        else:
            self._add_result(
                "Ownership data extracted",
                False,
                f"Extracted {count} entries but data appears incomplete",
                "warning",
                {"total": count, "with_owner": with_owner, "with_interest": with_interest}
            )
    
    def _check_data_completeness(self, bundle: ReportBundle) -> None:
        """Check data completeness.
        
        Args:
            bundle: Report bundle
        """
        if not bundle.ownership_entries:
            return
        
        # Check for missing legal descriptions
        missing_legal = sum(1 for e in bundle.ownership_entries if not e.legal_description)
        if missing_legal > 0:
            pct = (missing_legal / len(bundle.ownership_entries)) * 100
            self._add_result(
                "Legal descriptions present",
                pct < 50,
                f"{missing_legal} entries ({pct:.1f}%) missing legal description",
                "warning" if pct < 50 else "error"
            )
        else:
            self._add_result(
                "Legal descriptions present",
                True,
                "All entries have legal descriptions",
                "info"
            )
        
        # Check for missing recording info
        missing_recording = sum(1 for e in bundle.ownership_entries 
                               if not e.recording_info and not e.book and not e.recording_date)
        if missing_recording > 0:
            pct = (missing_recording / len(bundle.ownership_entries)) * 100
            self._add_result(
                "Recording information present",
                pct < 50,
                f"{missing_recording} entries ({pct:.1f}%) missing recording info",
                "warning"
            )
        else:
            self._add_result(
                "Recording information present",
                True,
                "All entries have recording information",
                "info"
            )
    
    def _check_data_integrity(self, bundle: ReportBundle) -> None:
        """Check data integrity.
        
        Args:
            bundle: Report bundle
        """
        if not bundle.ownership_entries:
            return
        
        # Check interest totals
        entries_with_interest = [e for e in bundle.ownership_entries if e.interest_decimal]
        if entries_with_interest:
            total_interest = sum(e.interest_decimal for e in entries_with_interest)
            
            # Check if close to 100% (1.0)
            if abs(total_interest - 1.0) < 0.05:  # Within 5%
                self._add_result(
                    "Interest total check",
                    True,
                    f"Total interest is {total_interest:.4f} (within 5% of 100%)",
                    "info",
                    {"total": total_interest}
                )
            else:
                severity = "warning" if abs(total_interest - 1.0) < 0.20 else "error"
                self._add_result(
                    "Interest total check",
                    False,
                    f"Total interest is {total_interest:.4f} (variance from 100%)",
                    severity,
                    {"total": total_interest, "variance": abs(total_interest - 1.0)}
                )
        
        # Check for duplicate owners
        owner_names = [e.owner_name for e in bundle.ownership_entries if e.owner_name]
        if owner_names:
            unique_owners = set(owner_names)
            if len(owner_names) > len(unique_owners):
                duplicates = len(owner_names) - len(unique_owners)
                self._add_result(
                    "Duplicate owner check",
                    False,
                    f"Found {duplicates} duplicate owner names",
                    "warning",
                    {"duplicates": duplicates}
                )
            else:
                self._add_result(
                    "Duplicate owner check",
                    True,
                    "No duplicate owner names found",
                    "info"
                )
    
    def _check_no_secrets_in_output(self, bundle: ReportBundle) -> None:
        """Check that no secrets are exposed in output.
        
        Args:
            bundle: Report bundle
        """
        # This is a basic check - in production, would be more sophisticated
        suspicious_patterns = ['sk-', 'ya29.', 'AIza', 'AKIA', 'token', 'secret', 'password']
        
        # Check in analyst notes
        for note in bundle.analyst_notes:
            for pattern in suspicious_patterns:
                if pattern in note.lower():
                    self._add_result(
                        "Secret exposure check",
                        False,
                        f"Potential secret found in analyst notes: {pattern}",
                        "critical"
                    )
                    return
        
        self._add_result(
            "Secret exposure check",
            True,
            "No obvious secrets found in output",
            "info"
        )
    
    def _check_google_drive_status(self, generated_files: dict) -> None:
        """Check Google Drive sync status.
        
        Args:
            generated_files: Generated files dictionary
        """
        drive_status = generated_files.get('drive_sync_status', 'unknown')
        
        if drive_status == 'success':
            self._add_result(
                "Google Drive sync",
                True,
                "Files successfully synced to Google Drive",
                "info"
            )
        elif drive_status == 'skipped':
            self._add_result(
                "Google Drive sync",
                True,
                "Google Drive sync skipped (not configured)",
                "info"
            )
        elif drive_status == 'failed':
            self._add_result(
                "Google Drive sync",
                False,
                "Google Drive sync failed",
                "warning"
            )
    
    def _calculate_confidence_score(self, bundle: ReportBundle) -> float:
        """Calculate overall confidence score.
        
        Args:
            bundle: Report bundle
            
        Returns:
            Confidence score 0-100
        """
        score = 0.0
        
        # Base score for having data (30 points)
        if bundle.ownership_entries:
            score += 20
            
            # Quality of ownership data (10 points)
            with_owner = sum(1 for e in bundle.ownership_entries if e.owner_name)
            with_interest = sum(1 for e in bundle.ownership_entries if e.interest_decimal or e.interest_fraction)
            completeness = (with_owner + with_interest) / (2 * len(bundle.ownership_entries))
            score += 10 * completeness
        
        # Source documents found (15 points)
        if bundle.source_documents:
            score += min(15, len(bundle.source_documents) * 3)
        
        # Legal descriptions present (10 points)
        if bundle.ownership_entries:
            with_legal = sum(1 for e in bundle.ownership_entries if e.legal_description)
            legal_pct = with_legal / len(bundle.ownership_entries)
            score += 10 * legal_pct
        
        # Recording info present (10 points)
        if bundle.ownership_entries:
            with_recording = sum(1 for e in bundle.ownership_entries 
                                if e.recording_info or e.book or e.recording_date)
            recording_pct = with_recording / len(bundle.ownership_entries)
            score += 10 * recording_pct
        
        # Chain of title data (10 points)
        if bundle.chain_events:
            score += min(10, len(bundle.chain_events) * 2)
        
        # QA checks passed (15 points)
        if self.results:
            passed = sum(1 for r in self.results if r.passed)
            qa_pct = passed / len(self.results)
            score += 15 * qa_pct
        
        # Penalties
        # Critical failures
        critical_failures = sum(1 for r in self.results if not r.passed and r.severity == 'critical')
        score -= critical_failures * 20
        
        # No ownership data is a major penalty
        if not bundle.ownership_entries:
            score = min(score, 40)  # Cap at 40 if no data
        
        # Clamp to 0-100
        return max(0.0, min(100.0, score))
    
    def generate_qa_report(self, output_path: Path, bundle: ReportBundle, confidence: float) -> None:
        """Generate QA report in markdown.
        
        Args:
            output_path: Output file path
            bundle: Report bundle
            confidence: Confidence score
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# Quality Assurance Report\n\n")
            f.write(f"**Generated:** {bundle.metadata.generated_at.strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            
            # Confidence score
            f.write("## Confidence Score\n\n")
            f.write(f"**Overall Score: {confidence:.1f}/100**\n\n")
            
            if confidence >= 80:
                f.write("✅ **HIGH CONFIDENCE** - Report data appears reliable\n\n")
            elif confidence >= 60:
                f.write("⚠️ **MEDIUM CONFIDENCE** - Report has some gaps or issues\n\n")
            else:
                f.write("❌ **LOW CONFIDENCE** - Report requires significant human review\n\n")
            
            # Summary
            passed = sum(1 for r in self.results if r.passed)
            total = len(self.results)
            f.write("## Check Summary\n\n")
            f.write(f"- **Passed:** {passed}/{total} checks\n")
            f.write(f"- **Failed:** {total - passed} checks\n\n")
            
            # Details by severity
            for severity in ['critical', 'error', 'warning', 'info']:
                severity_results = [r for r in self.results if r.severity == severity]
                if not severity_results:
                    continue
                
                icon = {'critical': '🔴', 'error': '❌', 'warning': '⚠️', 'info': 'ℹ️'}[severity]
                f.write(f"## {icon} {severity.upper()} ({len(severity_results)})\n\n")
                
                for result in severity_results:
                    status = "✅ PASS" if result.passed else "❌ FAIL"
                    f.write(f"### {status}: {result.check_name}\n\n")
                    f.write(f"{result.message}\n\n")
                    
                    if result.details:
                        f.write("**Details:**\n")
                        for key, value in result.details.items():
                            f.write(f"- {key}: {value}\n")
                        f.write("\n")
        
        self.logger.info(f"QA report generated: {output_path}")
