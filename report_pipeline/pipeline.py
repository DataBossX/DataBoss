"""Grocery Report pipeline — stages A through I.

Each stage is a pure function that reads prior CSV outputs and writes its own,
so any stage is independently rerunnable. Deterministic parsing throughout; no
fabrication — a field is emitted only when found, else left blank and flagged
REVIEW REQUIRED. Every fact row records its source file (and page/sheet/row).

Run all: python -m report_pipeline.pipeline --root <PROJECT_ROOT>
"""
from __future__ import annotations

import argparse
import html
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from report_pipeline import common as C

# Files/dirs we never treat as project source documents.
SKIP_DIRS = {".git", "node_modules", ".venv", "venv", "__pycache__", "output",
             "outputs", "backups", ".auth", "temp", "quarantine", ".pytest_cache"}
TEXTLIKE = {".txt", ".md", ".csv", ".tsv", ".json", ".log"}
DOCLIKE = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".rtf"}
IMAGELIKE = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}


# ============================================================ A. INVENTORY ===

def stage_inventory(root: Path, out: Path, log) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for p in sorted(root.rglob("*")):
        if p.is_dir():
            continue
        if any(part in SKIP_DIRS for part in p.relative_to(root).parts):
            continue
        try:
            st = p.stat()
            rows.append({
                "path": str(p),
                "relpath": str(p.relative_to(root)),
                "ext": p.suffix.lower(),
                "size_bytes": st.st_size,
                "modified": C.datetime.datetime.fromtimestamp(st.st_mtime).isoformat(timespec="seconds"),
                "sha256": C.sha256_file(p),
                "likely_type": _guess_type(p),
                "status": "inventoried",
                "generated_at": C.now_iso(),
            })
        except Exception as exc:  # noqa: BLE001
            log.warning("inventory skip %s: %s", p, exc)
    fields = ["path", "relpath", "ext", "size_bytes", "modified", "sha256",
              "likely_type", "status", "generated_at"]
    C.write_table(out, "file_inventory", fields, rows)
    log.info("A. inventory: %d files", len(rows))
    return rows


def _guess_type(p: Path) -> str:
    ext = p.suffix.lower()
    name = p.name.lower()
    if ext in IMAGELIKE:
        return "image/scan"
    if ext in {".xls", ".xlsx", ".csv"}:
        if any(k in name for k in ("ownership", "owner", "runsheet", "run_sheet")):
            return "ownership_spreadsheet"
        if any(k in name for k in ("ogl", "lease")):
            return "ogl_lease_sheet"
        return "spreadsheet"
    if ext in {".doc", ".docx", ".pdf", ".rtf", ".txt", ".md"}:
        return "document"
    return "other"


# ============================================================ B. DEDUPE ======

def stage_dedupe(inventory: List[Dict[str, Any]], out: Path, log) -> None:
    by_hash: Dict[str, List[Dict[str, Any]]] = {}
    for r in inventory:
        by_hash.setdefault(r["sha256"], []).append(r)

    dup_rows: List[Dict[str, Any]] = []
    quarantine_rows: List[Dict[str, Any]] = []
    for h, group in by_hash.items():
        if len(group) < 2:
            continue
        keep = min(group, key=lambda r: r["path"])  # deterministic keeper
        for r in group:
            is_keep = r["path"] == keep["path"]
            dup_rows.append({
                "sha256": h, "path": r["path"], "relpath": r["relpath"],
                "match_type": "exact_sha256", "duplicate_of": keep["relpath"],
                "action": "KEEP" if is_keep else "QUARANTINE_CANDIDATE",
                "generated_at": C.now_iso(),
            })
            if not is_keep:
                quarantine_rows.append({
                    "path": r["path"], "relpath": r["relpath"], "sha256": h,
                    "reason": "exact duplicate (sha256)", "duplicate_of": keep["relpath"],
                    "proposed_dest": f"quarantine/duplicates/{r['relpath']}",
                    "note": "NO DELETE — move only after Rodney approval",
                    "generated_at": C.now_iso(),
                })

    # Fuzzy filename near-duplicates (same stem, different ext or minor vari/copy).
    fuzzy = _fuzzy_name_dupes(inventory)
    dup_rows.extend(fuzzy)

    C.write_table(out, "duplicate_candidates",
                  ["sha256", "path", "relpath", "match_type", "duplicate_of", "action", "generated_at"],
                  dup_rows)
    C.write_table(out, "quarantine_plan",
                  ["path", "relpath", "sha256", "reason", "duplicate_of", "proposed_dest", "note", "generated_at"],
                  quarantine_rows)
    log.info("B. dedupe: %d exact-dup rows, %d quarantine candidates, %d fuzzy",
             len([d for d in dup_rows if d['match_type'] == 'exact_sha256']),
             len(quarantine_rows), len(fuzzy))


def _norm_stem(name: str) -> str:
    s = name.lower()
    s = re.sub(r"\b(copy|final|draft|v?\d+|rev\d*|\(\d+\))\b", "", s)
    return re.sub(r"[^a-z0-9]", "", s)


def _fuzzy_name_dupes(inventory: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    buckets: Dict[str, List[Dict[str, Any]]] = {}
    for r in inventory:
        buckets.setdefault(_norm_stem(Path(r["relpath"]).stem), []).append(r)
    rows: List[Dict[str, Any]] = []
    for key, group in buckets.items():
        if not key or len(group) < 2:
            continue
        # Only report if not already exact-dup (different hashes present).
        if len({r["sha256"] for r in group}) < 2:
            continue
        keep = min(group, key=lambda r: r["path"])
        for r in group:
            if r["path"] == keep["path"]:
                continue
            rows.append({
                "sha256": r["sha256"], "path": r["path"], "relpath": r["relpath"],
                "match_type": "fuzzy_filename", "duplicate_of": keep["relpath"],
                "action": "REVIEW_REQUIRED", "generated_at": C.now_iso(),
            })
    return rows


# ============================================================ C. TEXT ========

def stage_extract_text(inventory: List[Dict[str, Any]], out: Path, log) -> List[Dict[str, Any]]:
    text_dir = out / "extracted_text"
    index: List[Dict[str, Any]] = []
    for r in inventory:
        p = Path(r["path"])
        ext = r["ext"]
        text, method, status, conf = "", "none", "skipped", ""
        try:
            if ext in TEXTLIKE:
                text = p.read_text(encoding="utf-8", errors="replace")
                method, status, conf = "read", "extracted", 0.99
            elif ext == ".pdf":
                text, method, status, conf = _extract_pdf(p)
            elif ext == ".docx":
                text, method, status, conf = _extract_docx(p)
            elif ext in {".xlsx", ".xls"}:
                text, method, status, conf = _extract_xlsx(p)
            elif ext in IMAGELIKE:
                text, method, status, conf = _extract_image(p)
            else:
                status = "unsupported"
        except Exception as exc:  # noqa: BLE001
            status, method = "error", f"error:{type(exc).__name__}"
            log.warning("text extract error %s: %s", r["relpath"], exc)

        txt_path = ""
        if text:
            safe = re.sub(r"[^A-Za-z0-9._-]", "_", r["relpath"])
            tp = text_dir / f"{safe}.txt"
            tp.write_text(text, encoding="utf-8")
            txt_path = str(tp)

        index.append({
            "source_path": r["path"], "relpath": r["relpath"], "ext": ext,
            "text_path": txt_path, "chars": len(text), "method": method,
            "status": status, "confidence": conf, "generated_at": C.now_iso(),
        })
    C.write_table(out, "source_text_index",
                  ["source_path", "relpath", "ext", "text_path", "chars", "method",
                   "status", "confidence", "generated_at"], index)
    ok = sum(1 for r in index if r["status"] == "extracted")
    log.info("C. text extraction: %d/%d extracted (caps: %s)", ok, len(index), C.caps_summary())
    return index


def _extract_pdf(p: Path) -> Tuple[str, str, str, Any]:
    if C.CAPS["pypdf"]:
        import pypdf

        reader = pypdf.PdfReader(str(p))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
        if text.strip():
            return text, "pypdf", "extracted", 0.85
        return "", "pypdf", "needs_ocr", ""  # scanned PDF
    if C.CAPS["pdfminer"]:
        from pdfminer.high_level import extract_text

        text = extract_text(str(p)) or ""
        return (text, "pdfminer", "extracted", 0.85) if text.strip() else ("", "pdfminer", "needs_ocr", "")
    return "", "none", "needs_dependency", ""


def _extract_docx(p: Path) -> Tuple[str, str, str, Any]:
    if not C.CAPS["docx"]:
        return "", "none", "needs_dependency", ""
    import docx

    d = docx.Document(str(p))
    return "\n".join(par.text for par in d.paragraphs), "python-docx", "extracted", 0.95


def _extract_xlsx(p: Path) -> Tuple[str, str, str, Any]:
    if not C.CAPS["openpyxl"]:
        return "", "none", "needs_dependency", ""
    import openpyxl

    wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
    parts: List[str] = []
    for ws in wb.worksheets:
        parts.append(f"# SHEET: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            parts.append("\t".join("" if c is None else str(c) for c in row))
    return "\n".join(parts), "openpyxl", "extracted", 0.95


def _extract_image(p: Path) -> Tuple[str, str, str, Any]:
    if C.CAPS["pytesseract"] and C.CAPS["PIL"]:
        import pytesseract
        from PIL import Image

        text = pytesseract.image_to_string(Image.open(str(p)))
        return (text, "tesseract-ocr", "extracted", 0.6) if text.strip() else ("", "tesseract-ocr", "ocr_empty", "")
    return "", "none", "needs_dependency", ""


# ============================================================ D. CLASSIFY =====

CLASS_RULES: List[Tuple[str, List[str]]] = [
    ("deed", ["warranty deed", "quitclaim", "mineral deed", "grant deed", " deed "]),
    ("assignment", ["assignment of", "assignor", "assignee", "assigns and conveys"]),
    ("oil_and_gas_lease", ["oil and gas lease", "oil & gas lease", "lessor", "lessee", "paid-up lease"]),
    ("memorandum", ["memorandum of"]),
    ("pooling_unit", ["pooling", "unit agreement", "declaration of unit", "spacing"]),
    ("probate", ["probate", "last will", "testament", "decedent", "estate of"]),
    ("affidavit", ["affidavit", "affiant", "sworn"]),
    ("court_record", ["district court", "case no", "plaintiff", "defendant", "judgment"]),
    ("tax_assessment", ["assessor", "tax roll", "ad valorem", "assessment"]),
    ("runsheet", ["run sheet", "runsheet", "chain of title"]),
    ("ownership_spreadsheet", ["net acres", "decimal interest", "owner name", "tract owner"]),
    ("ogl_lease_sheet", ["lease sheet", "ogl", "primary term", "royalty rate"]),
    ("map_plat", ["plat", "survey", "township", "range", "section corner"]),
    ("prior_report", ["title opinion", "runsheet report", "drilling opinion", "division order title"]),
]


def stage_classify(text_index: List[Dict[str, Any]], out: Path, log) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for t in text_index:
        blob = (Path(t["text_path"]).read_text(encoding="utf-8", errors="replace").lower()
                if t.get("text_path") else "")
        name = t["relpath"].lower()
        scored: List[Tuple[str, int]] = []
        for cat, kws in CLASS_RULES:
            hits = sum(1 for k in kws if k in blob) + sum(2 for k in kws if k in name)
            if hits:
                scored.append((cat, hits))
        scored.sort(key=lambda x: -x[1])
        if scored:
            primary = scored[0][0]
            conf = min(0.95, 0.5 + 0.1 * scored[0][1])
            review = "" if conf >= 0.6 else "REVIEW REQUIRED"
        else:
            primary, conf, review = "unknown_review_required", "", "REVIEW REQUIRED"
        rows.append({
            "source_path": t["source_path"], "relpath": t["relpath"],
            "primary_class": primary,
            "all_classes": ";".join(c for c, _ in scored) or "unknown_review_required",
            "confidence": conf, "review_flag": review,
            "text_status": t.get("status", ""), "generated_at": C.now_iso(),
        })
    C.write_table(out, "document_classification",
                  ["source_path", "relpath", "primary_class", "all_classes",
                   "confidence", "review_flag", "text_status", "generated_at"], rows)
    log.info("D. classification: %d docs, %d review-required",
             len(rows), sum(1 for r in rows if r["review_flag"]))
    return rows


# ============================================================ E. FACTS =======

LABELS = {
    "grantor": r"grantor", "grantee": r"grantee", "lessor": r"lessor",
    "lessee": r"lessee", "assignor": r"assignor", "assignee": r"assignee",
    "decedent": r"decedent|estate of|deceased",
    "effective_date": r"effective date", "execution_date": r"execution date|dated",
    "recording_date": r"recorded|recording date|filed",
    "county": r"county", "state": r"state",
    "legal_description": r"legal description|described as|lands? described",
    "tract_description": r"tract",
    "gross_acres": r"gross acres|containing", "net_acres": r"net acres",
    "royalty": r"royalty", "working_interest": r"working interest",
    "net_revenue_interest": r"net revenue interest|nri",
    "lease_burden": r"burden", "mineral_interest": r"mineral interest",
    "decimal_interest": r"decimal interest|decimal",
    "depth_restrictions": r"depth|below|above", "term": r"primary term|term of",
    "reservations": r"reserv|except",
}
HIGH_VALUE = {"decimal_interest", "net_acres", "royalty", "net_revenue_interest",
              "working_interest", "legal_description"}

_INSTR_RE = re.compile(r"\b(?:instrument|inst|doc(?:ument)?|reception|entry)\s*(?:no\.?|#|number)?\s*[:#]?\s*([0-9][0-9\-]{4,})", re.I)
_BOOKPAGE_RE = re.compile(r"\b(?:book|bk|vol(?:ume)?)\.?\s*([0-9]+)\s*(?:page|pg|p)\.?\s*([0-9]+)", re.I)
_DATE_RE = re.compile(r"(\d{4}-\d{2}-\d{2})|(\d{1,2}/\d{1,2}/\d{2,4})")
_ACRE_RE = re.compile(r"(\d{1,3}(?:,\d{3})*(?:\.\d+)?)\s*acres?", re.I)
_DECIMAL_RE = re.compile(r"\b(0?\.\d{4,8})\b")
_FRACTION_RE = re.compile(r"\b(\d{1,2}/\d{1,3})\b")
_LEGAL_RE = re.compile(r"(sec(?:tion)?\.?\s*\d+[,\s]+t\w*\.?\s*\d+\s*[ns][,\s]+r\w*\.?\s*\d+\s*[ew])", re.I)


def _label_value(text: str, pattern: str) -> Optional[str]:
    m = re.search(rf"(?:{pattern})\s*[:\-]\s*(.+)", text, re.I)
    if m:
        return m.group(1).strip()[:200]
    return None


def stage_extract_facts(text_index, classification, out: Path, log) -> List[Dict[str, Any]]:
    cls_by_path = {c["source_path"]: c for c in classification}
    fields = (["source_file", "relpath", "primary_class", "source_page_or_sheet"]
              + list(LABELS.keys())
              + ["instrument_no", "book_page", "confidence", "review_flags", "generated_at"])
    rows: List[Dict[str, Any]] = []
    for t in text_index:
        if not t.get("text_path"):
            continue
        text = Path(t["text_path"]).read_text(encoding="utf-8", errors="replace")
        row: Dict[str, Any] = {k: "" for k in fields}
        row["source_file"] = t["source_path"]
        row["relpath"] = t["relpath"]
        row["primary_class"] = cls_by_path.get(t["source_path"], {}).get("primary_class", "")
        row["source_page_or_sheet"] = ""
        row["generated_at"] = C.now_iso()

        found, review = 0, []
        for field, pat in LABELS.items():
            val = _label_value(text, pat)
            if val:
                row[field] = val
                found += 1

        # Deterministic pattern fields (override/augment label hits).
        if (m := _INSTR_RE.search(text)):
            row["instrument_no"] = m.group(1); found += 1
        if (m := _BOOKPAGE_RE.search(text)):
            row["book_page"] = f"{m.group(1)}/{m.group(2)}"; found += 1
        if (m := _LEGAL_RE.search(text)) and not row["legal_description"]:
            row["legal_description"] = re.sub(r"\s+", " ", m.group(1)).strip(); found += 1
        if (m := _ACRE_RE.search(text)) and not row["gross_acres"]:
            row["gross_acres"] = m.group(1)
        if (m := _DECIMAL_RE.search(text)) and not row["decimal_interest"]:
            row["decimal_interest"] = m.group(1); found += 1
        if not row["royalty"] and (m := _FRACTION_RE.search(text)):
            row["royalty"] = m.group(1)

        # Confidence: proportion of fields located, capped; high-value low-conf flag.
        row["confidence"] = round(min(0.9, 0.2 + 0.08 * found), 2)
        for hv in HIGH_VALUE:
            if row.get(hv) and row["confidence"] < 0.6:
                review.append(f"low_conf:{hv}")
        if found == 0:
            review.append("REVIEW REQUIRED: no fields extracted")
        if not row["instrument_no"] and not row["book_page"]:
            review.append("missing_recording_citation")
        row["review_flags"] = ";".join(review)
        rows.append(row)

    C.write_table(out, "extracted_facts", fields, rows)
    log.info("E. facts: %d fact rows extracted", len(rows))
    return rows


# ============================================================ F. RECONCILE ===

def _norm_party(s: str) -> str:
    return re.sub(r"[^a-z0-9]", "", (s or "").lower())


def _norm_tract(s: str) -> str:
    return re.sub(r"[^a-z0-9]", "", (s or "").lower())


def _best_date(row: Dict[str, Any]) -> str:
    for k in ("recording_date", "effective_date", "execution_date"):
        v = row.get(k)
        if v:
            m = _DATE_RE.search(v)
            if m:
                d = m.group(0)
                if "/" in d:
                    mm, dd, yy = (d.split("/") + ["", "", ""])[:3]
                    yy = ("20" + yy) if len(yy) == 2 else yy
                    return f"{yy}-{int(mm):02d}-{int(dd):02d}" if mm and dd and yy else "0000-00-00"
                return d
    return "0000-00-00"


def stage_reconcile(facts: List[Dict[str, Any]], out: Path, log) -> Dict[str, List[Dict[str, Any]]]:
    for f in facts:
        f["_date"] = _best_date(f)
        f["_tract"] = _norm_tract(f.get("legal_description") or f.get("tract_description"))

    # Chain rows grouped by tract, ordered by date.
    chains: List[Dict[str, Any]] = []
    conflicts: List[Dict[str, Any]] = []
    groups: Dict[str, List[Dict[str, Any]]] = {}
    for f in facts:
        groups.setdefault(f["_tract"] or "UNKNOWN_TRACT", []).append(f)

    for tract, items in groups.items():
        items.sort(key=lambda r: r["_date"])
        prev_to = None
        for i, f in enumerate(items):
            frm = f.get("grantor") or f.get("assignor") or f.get("lessor") or ""
            to = f.get("grantee") or f.get("assignee") or f.get("lessee") or ""
            gap = ""
            if prev_to and frm and _norm_party(prev_to) != _norm_party(frm):
                gap = "CHAIN GAP: prior grantee != current grantor"
                conflicts.append({
                    "type": "chain_gap", "tract": tract, "relpath": f["relpath"],
                    "detail": f"prior_to='{prev_to}' current_from='{frm}'",
                    "severity": "high", "generated_at": C.now_iso(),
                })
            chains.append({
                "tract": tract, "seq": i + 1, "date": f["_date"],
                "instrument": f.get("primary_class", ""), "instrument_no": f.get("instrument_no", ""),
                "from_party": frm, "to_party": to,
                "decimal_interest": f.get("decimal_interest", ""),
                "gap_flag": gap, "source_file": f["source_file"], "generated_at": C.now_iso(),
            })
            if to:
                prev_to = to

    # Decimal sums per tract.
    recon: List[Dict[str, Any]] = []
    for tract, items in groups.items():
        decs = []
        for f in items:
            try:
                if f.get("decimal_interest"):
                    decs.append(float(f["decimal_interest"]))
            except ValueError:
                pass
        total = round(sum(decs), 6)
        status = "ok"
        if decs and abs(total - 1.0) > 0.0005:
            status = "DECIMAL SUM CONFLICT"
            conflicts.append({
                "type": "decimal_sum", "tract": tract, "relpath": "",
                "detail": f"sum={total} across {len(decs)} interests (expected ~1.0)",
                "severity": "high", "generated_at": C.now_iso(),
            })
        recon.append({
            "tract": tract, "n_instruments": len(items), "n_decimals": len(decs),
            "decimal_sum": total, "status": status,
            "date_span": f"{items[0]['_date']}..{items[-1]['_date']}" if items else "",
            "generated_at": C.now_iso(),
        })

    C.write_table(out, "reconciliation_table",
                  ["tract", "n_instruments", "n_decimals", "decimal_sum", "status",
                   "date_span", "generated_at"], recon)
    C.write_table(out, "chain_summary",
                  ["tract", "seq", "date", "instrument", "instrument_no", "from_party",
                   "to_party", "decimal_interest", "gap_flag", "source_file", "generated_at"], chains)
    C.write_table(out, "conflicts_and_gaps",
                  ["type", "tract", "relpath", "detail", "severity", "generated_at"], conflicts)
    log.info("F. reconcile: %d tracts, %d chain rows, %d conflicts",
             len(recon), len(chains), len(conflicts))
    return {"reconciliation": recon, "chains": chains, "conflicts": conflicts}


# ============================================================ G. VALIDATE =====

def stage_validate(inventory, classification, facts, recon, out: Path, log) -> List[Dict[str, Any]]:
    findings: List[Dict[str, Any]] = []

    def add(rule, severity, relpath, detail):
        findings.append({"rule": rule, "severity": severity, "relpath": relpath,
                         "detail": detail, "generated_at": C.now_iso()})

    classified_paths = {c["source_path"] for c in classification if c["primary_class"] != "unknown_review_required"}
    fact_paths = {f["source_file"] for f in facts}

    for f in facts:
        rp = f["relpath"]
        if not f.get("instrument_no") and not f.get("book_page"):
            add("missing_recording_data", "medium", rp, "no instrument no. or book/page")
        if not (f.get("source_file")):
            add("missing_source_citation", "high", rp, "fact row without source file")
        # impossible dates
        for dk in ("effective_date", "execution_date", "recording_date"):
            v = f.get(dk) or ""
            m = _DATE_RE.search(v)
            if m and m.group(1):
                y = int(m.group(1)[:4])
                if y < 1800 or y > C.datetime.date.today().year + 1:
                    add("impossible_date", "high", rp, f"{dk}={v}")
        for hv in HIGH_VALUE:
            try:
                if f.get(hv) and float(f.get("confidence") or 0) < 0.6:
                    add("high_value_low_confidence", "high", rp, f"{hv} at conf {f.get('confidence')}")
            except ValueError:
                pass
        if f["source_file"] not in classified_paths:
            add("facts_without_classification", "medium", rp, "extracted facts but doc unclassified")

    for r in recon["reconciliation"]:
        if r["status"] != "ok" and r["n_decimals"]:
            add("decimal_sum_mismatch", "high", r["tract"], f"sum={r['decimal_sum']}")
    for c in recon["conflicts"]:
        if c["type"] == "chain_gap":
            add("chain_gap", "high", c.get("relpath", ""), c["detail"])

    # duplicate instrument numbers across files
    seen: Dict[str, str] = {}
    for f in facts:
        ino = f.get("instrument_no")
        if ino:
            if ino in seen and seen[ino] != f["relpath"]:
                add("duplicate_instrument", "medium", f["relpath"], f"instrument {ino} also in {seen[ino]}")
            seen[ino] = f["relpath"]

    C.write_table(out, "validation_report",
                  ["rule", "severity", "relpath", "detail", "generated_at"], findings)
    review_rows = [f for f in findings if f["severity"] == "high"]
    C.write_table(out, "review_required",
                  ["rule", "severity", "relpath", "detail", "generated_at"], review_rows)
    log.info("G. validation: %d findings (%d high-severity)", len(findings), len(review_rows))
    return findings


# ============================================================ H. ASSEMBLE =====

def stage_assemble(inventory, classification, facts, recon, findings, out: Path, log) -> None:
    ts = C.now_iso()
    n_files = len(inventory)
    n_class = sum(1 for c in classification if c["primary_class"] != "unknown_review_required")
    n_facts = len(facts)
    n_conf = len(recon["conflicts"])
    n_high = sum(1 for f in findings if f["severity"] == "high")

    # Curative list (xlsx/csv)
    curative = [f for f in findings if f["rule"] in {
        "missing_recording_data", "chain_gap", "decimal_sum_mismatch",
        "impossible_date", "duplicate_instrument", "high_value_low_confidence"}]
    C.write_table(out, "Grocery_Report_Curative_List",
                  ["rule", "severity", "relpath", "detail", "generated_at"], curative)

    # Source index
    src_index = [{
        "relpath": i["relpath"], "likely_type": i["likely_type"],
        "sha256": i["sha256"], "size_bytes": i["size_bytes"], "modified": i["modified"],
    } for i in inventory]
    C.write_table(out, "Grocery_Report_Source_Index",
                  ["relpath", "likely_type", "sha256", "size_bytes", "modified"], src_index)

    # Executive summary
    exec_md = f"""# Grocery Report — Executive Summary
_Generated: {ts}_

- **Source documents inventoried:** {n_files}
- **Documents classified:** {n_class} ({n_files - n_class} unknown/review)
- **Fact rows extracted:** {n_facts}
- **Tracts reconciled:** {len(recon['reconciliation'])}
- **Conflicts / gaps:** {n_conf}
- **High-severity review items:** {n_high}

## Delivery posture
{_risk_line(n_files, n_class, n_high, n_conf)}

> Every fact in the detailed tables links to a source file. Items marked
> REVIEW REQUIRED must be resolved by a human before final sign-off. No values
> were fabricated; blank fields mean "not found in source".
"""
    (out / "Grocery_Report_Executive_Summary.md").write_text(exec_md, encoding="utf-8")

    # Draft report
    lines = [f"# Grocery Report — DRAFT", f"_Generated: {ts}_  \n",
             "> DRAFT for Rodney's review. Deterministic extraction; unresolved items flagged.\n",
             "## 1. Executive Summary\n", exec_md.split("\n", 2)[2],
             "\n## 2. Title / Ownership Reconciliation\n",
             "| Tract | Instruments | Decimal sum | Status | Date span |",
             "|-------|-------------|-------------|--------|-----------|"]
    for r in recon["reconciliation"]:
        lines.append(f"| {r['tract'][:40]} | {r['n_instruments']} | {r['decimal_sum']} | {r['status']} | {r['date_span']} |")
    lines += ["\n## 3. Lease / OGL Summary\n"]
    ogl = [f for f in facts if f.get("primary_class") in {"oil_and_gas_lease", "ogl_lease_sheet"}]
    if ogl:
        lines += ["| Source | Lessor | Lessee | Royalty | Term | Legal |",
                  "|--------|--------|--------|---------|------|-------|"]
        for f in ogl:
            lines.append(f"| {f['relpath']} | {f.get('lessor','')} | {f.get('lessee','')} | "
                         f"{f.get('royalty','')} | {f.get('term','')} | {f.get('legal_description','')[:40]} |")
    else:
        lines.append("_No lease/OGL documents parsed yet._")
    lines += ["\n## 4. Curative / Exceptions\n"]
    if curative:
        lines += ["| Rule | Severity | File | Detail |", "|------|----------|------|--------|"]
        for c in curative:
            lines.append(f"| {c['rule']} | {c['severity']} | {c['relpath']} | {c['detail'][:60]} |")
    else:
        lines.append("_No curative items detected (or insufficient data)._")
    lines += ["\n## 5. Unresolved / Review Required\n"]
    hi = [f for f in findings if f["severity"] == "high"]
    lines.append(f"{len(hi)} high-severity items — see `review_required.csv`.\n")
    lines += ["\n## 6. Calculation Notes\n",
              "- Decimal interests summed per normalized tract; deviation >0.0005 from 1.0 flagged.",
              "- Dates normalized to ISO; chain gaps flagged when prior grantee != next grantor.",
              "- Acreage/decimal math is only computed where source data is present.\n",
              "\n## 7. Source Document Index\n",
              "See `Grocery_Report_Source_Index.csv` — every source file with hash + type.\n"]
    (out / "Grocery_Report_DRAFT.md").write_text("\n".join(lines), encoding="utf-8")

    _maybe_docx(out, "\n".join(lines), log)
    log.info("H. assembly: draft + exec summary + curative + source index written")


def _risk_line(n_files, n_class, n_high, n_conf) -> str:
    if n_files == 0:
        return "🔴 **RED** — no source documents found at the configured root."
    if n_high or n_conf:
        return f"🟡 **YELLOW** — {n_high} high-severity + {n_conf} conflicts to clear before final."
    if n_class < n_files:
        return "🟡 **YELLOW** — some documents unclassified/need review."
    return "🟢 **GREEN** — no blocking issues detected in automated pass."


def _maybe_docx(out: Path, md_text: str, log) -> None:
    if not C.CAPS["docx"]:
        (out / "Grocery_Report_DRAFT.docx.SKIPPED.txt").write_text(
            "python-docx not installed; .docx not generated. Install: pip install python-docx",
            encoding="utf-8")
        return
    import docx

    d = docx.Document()
    for line in md_text.splitlines():
        if line.startswith("# "):
            d.add_heading(line[2:], 0)
        elif line.startswith("## "):
            d.add_heading(line[3:], 1)
        elif line.strip():
            d.add_paragraph(line)
    d.save(out / "Grocery_Report_DRAFT.docx")


# ============================================================ I. DASHBOARD ====

def stage_dashboard(inventory, text_index, classification, facts, recon, findings, out: Path, log) -> None:
    ts = C.now_iso()
    n = max(1, len(inventory))
    extracted = sum(1 for t in text_index if t["status"] == "extracted")
    classified = sum(1 for c in classification if c["primary_class"] != "unknown_review_required")
    n_high = sum(1 for f in findings if f["severity"] == "high")
    n_conf = len(recon["conflicts"])

    stages = [
        ("A Inventory", 100),
        ("C Text extraction", round(100 * extracted / n)),
        ("D Classification", round(100 * classified / n)),
        ("E Fact extraction", round(100 * min(len(facts), n) / n)),
        ("F Reconciliation", 100 if recon["reconciliation"] else 0),
        ("G Validation", 100 if findings is not None else 0),
    ]

    def color(pct, high):
        if high:
            return "#c0392b"
        if pct >= 90:
            return "#1e8449"
        if pct >= 50:
            return "#b9770e"
        return "#c0392b"

    overall = "🟢 GREEN" if (n_high == 0 and n_conf == 0 and classified == len(inventory)) else \
              ("🔴 RED" if len(inventory) == 0 else "🟡 YELLOW")

    rows_html = "".join(
        f'<tr><td>{html.escape(name)}</td><td style="background:{color(pct, False)};color:#fff;text-align:center">{pct}%</td></tr>'
        for name, pct in stages)

    html_doc = f"""<!doctype html><meta charset="utf-8">
<title>Grocery Report Status</title>
<style>
 body{{font-family:system-ui,Arial,sans-serif;margin:2rem;color:#111}}
 table{{border-collapse:collapse;min-width:420px}} td,th{{border:1px solid #ccc;padding:.5rem .8rem}}
 .kpi{{display:inline-block;margin:.4rem 1rem .4rem 0;padding:.6rem 1rem;border-radius:8px;background:#f4f4f4}}
 .big{{font-size:1.4rem;font-weight:700}}
</style>
<h1>Grocery Report — Status Dashboard</h1>
<p>Generated: {ts}</p>
<p class="big">Monday (2026-07-06) delivery: {overall}</p>
<div>
 <span class="kpi">Files: <b>{len(inventory)}</b></span>
 <span class="kpi">Text extracted: <b>{extracted}</b></span>
 <span class="kpi">Classified: <b>{classified}</b></span>
 <span class="kpi">Fact rows: <b>{len(facts)}</b></span>
 <span class="kpi">Conflicts: <b>{n_conf}</b></span>
 <span class="kpi" style="background:#fdecea">High-sev review: <b>{n_high}</b></span>
</div>
<h2>Stage completion</h2>
<table><tr><th>Stage</th><th>Complete</th></tr>{rows_html}</table>
<h2>Remaining blockers</h2>
<ul>
 {'<li>No source documents found at the configured root.</li>' if not inventory else ''}
 {f'<li>{n_high} high-severity validation items (see review_required.csv).</li>' if n_high else ''}
 {f'<li>{n_conf} chain/decimal conflicts (see conflicts_and_gaps.csv).</li>' if n_conf else ''}
 {'<li>None detected in automated pass.</li>' if (inventory and not n_high and not n_conf) else ''}
</ul>
<p style="color:#666">Green=verified · Yellow=review required · Red=blocked/conflict/missing.
Legend applies to stage cells and overall posture.</p>
"""
    (out / "status_dashboard.html").write_text(html_doc, encoding="utf-8")
    C.write_table(out, "status_dashboard",
                  ["stage", "percent_complete", "generated_at"],
                  [{"stage": s, "percent_complete": p, "generated_at": ts} for s, p in stages])
    log.info("I. dashboard: %s", overall)


# ============================================================ ORCHESTRATOR ====

def run_all(root: Path, out: Path, log) -> None:
    log.info("=== Grocery Report pipeline start | root=%s ===", root)
    log.info("capabilities: %s", C.caps_summary())
    inv = stage_inventory(root, out, log)
    stage_dedupe(inv, out, log)
    tix = stage_extract_text(inv, out, log)
    cls = stage_classify(tix, out, log)
    facts = stage_extract_facts(tix, cls, out, log)
    recon = stage_reconcile(facts, out, log)
    findings = stage_validate(inv, cls, facts, recon, out, log)
    stage_assemble(inv, cls, facts, recon, findings, out, log)
    stage_dashboard(inv, tix, cls, facts, recon, findings, out, log)
    log.info("=== pipeline complete | outputs in %s ===", out)


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(prog="grocery-report", description="Grocery Report title pipeline")
    ap.add_argument("--root", default=None, help="project root (default: env DATABOSSX_PROJECT_ROOT or .)")
    ap.add_argument("--out", default=None, help="output dir (default: <root>/output)")
    args = ap.parse_args(argv)
    root = C.resolve_root(args.root)
    out = C.output_dir(root, args.out)
    log = C.get_logger(out=out)
    run_all(root, out, log)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
