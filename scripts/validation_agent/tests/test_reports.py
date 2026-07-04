"""Report finder scores candidates, output generator never overwrites, and the
report improver labels provenance without fabricating legal facts."""

from pathlib import Path

from config.settings import Settings
from core.run_manager import RunManager
from reports.report_finder import ReportFinder
from reports.report_improver import (build_improved_markdown, LABEL_VERIFIED,
                                     LABEL_ESCALATED, LABEL_UNVERIFIED)
from reports.output_generator import OutputGenerator
from tests.fixtures.make_fixtures import make_report


def _settings(tmp_path):
    s = Settings()
    s.project_root = tmp_path
    s.outputs_dir = tmp_path / "outputs"
    s.extra_report_dirs = []
    return s


def test_report_finder_identifies_best(tmp_path):
    (tmp_path / "reports").mkdir()
    good = make_report(tmp_path / "reports" / "prospect_report.md")
    # A sparse, low-value competitor.
    (tmp_path / "reports" / "notes.txt").write_text("misc", encoding="utf-8")
    finder = ReportFinder(_settings(tmp_path), workbook_stem="prospect")
    best = finder.best()
    assert best is not None
    assert Path(best.path).name == "prospect_report.md"
    assert best.score > 0


def test_output_generator_never_overwrites(tmp_path):
    rm = RunManager(tmp_path / "outputs", timestamp="20260101_000000")
    gen = OutputGenerator(rm, db=None, run_id="r1")
    p1 = gen.write_markdown("# a", name="audit_report.md")
    p2 = gen.write_markdown("# b", name="audit_report.md")
    assert p1 != p2
    assert p1.exists() and p2.exists()


def test_docx_packet_export(tmp_path):
    rm = RunManager(tmp_path / "outputs", timestamp="20260101_000001")
    gen = OutputGenerator(rm, db=None, run_id="r1")
    p = gen.write_docx_packet("Packet", [("Status", "CERTIFIED"),
                                         ("Notes", "line1\nline2")])
    assert p is not None and p.exists() and p.suffix == ".docx"
    # It is a valid docx (a zip openable by python-docx).
    from docx import Document
    doc = Document(str(p))
    assert any("CERTIFIED" in par.text for par in doc.paragraphs)


def test_improver_reads_docx_prior_report(tmp_path):
    from docx import Document
    prior = tmp_path / "prior.docx"
    d = Document()
    d.add_paragraph("Prior examiner note: lease recorded.")
    d.save(str(prior))
    text = build_improved_markdown(
        run_id="r1", mode="dry_run", source_report=prior,
        manifest={"file_name": "wb.xlsx", "sha256": "x", "sheets": [],
                  "is_macro_enabled": False},
        validation_results=[], source_documents=[], escalations=[],
        spend={"cumulative": 0, "cap": 100, "remaining": 100},
        credentials_ok=True, timestamp="2026-01-01T00:00:00Z")
    assert "Prior examiner note: lease recorded." in text
    assert LABEL_UNVERIFIED in text


def test_improver_labels_and_does_not_fabricate(tmp_path):
    src = make_report(tmp_path / "prior.md")
    manifest = {"file_name": "wb.xlsx", "sha256": "abc", "sheets": [],
                "is_macro_enabled": False}
    validation = [{"gate_name": "Acreage Footing Gate", "status": "PASS",
                   "severity": "info", "repairable": False,
                   "affected_subject": "tract acreage"}]
    escalations = [{"category": "Chain Continuity Gate", "severity": "high",
                    "subject": "title chain", "reason": "grantor not vested",
                    "recommended_action": "examiner review"}]
    md = build_improved_markdown(
        run_id="r1", mode="dry_run", source_report=src, manifest=manifest,
        validation_results=validation, source_documents=[],
        escalations=escalations, spend={"cumulative": 0, "cap": 100,
                                        "remaining": 100},
        credentials_ok=False, timestamp="2026-01-01T00:00:00Z")
    assert LABEL_VERIFIED in md
    assert LABEL_ESCALATED in md
    assert LABEL_UNVERIFIED in md
    assert "grantor not vested" in md
    # Fabrication guard: a legal fact never provided must not appear.
    assert "Warranty Deed dated 1887" not in md
    # Prior report content is carried but explicitly labeled unverified.
    assert "Prior Report (preserved verbatim, unverified)" in md
