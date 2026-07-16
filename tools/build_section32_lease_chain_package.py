from __future__ import annotations

import hashlib
import shutil
import textwrap
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
AI_NAME = "GPT-5-6-SOL"
PACKAGE_NAME = f"SECTION32_DIVERSIFIED_LEASE_CHAIN_{AI_NAME}"
OUTPUT_ROOT = ROOT / "artifacts"
PACKAGE_DIR = OUTPUT_ROOT / PACKAGE_NAME
ZIP_PATH = OUTPUT_ROOT / f"{PACKAGE_NAME}.zip"
AS_OF = "2026-07-16"
SUBJECT = "Section 32, Township 11 North, Range 25 West, Beckham County, Oklahoma"
SOURCE_WORKBOOK_NAME = (
    "32-11N-25W_Beckham_Co_Diversified_Cursory_Title_Report_7-15-26_NHE_UPDATED2.xlsx"
)
SOURCE_WORKBOOK_ID = "1lVVqy37tFdY_me1AyMm2AjSE0fDgwLOx"
SOURCE_WORKBOOK_URL = f"https://drive.google.com/file/d/{SOURCE_WORKBOOK_ID}/view"
IMAGE_FOLDER_URL = (
    "https://drive.google.com/drive/folders/1JDmGBF7-3qF4VOa7lJx9tEf8DhsXcex1"
)

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "E7E6E6"
LIGHT_YELLOW = "FFF2CC"
LIGHT_RED = "F4CCCC"
LIGHT_GREEN = "D9EAD3"
WHITE = "FFFFFF"
THIN_GRAY = Side(style="thin", color="B7B7B7")


TRACTS = [
    {
        "tract": "Tract 1",
        "legal": "N/2 NE/4 and SW/4 NE/4",
        "gross": 120,
        "identified": 120,
        "quarter": "NE/4",
    },
    {
        "tract": "Tract 2",
        "legal": "SE/4 NE/4 and E/2 SE/4",
        "gross": 120,
        "identified": 120,
        "quarter": "NE/4 and SE/4",
    },
    {
        "tract": "Tract 3",
        "legal": "NW/4",
        "gross": 160,
        "identified": 80,
        "quarter": "NW/4",
    },
    {
        "tract": "Tract 4",
        "legal": "SW/4",
        "gross": 160,
        "identified": 160,
        "quarter": "SW/4",
    },
    {
        "tract": "Tract 5",
        "legal": "W/2 SE/4",
        "gross": 80,
        "identified": 80,
        "quarter": "SE/4",
    },
]


CURRENT_ASSET = {
    "lease_name": "OK48147.001.1 (record asset identifier; base OGL not identified)",
    "original_lessor": "Unknown — base OGL must be pulled",
    "original_lessee": "Unknown — base OGL must be pulled",
    "current_owner_chain": (
        "Staghorn Resources → Chesapeake → KL CHK SPV LLC → Diversified Production "
        "LLC/OCM Denali → DP Legacy Tapstone LLC → DP Sooner HoldCo LLC → "
        "Diversified Production LLC"
    ),
    "diversified_acquisition": (
        "Latest return into Diversified: DP Sooner HoldCo LLC → Diversified Production "
        "LLC, I-2023-000796, Bk 2400/Pg 551–567, recorded 2023-03-02. "
        "Schedule B identifies OK48147.001.1 at Bk 2400/Pg 566."
    ),
    "book_page": "2400/551–567 (asset at 2400/566)",
    "document_number": "I-2023-000796",
    "recording_date": "2023-03-02",
    "effective_date": "Not stated in reviewed evidence",
    "legal_description": (
        "All of Section 32-11N-25W is the most likely gross legal scope because the "
        "modern records carry all 16 quarter-quarter index ticks. The asset schedule's "
        "legal and interest columns are blank; this is a scope conclusion, not a "
        "conclusion that Diversified owns 640 net leasehold acres."
    ),
    "quarter_section": "All quarters (scope only)",
    "tract": "Tracts 1–5 (scope only)",
    "gross_acres": (
        "640 gross acres, maximum indexed legal scope; 560 unique gross acres are "
        "described by the 11 historical Union/Leede lease references"
    ),
    "net_acres": "Not reasonably supportable from available schedules",
    "lease_type": "Leasehold asset / base oil and gas lease not identified",
    "interest_type": "Leasehold; estate, depth allocation and quantum unproved",
    "confidence": (
        "MEDIUM for current asset identity and conveyance direction; LOW for base-lease "
        "identity, acreage attribution and quantum"
    ),
    "assumptions": (
        "Treats all 16 quarter-quarter index ticks as gross scope only. Does not allocate "
        "the 2022 co-assignee acquisition between Diversified and OCM Denali. Does not "
        "assume that every historical lease in Bk 872/Pg 279 is part of OK48147.001.1."
    ),
}


COMPONENTS = [
    {
        "id": "H-01",
        "name": "Garrett–McKinley Lease",
        "lessor": "Cora B. Garrett and S. A. Garrett",
        "lessee": "C. E. McKinley",
        "date": "1947-11-13",
        "book_page": "63/33",
        "legal": "E/2 SE/4, Section 32-11N-25W",
        "quarter": "SE/4",
        "tract": "Tract 2",
        "gross": 80,
        "source": "Bk 872/Pg 284, Exhibit B (Image 4393)",
        "confidence": "HIGH for historical lease identity/legal; LOW for current Diversified linkage",
        "notes": (
            "Original lease not isolated. Bk 872/Pg 279 assigned only rights from 11,100 "
            "feet to 19,480 feet for this lease."
        ),
    },
    {
        "id": "H-02",
        "name": "Crook–Union Lease",
        "lessor": "Charlie B. Crook and Ima Pearl Crook",
        "lessee": "Union Oil Company of California",
        "date": "1972-02-05",
        "book_page": "264/345",
        "legal": "NE/4, Section 32-11N-25W (also covers 80 acres in Section 28)",
        "quarter": "NE/4",
        "tract": "Tracts 1 and 2",
        "gross": 160,
        "source": "Original lease Image 1197; Bk 872/Pg 283, Exhibit A",
        "confidence": "HIGH for lease terms/legal; LOW for current Diversified linkage",
        "notes": (
            "Instrument 1131; recorded 1972-02-22. Whole lease states 240 acres; 160 acres "
            "are in Section 32. Five-year term commencing 1973-01-13 and thereafter while "
            "producing. Bk 872/Pg 279 excepts the Crook wellbore rights previously "
            "conveyed at Bk 502/Pg 267."
        ),
    },
    {
        "id": "H-03",
        "name": "LeGrand–Union Lease",
        "lessor": "Cora B. LeGrand, a widow",
        "lessee": "Union Oil Company of California",
        "date": "1976-04-29",
        "book_page": "307/31",
        "legal": "W/2 SE/4 and S/2 SW/4, Section 32-11N-25W",
        "quarter": "SE/4 and SW/4",
        "tract": "Tracts 4 and 5",
        "gross": 160,
        "source": "Bk 872/Pg 282, Exhibit A (Image 4391)",
        "confidence": "HIGH for exhibit recital; LOW for current Diversified linkage",
        "notes": "Historical deep-rights schedule; original lease terms/HBP not reviewed.",
    },
    {
        "id": "H-04",
        "name": "LaFortune et al.–Union Lease",
        "lessor": "Gertrude L. LaFortune et al.",
        "lessee": "Union Oil Company of California",
        "date": "1977-10-14",
        "book_page": "340/302",
        "legal": "N/2 SW/4, Section 32-11N-25W",
        "quarter": "SW/4",
        "tract": "Tract 4",
        "gross": 80,
        "source": "Bk 872/Pg 282, Exhibit A (Image 4391)",
        "confidence": "HIGH for exhibit recital; LOW for current Diversified linkage",
        "notes": "Overlaps other N/2 SW/4 leases; acreage must not be added.",
    },
    {
        "id": "H-05",
        "name": "LaFortune Widow–Union Lease",
        "lessor": "Gertrude LaFortune, a widow",
        "lessee": "Union Oil Company of California",
        "date": "1977-10-14",
        "book_page": "340/304",
        "legal": "N/2 SW/4, Section 32-11N-25W",
        "quarter": "SW/4",
        "tract": "Tract 4",
        "gross": 80,
        "source": "Bk 872/Pg 282, Exhibit A (Image 4391)",
        "confidence": "HIGH for exhibit recital; LOW for current Diversified linkage",
        "notes": "Relationship to Bk 340/Pg 302 is unresolved; do not double count.",
    },
    {
        "id": "H-06",
        "name": "Harrison–Union Lease",
        "lessor": "George T. Harrison Jr. and Renata Harrison",
        "lessee": "Union Oil Company of California",
        "date": "1977-09-29",
        "book_page": "340/323",
        "legal": "N/2 NW/4 and N/2 SW/4, Section 32-11N-25W",
        "quarter": "NW/4 and SW/4",
        "tract": "Tracts 3 and 4",
        "gross": 160,
        "source": "Bk 872/Pg 282, Exhibit A (Image 4391)",
        "confidence": "HIGH for exhibit recital; LOW for current Diversified linkage",
        "notes": "Historical deep-rights schedule; original lease terms/HBP not reviewed.",
    },
    {
        "id": "H-07",
        "name": "Leach–Union Lease",
        "lessor": "Carol Seidenbach Leach",
        "lessee": "Union Oil Company of California",
        "date": "1977-09-29",
        "book_page": "342/564",
        "legal": "N/2 SW/4, Section 32-11N-25W",
        "quarter": "SW/4",
        "tract": "Tract 4",
        "gross": 80,
        "source": "Bk 872/Pg 282, Exhibit A (Image 4391)",
        "confidence": "HIGH for exhibit recital; LOW for current Diversified linkage",
        "notes": "Overlaps other N/2 SW/4 leases; acreage must not be added.",
    },
    {
        "id": "H-08",
        "name": "Kruse–Union Lease",
        "lessor": "Cynthia Seidenbach Kruse",
        "lessee": "Union Oil Company of California",
        "date": "1977-11-01",
        "book_page": "342/566",
        "legal": "N/2 SW/4, Section 32-11N-25W",
        "quarter": "SW/4",
        "tract": "Tract 4",
        "gross": 80,
        "source": "Bk 872/Pg 282, Exhibit A (Image 4391)",
        "confidence": "HIGH for exhibit recital; LOW for current Diversified linkage",
        "notes": "Overlaps other N/2 SW/4 leases; acreage must not be added.",
    },
    {
        "id": "H-09",
        "name": "F. L. Randel–Union Lease",
        "lessor": "F. L. Randel",
        "lessee": "Union Oil Company of California",
        "date": "1978-05-16",
        "book_page": "352/719",
        "legal": "NE/4, Section 32-11N-25W",
        "quarter": "NE/4",
        "tract": "Tracts 1 and 2",
        "gross": 160,
        "source": "Bk 872/Pgs 282–283, Exhibit A (Images 4391–4392)",
        "confidence": "HIGH for exhibit recital; LOW for current Diversified linkage",
        "notes": "Original lease terms/HBP not reviewed.",
    },
    {
        "id": "H-10",
        "name": "J. L. Randel–Union Lease",
        "lessor": "J. L. Randel",
        "lessee": "Union Oil Company of California",
        "date": "1978-05-16",
        "book_page": "352/721",
        "legal": "NE/4, Section 32-11N-25W",
        "quarter": "NE/4",
        "tract": "Tracts 1 and 2",
        "gross": 160,
        "source": "Bk 872/Pg 283, Exhibit A (Image 4392)",
        "confidence": "HIGH for exhibit recital; LOW for current Diversified linkage",
        "notes": "Original lease terms/HBP not reviewed.",
    },
    {
        "id": "H-11",
        "name": "Great Western–Union Lease",
        "lessor": "Great Western Oil & Gas, Inc.",
        "lessee": "Union Oil Company of California",
        "date": "1979-03-21",
        "book_page": "376/369; corrected lease 1039/20–23",
        "legal": "NE/4, Section 32-11N-25W",
        "quarter": "NE/4",
        "tract": "Tracts 1 and 2",
        "gross": 160,
        "source": "Bk 872/Pg 283, Exhibit A (Image 4392); county metadata",
        "confidence": "HIGH for exhibit recital; MEDIUM for correction metadata; LOW for current linkage",
        "notes": "Corrected lease metadata recorded 1988-09-16; operative copy not reviewed.",
    },
]


CHAIN = [
    {
        "order": 1,
        "role": "Latest vesting instrument into current named owner",
        "instrument": "Confirmation of Merger / Confirmatory Conveyance",
        "assignor": "DP Sooner HoldCo LLC",
        "assignee": "Diversified Production LLC",
        "book_page": "2400/551–567",
        "document_number": "I-2023-000796",
        "recording_date": "2023-03-02",
        "effective_date": "Not stated in reviewed evidence",
        "legal_or_asset": "Schedule B includes OK48147.001.1 at 2400/566; legal/interest fields blank",
        "evidence": "Source workbook, Runsheet row 3; described as image-derived",
        "confidence": "MEDIUM-HIGH for parties/record/asset identity; LOW for quantum",
        "issue": "Pull full instrument to confirm merger effective date and operative title effect.",
    },
    {
        "order": 2,
        "role": "Prior division / allocation",
        "instrument": "Confirmation of Division / Confirmatory Conveyance",
        "assignor": "DP Legacy Tapstone LLC",
        "assignee": "Diversified ABS VI Upstream LLC and DP Sooner HoldCo LLC",
        "book_page": "2395/415–464",
        "document_number": "I-2022-004838",
        "recording_date": "2022-11-18",
        "effective_date": "Not stated in reviewed evidence",
        "legal_or_asset": "OK48147.001.1 at 2395/462; legal/interest fields blank; all 16 Q-Q index ticks",
        "evidence": "Source workbook, Runsheet row 2; described as image-derived",
        "confidence": "MEDIUM-HIGH for record/asset identity; LOW for allocation",
        "issue": "Schedule does not show what portion went to each co-grantee.",
    },
    {
        "order": 3,
        "role": "Corporate predecessor consolidation",
        "instrument": "Merger",
        "assignor": "Diversified Production LLC and other DP merger entities",
        "assignee": "DP Legacy Tapstone LLC",
        "book_page": "2389/500–580",
        "document_number": "I-2022-003506",
        "recording_date": "2022-08-24",
        "effective_date": "Not stated in reviewed evidence",
        "legal_or_asset": "All Section 32 index ticks",
        "evidence": "County metadata and Section 32 index, transcribed in source workbook",
        "confidence": "MEDIUM for record identity; LOW for asset-specific merger effect",
        "issue": "Merger plan/certificate and survivor schedule not reviewed.",
    },
    {
        "order": 4,
        "role": "Diversified acquisition",
        "instrument": "Assignment",
        "assignor": "KL CHK SPV LLC (local handwriting may read K3 CHK)",
        "assignee": "Diversified Production LLC and OCM Denali Holdings LLC",
        "book_page": "2371/470–494",
        "document_number": "I-2022-000152",
        "recording_date": "2022-01-12",
        "effective_date": "Possible 2021-12-02 instrument date; not independently reconciled",
        "legal_or_asset": "All 16 Q-Q index ticks / long description; schedule and fractions absent",
        "evidence": "County metadata and Section 32 index, transcribed in source workbook",
        "confidence": "MEDIUM for conveyance direction; LOW for acreage and co-assignee allocation",
        "issue": "Do not attribute OCM Denali's reported 48.75% to Diversified.",
    },
    {
        "order": 5,
        "role": "Immediate Chesapeake predecessor",
        "instrument": "Assignment",
        "assignor": "Chesapeake affiliates",
        "assignee": "KL CHK SPV LLC",
        "book_page": "2340/490–519",
        "document_number": "I-2020-004462",
        "recording_date": "2020-12-14",
        "effective_date": "Not stated in reviewed evidence",
        "legal_or_asset": "Section 32 index marks; schedule/estate/quantum not reviewed",
        "evidence": "County metadata and Section 32 index, transcribed in source workbook",
        "confidence": "MEDIUM for parties/record; LOW for asset-specific bridge",
        "issue": "Pull schedule proving OK48147.001.1 moved into KL CHK SPV LLC.",
    },
    {
        "order": 6,
        "role": "Earlier asset-specific predecessor reference",
        "instrument": "Assignment",
        "assignor": "Staghorn Resources LLC",
        "assignee": "Chesapeake",
        "book_page": "1697/236",
        "document_number": "I-2001-004316",
        "recording_date": "Not verified from reviewed evidence",
        "effective_date": "Not verified from reviewed evidence",
        "legal_or_asset": "Cited by the OK48147.001.1 schedule as an assignment, not a base OGL",
        "evidence": "Source workbook OGL open-item note; schedule transcription",
        "confidence": "MEDIUM for asset-reference citation; LOW for operative scope",
        "issue": (
            "Do not substitute the separate Section 32 index assignment at 1697/574, "
            "recorded 2001-06-14. Pull both records and reconcile."
        ),
    },
    {
        "order": 7,
        "role": "First reviewed instrument that defines historical lease acreage",
        "instrument": "Assignment of Oil and Gas Leases",
        "assignor": "Union Oil Company of California",
        "assignee": "Leede Exploration",
        "book_page": "872/279–284",
        "document_number": "1517 (stamped)",
        "recording_date": "1986-03-06",
        "effective_date": "Date of first production",
        "legal_or_asset": (
            "Exhibit A lists ten Section 32 leases from top Morrow to 19,480 feet; "
            "Exhibit B lists Garrett E/2 SE/4 from 11,100 to 19,480 feet. "
            "Crook wellbore rights previously conveyed at 502/267 are excepted."
        ),
        "evidence": "Direct images 4388–4393",
        "confidence": "HIGH for historical leases/legal/depth; LOW for bridge to OK48147.001.1",
        "issue": "The Leede/Staghorn bridge is not proved; this is a historical definition set.",
    },
]


IMAGE_IDS = {
    "1197.jpg": "1XFQldQn-yMS5PYOMXkFp9s3zl-75OJ98",
    "1575.jpg": "13fvCItA8pLaDTwEzp2kiCfKOuiEBx1qe",
    "1940.jpg": "1_H89iSSu7AoLIe8jjzNYXF1FASaIrs_5",
    "1941.jpg": "1ecekNtbmX5PYH6LJLntUzmWgsYrgkPiH",
    "4388.jpg": "1mjKMxQHBWDkdLJoWbcowS4U4B_l_8f6C",
    "4389.jpg": "1betV--AQaJ2h9ZX3oIwXkxwrp9faLhEA",
    "4390.jpg": "1gYwQWEXyNDTu4qxrfO8JyRJaDEt2kZAD",
    "4391.jpg": "1A9l0EDO69rCS9fjRwByX5f3p2ZHtue0v",
    "4392.jpg": "1GuufvrCJDSaR-PKW-aqbIX3W6mYx9Bov",
    "4393.jpg": "18ab8XOnjLkSjSAWgU9Wct5YXQ6R_CjWR",
}


def image_url(image_name: str) -> str:
    return f"https://drive.google.com/file/d/{IMAGE_IDS[image_name]}/view"


def prepare_output() -> None:
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    if PACKAGE_DIR.exists():
        shutil.rmtree(PACKAGE_DIR)
    PACKAGE_DIR.mkdir()
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()


def style_sheet(ws, freeze: str = "A2", filter_table: bool = True) -> None:
    ws.freeze_panes = freeze
    ws.sheet_view.showGridLines = False
    if filter_table and ws.max_row > 1 and ws.max_column > 1:
        ws.auto_filter.ref = ws.dimensions
    for cell in ws[1]:
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(color=WHITE, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(bottom=Side(style="medium", color=BLUE))
    ws.row_dimensions[1].height = 36
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(bottom=THIN_GRAY)
    for col_idx in range(1, ws.max_column + 1):
        values = [
            str(ws.cell(row=row_idx, column=col_idx).value or "")
            for row_idx in range(1, min(ws.max_row, 75) + 1)
        ]
        longest = max((len(line) for value in values for line in value.splitlines()), default=10)
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max(longest + 2, 12), 44)


def add_title_sheet(wb: Workbook) -> None:
    ws = wb.active
    ws.title = "Executive Summary"
    ws.sheet_view.showGridLines = False
    ws.merge_cells("A1:H2")
    title = ws["A1"]
    title.value = "SECTION 32 — DIVERSIFIED LEASE CHAIN"
    title.fill = PatternFill("solid", fgColor=NAVY)
    title.font = Font(color=WHITE, bold=True, size=18)
    title.alignment = Alignment(horizontal="center", vertical="center")
    rows = [
        ("Subject", SUBJECT),
        ("Client", "Diversified Energy"),
        ("Prepared by", AI_NAME),
        ("Cursory as-of", AS_OF),
        ("Confirmed current asset identity", "OK48147.001.1"),
        (
            "Current owner conclusion",
            "Diversified Production LLC is supported as the latest named record owner of "
            "OK48147.001.1 through I-2023-000796, subject to the limitations below.",
        ),
        (
            "Gross acreage conclusion",
            "640 gross acres is the most likely maximum legal scope from all 16 modern "
            "quarter-quarter index ticks. It is NOT Diversified net acreage. The 11 "
            "historical Union/Leede lease references describe 560 unique gross Section "
            "32 acres and overlap heavily.",
        ),
        (
            "Net acreage conclusion",
            "Not reasonably supportable. Modern schedules leave legal/interest columns "
            "blank and the 2022 acquisition names Diversified and OCM Denali as co-assignees.",
        ),
        (
            "Historical component conclusion",
            "Eleven original lease references are directly supported by Bk 872/Pg 279–284, "
            "but none is proved to be the base OGL for OK48147.001.1.",
        ),
        (
            "Overall confidence",
            "MEDIUM for current asset identity and conveyance direction; HIGH for the "
            "historical lease recitals and acreage calls; LOW for tying those leases to "
            "the current asset and for any net-acre/depth allocation.",
        ),
        (
            "Release posture",
            "CURSORY / HOLD FOR SOURCE PULL — suitable as a lease-chain foundation, not a "
            "title opinion and not a working-interest calculation.",
        ),
    ]
    row = 4
    for label, value in rows:
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=8)
        ws.cell(row=row, column=1, value=label)
        ws.cell(row=row, column=2, value=value)
        ws.cell(row=row, column=1).font = Font(bold=True, color=NAVY)
        ws.cell(row=row, column=1).fill = PatternFill("solid", fgColor=LIGHT_BLUE)
        for col in range(1, 9):
            ws.cell(row=row, column=col).alignment = Alignment(vertical="top", wrap_text=True)
            ws.cell(row=row, column=col).border = Border(bottom=THIN_GRAY)
        ws.row_dimensions[row].height = max(30, 15 * (len(value) // 95 + 1))
        row += 1
    ws.column_dimensions["A"].width = 29
    for col in "BCDEFGH":
        ws.column_dimensions[col].width = 15
    ws.freeze_panes = "A4"


def add_table_sheet(wb: Workbook, title: str, headers: list[str], rows: list[list]) -> None:
    ws = wb.create_sheet(title)
    ws.append(headers)
    for row in rows:
        ws.append(row)
    style_sheet(ws)


def build_master_workbook() -> Path:
    wb = Workbook()
    add_title_sheet(wb)

    inventory_headers = [
        "Status",
        "Lease Name",
        "Original Lessor",
        "Original Lessee",
        "Current Owner Chain",
        "Diversified Acquisition",
        "Book/Page",
        "Document Number",
        "Recording Date",
        "Effective Date",
        "Legal Description",
        "Quarter Section",
        "Tract",
        "Gross Acres",
        "Net Acres",
        "Lease Type",
        "Interest Type",
        "Evidence Confidence",
        "Assumptions / Limitations",
    ]
    asset_row = [
        "SUPPORTED CURRENT ASSET; BASE OGL OPEN",
        CURRENT_ASSET["lease_name"],
        CURRENT_ASSET["original_lessor"],
        CURRENT_ASSET["original_lessee"],
        CURRENT_ASSET["current_owner_chain"],
        CURRENT_ASSET["diversified_acquisition"],
        CURRENT_ASSET["book_page"],
        CURRENT_ASSET["document_number"],
        CURRENT_ASSET["recording_date"],
        CURRENT_ASSET["effective_date"],
        CURRENT_ASSET["legal_description"],
        CURRENT_ASSET["quarter_section"],
        CURRENT_ASSET["tract"],
        CURRENT_ASSET["gross_acres"],
        CURRENT_ASSET["net_acres"],
        CURRENT_ASSET["lease_type"],
        CURRENT_ASSET["interest_type"],
        CURRENT_ASSET["confidence"],
        CURRENT_ASSET["assumptions"],
    ]
    add_table_sheet(wb, "Current Lease Inventory", inventory_headers, [asset_row])

    chain_headers = [
        "Backward Order",
        "Chain Role",
        "Instrument",
        "Assignor",
        "Assignee",
        "Book/Page",
        "Document Number",
        "Recording Date",
        "Effective Date",
        "Legal / Asset Definition",
        "Evidence Basis",
        "Confidence",
        "Open Issue",
    ]
    chain_rows = [
        [
            row["order"],
            row["role"],
            row["instrument"],
            row["assignor"],
            row["assignee"],
            row["book_page"],
            row["document_number"],
            row["recording_date"],
            row["effective_date"],
            row["legal_or_asset"],
            row["evidence"],
            row["confidence"],
            row["issue"],
        ]
        for row in CHAIN
    ]
    add_table_sheet(wb, "Backward Chain", chain_headers, chain_rows)

    component_headers = [
        "Review ID",
        "Current Attribution Status",
        "Lease Name",
        "Original Lessor",
        "Original Lessee",
        "Lease Date",
        "Book/Page",
        "Legal Description",
        "Quarter",
        "Tract",
        "Section 32 Gross Acres",
        "Lease Type",
        "Interest Type",
        "Source",
        "Confidence",
        "Notes / Assumptions",
    ]
    component_rows = [
        [
            c["id"],
            "POTENTIAL HISTORICAL COMPONENT — NOT PROVED IN OK48147.001.1",
            c["name"],
            c["lessor"],
            c["lessee"],
            c["date"],
            c["book_page"],
            c["legal"],
            c["quarter"],
            c["tract"],
            c["gross"],
            "Oil and Gas Lease",
            "Historical leasehold; current estate unproved",
            c["source"],
            c["confidence"],
            c["notes"],
        ]
        for c in COMPONENTS
    ]
    add_table_sheet(wb, "Historical Lease Candidates", component_headers, component_rows)

    acreage_headers = [
        "Tract",
        "Control Legal",
        "Quarter",
        "Control Gross Acres",
        "Unique Acres Described by 11 Historical Leases",
        "Uncovered by Historical Definition Set",
        "Diversified Net Acres",
        "Conclusion",
    ]
    acreage_rows = [
        [
            t["tract"],
            t["legal"],
            t["quarter"],
            t["gross"],
            t["identified"],
            t["gross"] - t["identified"],
            "Not supportable",
            (
                "Historical lease schedule coverage only; current attribution unproved."
                if t["identified"] == t["gross"]
                else "Only N/2 NW/4 (80 acres) appears in the 11-lease definition set."
            ),
        ]
        for t in TRACTS
    ]
    acreage_rows.append(
        [
            "TOTAL",
            "Section 32",
            "All quarters",
            640,
            560,
            80,
            "Not supportable",
            "Modern index scope is 640 acres; historical definition set is 560 unique acres.",
        ]
    )
    add_table_sheet(wb, "Acreage Analysis", acreage_headers, acreage_rows)

    qc_headers = ["QC Item", "Result", "Conflict / Finding", "Required Resolution"]
    qc_rows = [
        [
            "Latest Diversified vesting",
            "SUPPORTED",
            "I-2023-000796 / 2400/551–567 names DP Sooner HoldCo → Diversified.",
            "Pull instrument and merger documents to confirm effective date/title effect.",
        ],
        [
            "Asset identity",
            "SUPPORTED",
            "OK48147.001.1 appears at 2395/462 and 2400/566.",
            "Pull both schedules; current source says legal/interest fields are blank.",
        ],
        [
            "Modern gross acreage",
            "CURSORY CONCLUSION",
            "All 16 quarter-quarter ticks support entire-section scope.",
            "Treat 640 as gross scope only, never as net acres or WI.",
        ],
        [
            "Net acres",
            "NOT SUPPORTABLE",
            "No reviewed schedule allocates the 2022 co-assignee acquisition.",
            "Obtain 2371/470–494 schedules and transaction allocation.",
        ],
        [
            "Base OGL",
            "OPEN",
            "1697/236 is an assignment reference, not the base lease.",
            "Pull I-2001-004316 and follow its schedules to the originating OGL.",
        ],
        [
            "1697 citation conflict",
            "FLAGGED",
            "Asset schedule cites 1697/236; separate index record 1697/574 was also used.",
            "Pull and compare both; do not merge them.",
        ],
        [
            "Historical component bridge",
            "OPEN",
            "Bk 872/Pg 279 defines 11 leases, but Leede → Staghorn is not proved.",
            "Pull 1988–2001 operative schedules, especially Leede/Kabala/Staghorn links.",
        ],
        [
            "Depth / wellbore",
            "FLAGGED",
            "872/279 is depth-limited; Crook wellbore was previously carved at 502/267.",
            "Map depth and wellbore exceptions before any WI work.",
        ],
        [
            "Later disposition",
            "OPEN",
            "2024 Teocalli wellbore assignment may affect Section 32; schedule absent.",
            "Pull I-2024-002930 / 2434/751–760 before calling all rights current.",
        ],
        [
            "Potential later acquisition",
            "OPEN",
            "2025 Stephens/Continental → Diversified metadata is countywide/ambiguous.",
            "Pull I-2025 record at 2464/200–217; exclude until Section 32 schedule is proven.",
        ],
    ]
    add_table_sheet(wb, "QC Conflicts", qc_headers, qc_rows)

    source_headers = ["Source Tier", "Source", "Use", "Reliability", "Location"]
    source_rows = [
        [
            "1 — direct image",
            "Bk 872/Pgs 279–284, Images 4388–4393",
            "Historical lease list, granting scope, depths and Crook exception",
            "High for visible text",
            IMAGE_FOLDER_URL,
        ],
        [
            "1 — direct image",
            "Crook OGL, Image 1197",
            "Original lessor/lessee, execution, recording, legal and term",
            "High for visible text",
            image_url("1197.jpg"),
        ],
        [
            "1 — direct image",
            "Crook Release and Agreement, Image 1575",
            "Lease recital, NE/4, HBP claim as of 1979",
            "High for recital; no present HBP conclusion",
            image_url("1575.jpg"),
        ],
        [
            "1 — direct image",
            "Union → Leede Correction Assignment, Images 1940–1941",
            "Crook wellbore/depth carve and reserved ORRI",
            "High for visible text",
            image_url("1940.jpg"),
        ],
        [
            "2 — image-derived workbook transcription",
            SOURCE_WORKBOOK_NAME,
            "Modern asset identity at 2395/462 and 2400/566",
            "Medium-high for cited record identity; source images not re-opened in this sprint",
            SOURCE_WORKBOOK_URL,
        ],
        [
            "3 — county metadata/index transcription",
            SOURCE_WORKBOOK_NAME,
            "Modern predecessor records and recording dates",
            "Medium for record identity; low for schedules/quantum",
            SOURCE_WORKBOOK_URL,
        ],
    ]
    add_table_sheet(wb, "Evidence Hierarchy", source_headers, source_rows)

    path = PACKAGE_DIR / "MASTER_LEASE_CHAIN.xlsx"
    wb.save(path)
    return path


def set_doc_defaults(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def shade_doc_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement("{http://schemas.openxmlformats.org/wml/2006/main}shd")
    shd.set("{http://schemas.openxmlformats.org/wml/2006/main}fill", fill)
    tc_pr.append(shd)


def add_doc_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        shade_doc_cell(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
            run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()


def build_summary_docx() -> Path:
    doc = Document()
    set_doc_defaults(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HORIZON ENERGY\nSECTION 32 DIVERSIFIED LEASE CHAIN")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(23, 54, 93)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{SUBJECT}\nCursory report as of {AS_OF}").italic = True

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "The strongest available evidence supports Diversified Production LLC as the "
        "latest named record owner in the chain for asset identifier OK48147.001.1. "
        "The latest vesting record is I-2023-000796, Bk 2400/Pgs 551–567, recorded "
        "March 2, 2023, from DP Sooner HoldCo LLC to Diversified Production LLC. "
        "Schedule B identifies OK48147.001.1 at Bk 2400/Pg 566."
    )
    doc.add_paragraph(
        "The modern records support a reasonable cursory conclusion that the asset's "
        "maximum gross legal scope is all 640 acres of Section 32 because all 16 "
        "quarter-quarter index boxes are marked. That is not proof that Diversified owns "
        "640 net leasehold acres. The operative schedule fields for legal and interest "
        "are blank, the 2022 acquisition names both Diversified and OCM Denali as "
        "co-assignees, and no allocation schedule was reviewed."
    )
    doc.add_paragraph(
        "Bk 872/Pgs 279–284 directly defines eleven historical oil-and-gas leases and "
        "560 unique gross acres in Section 32. Those leases are strong historical "
        "component candidates, but the missing Leede-to-Staghorn bridge prevents treating "
        "them as proven components of OK48147.001.1."
    )

    doc.add_heading("Diversified Lease Inventory", level=1)
    add_doc_table(
        doc,
        ["Asset / Lease", "Current Owner", "Latest Record", "Gross Scope", "Net Acres", "Confidence"],
        [
            [
                "OK48147.001.1; base OGL unknown",
                "Diversified Production LLC",
                "I-2023-000796; 2400/551–567; recorded 2023-03-02",
                "640 acres maximum indexed scope",
                "Not supportable",
                "Medium current identity; low acreage/quantum",
            ]
        ],
    )

    doc.add_heading("Backward Current-Owner Chain", level=1)
    add_doc_table(
        doc,
        ["Order", "From → To", "Instrument", "Recording", "Key Evidence / Limitation"],
        [
            [
                row["order"],
                f'{row["assignor"]} → {row["assignee"]}',
                row["instrument"],
                f'{row["document_number"]}; {row["book_page"]}; {row["recording_date"]}',
                row["legal_or_asset"],
            ]
            for row in CHAIN[:6]
        ],
    )

    doc.add_heading("Historical Lease Definition Set", level=1)
    doc.add_paragraph(
        "These eleven leases are reviewed historical candidates, not confirmed current "
        "Diversified leases. Acreages overlap and are not additive."
    )
    add_doc_table(
        doc,
        ["Lease", "Original Parties", "Book/Page; Date", "Section 32 Legal", "Gross Acres"],
        [
            [
                c["name"],
                f'{c["lessor"]} → {c["lessee"]}',
                f'{c["book_page"]}; {c["date"]}',
                c["legal"],
                c["gross"],
            ]
            for c in COMPONENTS
        ],
    )

    doc.add_heading("Acreage Summary", level=1)
    add_doc_table(
        doc,
        ["Tract", "Control Legal", "Gross", "Historical Definition Set", "Gap"],
        [
            [
                t["tract"],
                t["legal"],
                t["gross"],
                t["identified"],
                t["gross"] - t["identified"],
            ]
            for t in TRACTS
        ]
        + [["Total", "Section 32", 640, 560, 80]],
    )
    doc.add_paragraph(
        "The 80-acre difference is S/2 NW/4. Modern index marks may include that acreage, "
        "but the eleven leases in Bk 872/Pg 279 do not. No net-acre conclusion is made."
    )

    doc.add_heading("Confidence Notes", level=1)
    for item in [
        "High: visible lease names, legal descriptions and depth limits in Images 1197 and 4388–4393.",
        "Medium-high: cited asset identity at Bk 2395/Pg 462 and Bk 2400/Pg 566, derived from the strongest prior image-reviewed workbook.",
        "Medium: modern parties, document numbers, book/page ranges and recording dates transcribed from county metadata/indexes.",
        "Low: exact base OGL, co-assignee allocation, net leasehold acres, depth ownership and linkage of the eleven historical leases to OK48147.001.1.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Open Issues", level=1)
    for item in [
        "Pull I-2023-000796 and I-2022-004838 with all schedules; confirm effective dates and allocation.",
        "Pull I-2022-000152 and I-2020-004462 schedules; prove the Chesapeake/KL CHK/DP asset bridge.",
        "Pull I-2001-004316 at Bk 1697/Pg 236 and the separate Bk 1697/Pg 574 record; resolve the citation conflict.",
        "Follow the Staghorn asset backward to its base OGL. Do not stop at an all-right-title-and-interest assignment.",
        "Prove or disprove the Leede/Kabala/Staghorn bridge before attaching the eleven historical leases to the modern asset.",
        "Map the Bk 502/Pg 267 Crook wellbore carve and Bk 872/Pg 279 depth limitations.",
        "Pull the 2024 Teocalli and 2025 Stephens/Continental schedules for later disposition/acquisition screening.",
        "Verify lease-by-lease HBP, releases, extensions, depth clauses and production before WI work.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Cursory Reliance Statement", level=1)
    doc.add_paragraph(
        "This package is an evidence-supported cursory lease-chain work product, not a "
        "final title opinion. It intentionally does not calculate working interest, NRI, "
        "revenue ownership or decimal interests. Reasonable scope conclusions are labeled "
        "as such, and unresolved facts are not converted to zeros."
    )
    path = PACKAGE_DIR / "LEASE_SUMMARY.docx"
    doc.save(path)
    return path


def build_change_log() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Lease Review Log"
    ws.append(
        [
            "Review ID",
            "Lease / Asset",
            "Review Result",
            "Correction / Improvement",
            "Evidence Used",
            "Remaining Question",
        ]
    )
    ws.append(
        [
            "CURRENT-01",
            CURRENT_ASSET["lease_name"],
            "Supported current asset identity; base OGL open",
            "Separated gross legal scope (640) from unsupported net acres; rejected use of OCM Denali 48.75% as Diversified WI.",
            "Source workbook: Bk 2395/Pg 462 and Bk 2400/Pg 566; modern index chain",
            "Base OGL, allocation, effective dates, depth and net acres",
        ]
    )
    for c in COMPONENTS:
        correction = "Converted generic workbook recital to a named historical candidate."
        if c["id"] == "H-01":
            correction = "Corrected lessor from generic 'Garretts' to Cora B. Garrett and S. A. Garrett from the exhibit."
        elif c["id"] == "H-02":
            correction = (
                "Confirmed Charlie B. Crook and Ima Pearl Crook, not index variant 'Cook/Charles R.'; "
                "separated 240-acre whole lease from 160 Section 32 acres."
            )
        elif c["id"] in {"H-04", "H-05", "H-07", "H-08"}:
            correction = "Flagged overlapping N/2 SW/4 descriptions; prohibited additive acreage."
        elif c["id"] == "H-11":
            correction = "Added corrected-lease metadata at Bk 1039/Pgs 20–23 without treating metadata as operative proof."
        ws.append(
            [
                c["id"],
                c["name"],
                "Historical lease verified as a recital; current Diversified ownership not proved",
                correction,
                c["source"],
                "Is this lease included in OK48147.001.1 and currently HBP?",
            ]
        )
    style_sheet(ws)

    ws2 = wb.create_sheet("Corrections and Decisions")
    ws2.append(["ID", "Issue", "Prior Risk", "Decision / Correction", "Evidence", "Remaining Work"])
    decisions = [
        (
            "C-01",
            "1697 citation conflict",
            "Conflating 1697/236 and 1697/574",
            "Treat as separate records. Asset schedule cites 1697/236 / I-2001-004316; 1697/574 is a separate index assignment.",
            "Source workbook OGL note and Runsheet",
            "Pull both complete records.",
        ),
        (
            "C-02",
            "Current acreage",
            "Reporting 640 acres as Diversified-owned net acreage",
            "Report 640 only as maximum gross indexed scope; net acres not supportable.",
            "All 16 Q-Q ticks; blank schedule fields",
            "Obtain allocation and net-acre schedule.",
        ),
        (
            "C-03",
            "Historical schedule acreage",
            "Adding overlapping lease acreages",
            "Calculated unique described footprint of 560 acres; 80-acre S/2 NW/4 gap.",
            "Bk 872/Pgs 282–284",
            "Determine whether modern asset includes a separate S/2 NW/4 lease.",
        ),
        (
            "C-04",
            "Crook lease acreage",
            "Using 240 whole-lease acres as Section 32 acreage",
            "Use 160 gross Section 32 acres; remaining 80 acres lie in Section 28.",
            "Original lease Image 1197",
            "None for gross legal; current title remains open.",
        ),
        (
            "C-05",
            "Crook lease identity",
            "Index OCR name conflict",
            "Use image-confirmed Charlie B. Crook and Ima Pearl Crook.",
            "Images 1197 and 1575",
            "None for party names.",
        ),
        (
            "C-06",
            "Depth / wellbore scope",
            "Treating Bk 872/Pg 279 as a full-depth assignment",
            "Captured top-Morrow/11,100-foot to 19,480-foot limits and Crook wellbore exception.",
            "Images 4388–4393 and 1940–1941",
            "Map current depths and later assignments.",
        ),
        (
            "C-07",
            "Historical-to-modern bridge",
            "Assuming Union/Leede leases equal OK48147.001.1",
            "Classified all eleven as candidates only.",
            "No reviewed Leede/Staghorn schedule",
            "Prove chain with operative 1988–2001 assignments.",
        ),
        (
            "C-08",
            "Later transactions",
            "Assuming no later disposition or acquisition",
            "Flagged 2024 Teocalli and 2025 Stephens/Continental records.",
            "Countywide metadata in source workbook",
            "Pull schedules and test Section 32 applicability.",
        ),
    ]
    for decision in decisions:
        ws2.append(decision)
    style_sheet(ws2)
    path = PACKAGE_DIR / "CHANGE_LOG.xlsx"
    wb.save(path)
    return path


def build_evidence_index() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Evidence Index"
    ws.append(
        [
            "Lease / Asset",
            "Source Document",
            "Book/Page",
            "Document Number",
            "Image / PDF Reference",
            "Drive / File Reference",
            "Evidence Tier",
            "Fact Supported",
            "Limitations",
        ]
    )
    modern_sources = [
        (
            "OK48147.001.1",
            "Confirmation of Merger / Confirmatory Conveyance",
            "2400/551–567",
            "I-2023-000796",
            "Schedule B at 2400/566",
            SOURCE_WORKBOOK_URL,
            "Image-derived workbook transcription",
            "DP Sooner HoldCo → Diversified; asset identity",
            "Underlying image not re-opened; legal/interest fields reported blank.",
        ),
        (
            "OK48147.001.1",
            "Confirmation of Division / Confirmatory Conveyance",
            "2395/415–464",
            "I-2022-004838",
            "Exhibit E at 2395/462",
            SOURCE_WORKBOOK_URL,
            "Image-derived workbook transcription",
            "DP Legacy Tapstone → DP Sooner/ABS VI; asset identity",
            "Co-grantee allocation absent.",
        ),
        (
            "OK48147.001.1",
            "Merger",
            "2389/500–580",
            "I-2022-003506",
            "Section 32 index",
            SOURCE_WORKBOOK_URL,
            "Metadata/index",
            "DP merger branch → DP Legacy Tapstone",
            "Plan, certificate and schedule absent.",
        ),
        (
            "OK48147.001.1",
            "Assignment",
            "2371/470–494",
            "I-2022-000152",
            "Section 32 index / long description",
            SOURCE_WORKBOOK_URL,
            "Metadata/index",
            "KL CHK SPV → Diversified and OCM Denali",
            "Fractions and schedules absent.",
        ),
        (
            "OK48147.001.1",
            "Assignment",
            "2340/490–519",
            "I-2020-004462",
            "Section 32 index",
            SOURCE_WORKBOOK_URL,
            "Metadata/index",
            "Chesapeake → KL CHK SPV",
            "Asset-specific schedule absent.",
        ),
        (
            "OK48147.001.1",
            "Assignment reference",
            "1697/236",
            "I-2001-004316",
            "Asset schedule citation",
            SOURCE_WORKBOOK_URL,
            "Schedule transcription",
            "Staghorn → Chesapeake predecessor reference",
            "Recording date, operative legal and base OGL absent.",
        ),
    ]
    for row in modern_sources:
        ws.append(row)

    ws.append(
        [
            "Eleven historical lease candidates",
            "Assignment of Oil and Gas Leases, Union → Leede",
            "872/279–284",
            "1517 (stamped)",
            "Images 4388–4393",
            IMAGE_FOLDER_URL,
            "Direct recorded images",
            "Lease list, legal descriptions, depth scope, Crook exception",
            "No proved bridge to OK48147.001.1.",
        ]
    )
    for c in COMPONENTS:
        image = "4393.jpg" if c["id"] == "H-01" else ("1197.jpg" if c["id"] == "H-02" else "4391.jpg")
        if c["id"] in {"H-09", "H-10", "H-11"}:
            image = "4392.jpg"
        ws.append(
            [
                c["name"],
                "Original OGL" if c["id"] == "H-02" else "Lease recital in Bk 872/Pg 279 exhibit",
                c["book_page"],
                "1131" if c["id"] == "H-02" else "Not shown in reviewed evidence",
                image,
                image_url(image),
                "Direct image",
                f'{c["lessor"]} → {c["lessee"]}; {c["legal"]}; {c["gross"]} Section 32 gross acres',
                "Current Diversified ownership not proved.",
            ]
        )
    ws.append(
        [
            "Crook–Union Lease",
            "Release and Agreement",
            "395/618",
            "11651",
            "Image 1575",
            image_url("1575.jpg"),
            "Direct image",
            "1979 recital of Crook lease and historical Union HBP claim",
            "Not evidence of present HBP.",
        ]
    )
    ws.append(
        [
            "Crook wellbore/depth branch",
            "Correction Assignment, Union → Leede",
            "502/267–268",
            "13302",
            "Images 1940–1941",
            image_url("1940.jpg"),
            "Direct image",
            "Morrow-Springer wellbore assignment and reservations",
            "Historical only; current succession unproved.",
        ]
    )
    style_sheet(ws)

    ws2 = wb.create_sheet("Source File Register")
    ws2.append(["Source", "Identifier", "URL", "Use", "Access / Verification Note"])
    ws2.append(
        [
            SOURCE_WORKBOOK_NAME,
            SOURCE_WORKBOOK_ID,
            SOURCE_WORKBOOK_URL,
            "Strongest prior consolidated workbook and modern-chain citations",
            "Publicly downloaded and read during this sprint.",
        ]
    )
    for name, drive_id in IMAGE_IDS.items():
        ws2.append(
            [
                name,
                drive_id,
                image_url(name),
                "Direct lease/assignment evidence",
                "Publicly downloaded; key images 1197 and 4388–4393 visually reviewed.",
            ]
        )
    style_sheet(ws2)
    path = PACKAGE_DIR / "EVIDENCE_INDEX.xlsx"
    wb.save(path)
    return path


def build_readme() -> Path:
    text = f"""\
    # Section 32 Diversified Lease Chain

    Subject: {SUBJECT}

    Client: Diversified Energy

    Prepared by: {AI_NAME}

    Cursory as-of date: {AS_OF}

    ## Methodology

    This sprint did not restart the title examination. It used the strongest prior
    consolidated workbook, prior image-review work, county-index transcriptions, and
    direct review of the key lease and assignment images. The examination began with the
    latest named conveyance into Diversified Production LLC and worked backward:

    1. I-2023-000796, Bk 2400/Pgs 551–567;
    2. I-2022-004838, Bk 2395/Pgs 415–464;
    3. I-2022-003506, Bk 2389/Pgs 500–580;
    4. I-2022-000152, Bk 2371/Pgs 470–494;
    5. I-2020-004462, Bk 2340/Pgs 490–519; and
    6. I-2001-004316, Bk 1697/Pg 236.

    The last item is still an assignment, not a base OGL. The chain therefore remains
    open. Bk 872/Pgs 279–284 was separately reviewed because it is the first available
    instrument that directly defines historical lease names, acreage and depth scope.
    Its eleven leases were not silently attached to the modern asset because the
    Leede-to-Staghorn bridge is missing.

    ## Main Conclusion

    Diversified Production LLC is reasonably supported as the latest named record owner
    in the chain for asset identifier **OK48147.001.1**. The best supported maximum gross
    legal scope is all **640 acres** of Section 32 based on all 16 quarter-quarter index
    ticks. This is a gross scope conclusion only. It is not a statement that Diversified
    owns 640 net leasehold acres.

    The eleven historical leases listed in Bk 872/Pgs 279–284 describe **560 unique gross
    acres** in Section 32. They overlap heavily and leave S/2 NW/4 (80 acres) outside that
    historical definition set. They are potential components, not confirmed current
    Diversified leases.

    ## Overall Confidence

    - Medium for current asset identity and conveyance direction.
    - High for the visible historical lease names, legal descriptions and depth limits.
    - Low for the base OGL, modern acreage allocation, current depth ownership, net acres,
      and the connection between the eleven historical leases and OK48147.001.1.

    ## Important Limitations

    - The legal and interest columns for OK48147.001.1 are reported blank at 2395/462 and
      2400/566.
    - The 2022 acquisition names Diversified and OCM Denali as co-assignees. No reviewed
      schedule allocates their interests.
    - Bk 1697/Pg 236 and Bk 1697/Pg 574 are separate references and must not be conflated.
    - Bk 872/Pg 279 is depth-limited and excepts previously conveyed Crook wellbore rights.
    - Lease-by-lease HBP, releases, extensions and later dispositions were not proved.
    - No working interest, NRI, revenue ownership or decimal interest was calculated.

    ## Remaining Work Required

    1. Pull the complete modern instruments and schedules at 2400/551, 2395/415,
       2389/500, 2371/470 and 2340/490.
    2. Pull I-2001-004316 at 1697/236 and the separate 1697/574 instrument.
    3. Follow the Staghorn chain to the actual base OGL and prove or disprove the
       Leede/Kabala/Staghorn bridge.
    4. Map all depth and wellbore carve-outs, including 502/267 and the Carlson 7H records.
    5. Screen the 2024 Teocalli disposition and 2025 Stephens/Continental acquisition.
    6. Confirm current HBP and release status for each proved base lease before WI work.

    ## File Guide

    - `MASTER_LEASE_CHAIN.xlsx` — current asset conclusion, backward chain, historical
      candidates, acreage analysis and QC conflicts.
    - `LEASE_SUMMARY.docx` — concise reviewer-facing narrative and tables.
    - `CHANGE_LOG.xlsx` — lease-by-lease review and corrections.
    - `EVIDENCE_INDEX.xlsx` — source hierarchy and Drive references.
    - `README.md` — methodology, confidence and remaining work.

    ## Cursory Standard

    This package is an evidence-supported cursory lease-chain work product, not a final
    title opinion. “Not supportable” means the evidence does not justify a number; it
    does not mean zero.
    """
    path = PACKAGE_DIR / "README.md"
    path.write_text(textwrap.dedent(text), encoding="utf-8")
    return path


def qc_package(files: list[Path]) -> None:
    expected = {
        "MASTER_LEASE_CHAIN.xlsx",
        "LEASE_SUMMARY.docx",
        "CHANGE_LOG.xlsx",
        "EVIDENCE_INDEX.xlsx",
        "README.md",
    }
    actual = {path.name for path in files}
    if actual != expected:
        raise AssertionError(f"Deliverable mismatch: expected {expected}, got {actual}")
    for path in files:
        if not path.exists() or path.stat().st_size == 0:
            raise AssertionError(f"Missing or empty deliverable: {path}")

    master = load_workbook(PACKAGE_DIR / "MASTER_LEASE_CHAIN.xlsx", data_only=False)
    required_sheets = {
        "Executive Summary",
        "Current Lease Inventory",
        "Backward Chain",
        "Historical Lease Candidates",
        "Acreage Analysis",
        "QC Conflicts",
        "Evidence Hierarchy",
    }
    if not required_sheets.issubset(master.sheetnames):
        raise AssertionError("Master workbook is missing required sheets")
    if master["Current Lease Inventory"].max_row != 2:
        raise AssertionError("Current inventory must contain exactly one supported current asset")
    if master["Backward Chain"].max_row != len(CHAIN) + 1:
        raise AssertionError("Backward chain row count mismatch")
    if master["Historical Lease Candidates"].max_row != len(COMPONENTS) + 1:
        raise AssertionError("Historical candidate row count mismatch")
    if sum(t["gross"] for t in TRACTS) != 640:
        raise AssertionError("Tract control acreage does not total 640")
    if sum(t["identified"] for t in TRACTS) != 560:
        raise AssertionError("Unique historical definition acreage does not total 560")

    change_log = load_workbook(PACKAGE_DIR / "CHANGE_LOG.xlsx", read_only=True)
    if change_log["Lease Review Log"].max_row != len(COMPONENTS) + 2:
        raise AssertionError("Change log does not include every reviewed lease/asset")

    evidence = load_workbook(PACKAGE_DIR / "EVIDENCE_INDEX.xlsx", read_only=True)
    if evidence["Evidence Index"].max_row < 20:
        raise AssertionError("Evidence index is unexpectedly sparse")

    doc = Document(PACKAGE_DIR / "LEASE_SUMMARY.docx")
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for phrase in ["OK48147.001.1", "I-2023-000796", "not a final title opinion"]:
        if phrase.lower() not in full_text.lower():
            raise AssertionError(f"Summary is missing required phrase: {phrase}")


def package_zip(files: list[Path]) -> None:
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(files):
            archive.write(path, arcname=f"{PACKAGE_NAME}/{path.name}")
    with zipfile.ZipFile(ZIP_PATH) as archive:
        names = archive.namelist()
        if len(names) != 5:
            raise AssertionError(f"ZIP contains {len(names)} entries, expected 5")
        if any(not name.startswith(f"{PACKAGE_NAME}/") for name in names):
            raise AssertionError("ZIP must contain exactly one top-level folder")
        bad = archive.testzip()
        if bad:
            raise AssertionError(f"Corrupt ZIP member: {bad}")


def write_receipt(files: list[Path]) -> None:
    receipt = OUTPUT_ROOT / f"{PACKAGE_NAME}_SHA256.txt"
    lines = []
    for path in files + [ZIP_PATH]:
        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        lines.append(f"{digest}  {path.name}")
    receipt.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    prepare_output()
    files = [
        build_master_workbook(),
        build_summary_docx(),
        build_change_log(),
        build_evidence_index(),
        build_readme(),
    ]
    qc_package(files)
    package_zip(files)
    write_receipt(files)
    print(f"Built: {PACKAGE_DIR}")
    print(f"ZIP: {ZIP_PATH}")
    print(f"ZIP bytes: {ZIP_PATH.stat().st_size}")


if __name__ == "__main__":
    main()
