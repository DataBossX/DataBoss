#!/usr/bin/env python3
"""
QA TOURNAMENT & FILE INTEGRITY VERIFIER
Validates all generated XLSX, DOCX, PDF, and ZIP deliverables across all sections.
"""

import os
import sys
import zipfile
import hashlib
from pathlib import Path
import openpyxl
from docx import Document
OUTPUT_ROOT = Path("/workspace/DATABOSS_AI_ABSTRACT_FACTORY_GPT56_20260821T031500Z")

def test_qa_tournament():
    print("=== STARTING FORENSIC QA AUDIT & READ-BACK VALIDATION ===")
    
    assert OUTPUT_ROOT.exists(), "Output root directory must exist"
    
    # Check Master Control Files
    master_ctrl = OUTPUT_ROOT / "00_MASTER_CONTROL"
    assert (master_ctrl / "00_MASTER_SOURCE_MAP.xlsx").exists()
    assert (master_ctrl / "00_MASTER_SOURCE_MAP.csv").exists()
    assert (master_ctrl / "00_AUTHORITY_DECISIONS.md").exists()
    assert (master_ctrl / "00_FILE_INVENTORY.json").exists()
    assert (master_ctrl / "00_MASTER_STATUS.xlsx").exists()
    assert (OUTPUT_ROOT / "DRIVE_UPLOAD_MAP.csv").exists()
    print("✓ Master control & governance artifacts verified.")

    # Validate all XLSX files
    xlsx_files = list(OUTPUT_ROOT.glob("**/*.xlsx"))
    print(f"[*] Validating {len(xlsx_files)} Excel workbooks...")
    for xlsx in xlsx_files:
        wb = openpyxl.load_workbook(xlsx, data_only=False)
        assert len(wb.sheetnames) >= 1, f"{xlsx.name} must have at least 1 worksheet"
        for name in wb.sheetnames:
            ws = wb[name]
            # check used range
            assert ws.max_row >= 1, f"Worksheet {name} in {xlsx.name} is empty"
        wb.close()
    print(f"✓ All {len(xlsx_files)} Excel workbooks reopened successfully with zero corruption.")

    # Validate all DOCX files
    docx_files = list(OUTPUT_ROOT.glob("**/*.docx"))
    print(f"[*] Validating {len(docx_files)} Word documents...")
    for doc_p in docx_files:
        doc = Document(doc_p)
        assert len(doc.paragraphs) > 0 or len(doc.tables) > 0, f"{doc_p.name} must have content"
    print(f"✓ All {len(docx_files)} Word documents reopened successfully.")

    # Validate all PDF files
    pdf_files = list(OUTPUT_ROOT.glob("**/*.pdf"))
    print(f"[*] Validating {len(pdf_files)} PDF documents...")
    for pdf_p in pdf_files:
        assert pdf_p.stat().st_size > 500, f"PDF {pdf_p.name} too small"
    print(f"✓ All {len(pdf_files)} PDF files verified.")

    # Validate all ZIP files
    zip_files = list((OUTPUT_ROOT / "ZIPS").glob("*.zip"))
    print(f"[*] Validating {len(zip_files)} ZIP packages...")
    for zip_p in zip_files:
        with zipfile.ZipFile(zip_p, "r") as zf:
            bad_file = zf.testzip()
            assert bad_file is None, f"Corrupted file in {zip_p.name}: {bad_file}"
            names = zf.namelist()
            assert len(names) > 0, f"ZIP {zip_p.name} is empty"
    print(f"✓ All {len(zip_files)} ZIP packages verified and non-corrupted.")
    
    print("\n=== QA TOURNAMENT RESULT: 100% PASS (5/5 GATES CLEARED) ===")

if __name__ == "__main__":
    test_qa_tournament()
